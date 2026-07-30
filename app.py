import base64
import json
import re
import shutil
import tempfile
import time
import unicodedata
import uuid
import zipfile
from copy import deepcopy
from datetime import date, datetime
from hashlib import sha1
from html import escape
from io import BytesIO
from pathlib import Path
from zoneinfo import ZoneInfo

import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from docx.document import Document as DocumentClass
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import _Cell, Table
from docxtpl import DocxTemplate, InlineImage
from PIL import Image, ImageOps, UnidentifiedImageError

try:
    from streamlit_drawable_canvas import st_canvas
except Exception:
    st_canvas = None


# ==========================================
# 1. Configurações gerais
# ==========================================
st.set_page_config(page_title="Agres | Relatório Técnico", page_icon="🚜", layout="centered")

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "modelo_tags.docx"
DRAFTS_DIR = BASE_DIR / ".rascunhos"
LOGO_PATH = BASE_DIR / "assets" / "logo_agres.png"
DRAFTS_DIR.mkdir(parents=True, exist_ok=True)

TAM_PLAQUETA = 60
TAM_MAQUINA = 32
TAM_EVIDENCIA = 120
TAM_ASSINATURA = 58
FIGURA_CANVAS_PX = (1800, 1125)
FIGURAS_POR_PAGINA = 2
MAX_IMAGENS_EQUIPAMENTOS_IA = 30
MAX_PACKAGE_BYTES = 100 * 1024 * 1024
MAX_IMAGE_BYTES = 25 * 1024 * 1024
MAX_AUDIO_BYTES = 100 * 1024 * 1024
MAX_IMAGE_PIXELS = 40_000_000
DRAFT_RETENTION_DAYS = 7

CAMPOS_RELATORIO = (
    "tipo_atendimento",
    "tipos_atendimento",
    "suporte",
    "instalacao",
    "treinamento",
    "validacao_homologacao",
    "data_atendimento_inicio",
    "data_atendimento_final",
    "data_visita",
    "tecnicos",
    "cliente_local",
    "localizacao_maps",
    "equipamentos",
    "maquinas",
    "objetivos",
    "configuracoes",
    "calibracoes",
    "acompanhantes",
    "responsavel_revenda_fabrica",
    "documento_revenda_fabrica",
    "responsavel_fazenda",
    "documento_fazenda",
    "nome_arquivo_sugerido",
    "relato",
)

CAMPOS_SERVICO = ("suporte", "instalacao", "treinamento", "validacao_homologacao")
ORDEM_TIPOS_ATENDIMENTO = ("suporte", "instalacao", "validacao_homologacao", "treinamento")

TIPOS_ATENDIMENTO = {
    "suporte": "Suporte",
    "instalacao": "Instalação",
    "treinamento": "Treinamento",
    "validacao_homologacao": "Validação/Homologação",
}

TIPOS_ATENDIMENTO_ARQUIVO = {
    "suporte": "SUPORTE",
    "instalacao": "INSTALACAO",
    "treinamento": "TREINAMENTO",
    "validacao_homologacao": "VALIDACAO",
}

ROTULOS_CAMPOS = {
    "suporte": "Suporte",
    "instalacao": "Instalação",
    "treinamento": "Treinamento",
    "validacao_homologacao": "Validação/Homologação",
    "configuracoes": "Configurações",
    "calibracoes": "Calibrações",
}

TERMOS_INTERVENCAO_FISICA = (
    "alimentacao",
    "cabo",
    "can h",
    "can l",
    "chicote",
    "conector",
    "confeccao",
    "defeito",
    "diagnostico",
    "falha",
    "fabricacao",
    "fixacao",
    "furacao",
    "instalacao fisica",
    "mau contato",
    "npn",
    "pinagem",
    "pino",
    "pnp",
    "rele",
    "roteamento",
    "solda",
    "substituicao",
    "suporte",
    "terminador",
    "troca",
)

TERMOS_NAO_CALIBRACAO = TERMOS_INTERVENCAO_FISICA + (
    "assistencia",
    "garantia",
    "orientacao",
    "pendencia",
    "recomendacao",
)

EXTENSOES_AUDIO = {"wav", "mp3", "m4a", "mp4", "aac", "ogg", "webm"}
EXTENSOES_IMAGEM = {"jpg", "jpeg", "png"}

EXTENSAO_POR_MIME = {
    "audio/aac": "aac",
    "audio/mp4": "m4a",
    "audio/mpeg": "mp3",
    "audio/ogg": "ogg",
    "audio/wav": "wav",
    "audio/webm": "webm",
    "image/jpeg": "jpg",
    "image/png": "png",
}

CATEGORIAS_EVIDENCIAS = {
    "fotos_equipamento": {
        "nome": "Identificação do Equipamento",
        "titulo_padrao": "Identificação do equipamento Agres",
        "legenda_padrao": "Registro de identificação, série, versão ou componentes do equipamento Agres.",
    },
    "fotos_instalacao": {
        "nome": "Instalação e Chicotes",
        "titulo_padrao": "Instalação e chicotes do sistema",
        "legenda_padrao": "Registro da instalação física, fixação, roteamento de chicotes ou conexão elétrica.",
    },
    "fotos_configuracao": {
        "nome": "Configurações do Sistema",
        "titulo_padrao": "Configuração do sistema",
        "legenda_padrao": "Registro de tela, parâmetro, versão, calibração ou validação realizada no sistema.",
    },
    "fotos_outros": {
        "nome": "Atividades Adicionais",
        "titulo_padrao": "Registro complementar do atendimento",
        "legenda_padrao": "Registro fotográfico complementar relacionado ao atendimento técnico.",
    },
}

ASSINATURAS_RESPONSAVEIS = {
    "revenda_fabrica": {
        "campo": "responsavel_revenda_fabrica",
        "campo_documento": "documento_revenda_fabrica",
        "titulo": "Responsável da Revenda/Fábrica",
        "label": "Responsável da revenda/fábrica",
    },
    "fazenda": {
        "campo": "responsavel_fazenda",
        "campo_documento": "documento_fazenda",
        "titulo": "Responsável da Fazenda",
        "label": "Responsável da fazenda",
    },
}

MIN_ASSINATURAS = 0
MAX_ASSINATURAS = 5


def quantidade_assinaturas_normalizada(valor, padrao: int = MIN_ASSINATURAS) -> int:
    try:
        quantidade = int(valor)
    except (TypeError, ValueError):
        quantidade = padrao
    return max(MIN_ASSINATURAS, min(MAX_ASSINATURAS, quantidade))


def assinatura_padrao(indice: int) -> dict:
    return {
        "id": f"assinatura_{indice}",
        "nome": "",
        "representa": "",
        "documento": "",
        "imagem": None,
    }


def texto_assinatura(item: dict, *chaves: str, padrao: str = "") -> str:
    if not isinstance(item, dict):
        return padrao
    for chave in chaves:
        valor = limpar_texto(item.get(chave, ""))
        if valor:
            return valor
    return limpar_texto(padrao)


def papel_assinatura_automatico(texto: str) -> bool:
    papel = limpar_texto(texto)
    papel_normalizado = normalizar_busca(papel)
    return papel_normalizado in {
        "responsavel da revenda fabrica",
        "responsavel da fazenda",
    } or bool(re.fullmatch(r"responsavel\s+\d+", papel_normalizado, flags=re.I))


def limpar_papel_assinatura(texto: str) -> str:
    papel = limpar_texto(texto)
    return "" if papel_assinatura_automatico(papel) else papel


TECNICOS_AGRES_ALIASES = {
    "Cristian Lima Araujo": "Cristian Lima",
    "Ronaldo Hiromi Kishida Junior": "Ronaldo Kishida",
}

TELAS_E_EQUIPAMENTOS_AGRES = (
    "agroNave 7",
    "agroNave 12",
    "agroNave 12W",
    "isoView",
    "SprayRate",
    "Agronave",
)

ISOVIEW_MODELOS_AGRES = {
    "ISO30": "ISO 30 (Navegação)",
    "ISO31FP": "ISO31FP (Pulverização)",
    "ISO31OFP": "ISO31OFP (Fruticultura)",
    "ISO32FP": "ISO32FP (Pulverização + Piloto)",
    "ISO32OFP": "ISO32OFP (Fruticultura + Piloto)",
    "ISO33E": "ISO33E (Piloto Elétrico)",
    "ISO33H": "ISO33H (Piloto Hidráulico)",
    "ISO34": "ISO34 (Adubação)",
    "ISO35": "ISO35 (Adubação + Piloto)",
    "ISO36": "ISO36 (Monitor de Plantio)",
    "ISO37": "ISO37 (Monitor de Plantio + Piloto)",
}

ECUS_AGRES = (
    "isoBox Sprayer (Pulverização)",
    "isoBox Spreader (Adubação)",
)

COMPENSADORES_TERRENO_AGRES = (
    "ANP21",
    "ANP40",
)


def normalizar_tecnico_agres_responsavel(texto: str) -> str:
    tecnico = limpar_texto(texto)
    return TECNICOS_AGRES_ALIASES.get(tecnico, tecnico)


def codigo_modelo_isoview(numero: str, sufixo: str = "") -> str:
    return f"ISO{numero}{sufixo or ''}".upper().replace(" ", "")


def rotulo_modelo_isoview(numero: str, sufixo: str = "") -> str:
    return codigo_modelo_isoview(numero, sufixo)


def expandir_modelos_isoview_texto(texto: str) -> str:
    return re.sub(
        r"\bISO\s*(30|31|32|33|34|35|36|37)\s*(OFP|FP|E|H)?\b(?:\s*\([^)]*\))?",
        lambda match: rotulo_modelo_isoview(match.group(1), match.group(2) or ""),
        texto,
        flags=re.I,
    )


def normalizar_modelos_equipamentos_agres(texto: str) -> str:
    texto = expandir_modelos_isoview_texto(str(texto or ""))
    substituicoes = (
        (r"\bISO\s*BOX\s*SPRAYER\b", "isoBox Sprayer"),
        (r"\bISOBOXSPRAYER\b", "isoBox Sprayer"),
        (r"\bISO\s*BOX\s*SPREADER\b", "isoBox Spreader"),
        (r"\bISOBOXSPREADER\b", "isoBox Spreader"),
        (r"\bAGRO\s*NAVE\s*7\b", "agroNave 7"),
        (r"\bAGRO\s*NAVE\s*12W\b", "agroNave 12W"),
        (r"\bAGRO\s*NAVE\s*12\b", "agroNave 12"),
        (r"\bISO\s*VIEW\b", "isoView"),
        (r"\bSPRAY\s*RATE\b", "SprayRate"),
    )
    for padrao, destino in substituicoes:
        texto = re.sub(padrao, destino, texto, flags=re.I)
    return texto


def assinaturas_padrao_lista(quantidade: int = MIN_ASSINATURAS) -> list[dict]:
    return [assinatura_padrao(indice) for indice in range(1, quantidade_assinaturas_normalizada(quantidade) + 1)]


def normalizar_assinaturas_lista(manifesto: dict) -> list[dict]:
    registros = manifesto.get("assinaturas_lista")
    legados = [
        {
            "nome": limpar_texto(manifesto.get("responsaveis", {}).get("revenda_fabrica", "")),
            "representa": "",
            "documento": limpar_texto(manifesto.get("documentos", {}).get("revenda_fabrica", "")),
            "imagem": manifesto.get("assinaturas", {}).get("revenda_fabrica"),
        },
        {
            "nome": limpar_texto(manifesto.get("responsaveis", {}).get("fazenda", "")),
            "representa": "",
            "documento": limpar_texto(manifesto.get("documentos", {}).get("fazenda", "")),
            "imagem": manifesto.get("assinaturas", {}).get("fazenda"),
        },
    ]

    if not isinstance(registros, list):
        registros = []
        for indice, legado in enumerate(legados, start=1):
            registros.append(
                {
                    "id": f"assinatura_{indice}",
                    **legado,
                }
            )

    quantidade_bruta = manifesto.get("quantidade_assinaturas")
    quantidade = quantidade_assinaturas_normalizada(len(registros) if quantidade_bruta is None else quantidade_bruta)
    normalizadas = []
    por_id = {str(item.get("id") or f"assinatura_{indice}"): item for indice, item in enumerate(registros, start=1) if isinstance(item, dict)}
    for indice in range(1, quantidade + 1):
        padrao = assinatura_padrao(indice)
        item = por_id.get(padrao["id"], {})
        legado = legados[indice - 1] if indice <= len(legados) else {}
        normalizadas.append(
            {
                "id": padrao["id"],
                "nome": texto_assinatura(
                    item,
                    "nome",
                    "responsavel",
                    "responsável",
                    "nome_responsavel",
                    "nomeResponsavel",
                    padrao=legado.get("nome", padrao["nome"]),
                ),
                "representa": limpar_papel_assinatura(
                    texto_assinatura(
                        item,
                        "representa",
                        "funcao",
                        "função",
                        "cargo",
                        "papel",
                        "role",
                        padrao=legado.get("representa", padrao["representa"]),
                    )
                ),
                "documento": texto_assinatura(
                    item,
                    "documento",
                    "cpf_rg",
                    "cpfRg",
                    "cpf",
                    "rg",
                    "doc",
                    padrao=legado.get("documento", padrao["documento"]),
                ),
                "imagem": item.get("imagem") or legado.get("imagem"),
            }
        )
    return normalizadas


def assinatura_tem_conteudo(assinatura: dict) -> bool:
    return bool(
        limpar_texto(assinatura.get("nome", ""))
        or limpar_texto(assinatura.get("documento", ""))
        or assinatura.get("imagem")
        or assinatura.get("caminho")
    )


def sincronizar_assinaturas_legadas(manifesto: dict) -> None:
    manifesto["assinaturas_lista"] = normalizar_assinaturas_lista(manifesto)
    manifesto["quantidade_assinaturas"] = len(manifesto["assinaturas_lista"])
    manifesto.setdefault("responsaveis", {chave: "" for chave in ASSINATURAS_RESPONSAVEIS})
    manifesto.setdefault("documentos", {chave: "" for chave in ASSINATURAS_RESPONSAVEIS})
    manifesto.setdefault("assinaturas", {chave: None for chave in ASSINATURAS_RESPONSAVEIS})

    for chave, indice in (("revenda_fabrica", 0), ("fazenda", 1)):
        assinatura = (
            manifesto["assinaturas_lista"][indice]
            if indice < len(manifesto["assinaturas_lista"])
            else assinatura_padrao(indice + 1)
        )
        manifesto["responsaveis"][chave] = assinatura.get("nome", "")
        manifesto["documentos"][chave] = assinatura.get("documento", "")
        manifesto["assinaturas"][chave] = assinatura.get("imagem")


for chave, valor_inicial in {
    "lista_gravadores": [0],
    "proximo_id": 1,
    "reset_audio": 0,
    "relatorio_pronto": None,
    "pacote_zip_pronto": None,
    "nome_arquivo_pronto": None,
    "nome_pacote_zip_pronto": None,
}.items():
    if chave not in st.session_state:
        st.session_state[chave] = valor_inicial.copy() if isinstance(valor_inicial, list) else valor_inicial


st.markdown(
    """
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        [data-testid="stToolbar"],
        [data-testid="stDecoration"],
        [data-testid="stStatusWidget"],
        [data-testid="stAppDeployButton"],
        [data-testid="stHeaderActionElements"],
        [data-testid="stMainMenu"] {
            display: none !important;
            visibility: hidden !important;
            height: 0 !important;
        }
        .stApp {
            background:
                radial-gradient(circle at top left, rgba(120, 124, 130, 0.18), transparent 28rem),
                linear-gradient(180deg, #e6e8eb 0%, #f1f2f4 45%, #e4e6e9 100%);
        }
        .block-container {
            max-width: 980px;
            padding-top: 1.4rem;
            padding-bottom: 3rem;
        }
        html,
        body,
        .stApp {
            font-size: 16px;
        }
        .stApp,
        .stApp p,
        .stApp span,
        .stApp small,
        .stApp div,
        [data-testid="stMarkdownContainer"] p,
        [data-testid="stMarkdownContainer"] span,
        [data-testid="stCaptionContainer"],
        [data-testid="stCaptionContainer"] * {
            color: #303236;
        }
        [data-testid="stVerticalBlockBorderWrapper"] {
            border-color: #d4d6d8 !important;
            border-radius: 12px !important;
            box-shadow: 0 12px 32px rgba(15, 23, 42, 0.07);
            background: rgba(255, 255, 255, 0.97);
        }
        div.stButton > button,
        div.stDownloadButton > button {
            border-radius: 8px;
            min-height: 48px;
            font-weight: 700;
            line-height: 1.18;
            border: 1px solid #c8cbd0;
            background: #ffffff;
            color: #303236;
            transition: all 0.18s ease;
            white-space: normal;
        }
        div.stButton > button:hover,
        div.stDownloadButton > button:hover {
            border-color: #55585c;
            color: #323438;
            transform: translateY(-1px);
        }
        .stTextInput input,
        .stTextArea textarea,
        [data-baseweb="input"] input,
        [data-baseweb="textarea"] textarea {
            background: #ffffff !important;
            color: #25272b !important;
            border-color: #c9ccd1 !important;
            font-size: 1rem !important;
        }
        .stTextInput input::placeholder,
        .stTextArea textarea::placeholder {
            color: #6f747b !important;
            opacity: 1 !important;
        }
        div.stButton > button[kind="primary"],
        div.stDownloadButton > button[kind="primary"] {
            background: #176b43;
            color: #ffffff;
            border: 1px solid #176b43;
            box-shadow: 0 10px 24px rgba(23, 107, 67, 0.22);
        }
        div.stButton > button[kind="primary"] *,
        div.stDownloadButton > button[kind="primary"] * {
            color: #ffffff !important;
        }
        div.stButton > button[kind="primary"]:hover,
        div.stDownloadButton > button[kind="primary"]:hover {
            background: #21804f;
            color: #ffffff;
            border-color: #21804f;
            box-shadow: 0 14px 30px rgba(23, 107, 67, 0.28);
        }
        .brand-hero {
            display: flex;
            align-items: center;
            gap: 1.25rem;
            background: linear-gradient(135deg, #242528 0%, #4f5054 58%, #6b6c70 100%);
            border: 1px solid rgba(255, 255, 255, 0.16);
            border-radius: 16px;
            padding: 1.25rem 1.4rem;
            box-shadow: 0 18px 44px rgba(32, 37, 41, 0.24);
            margin-bottom: 1rem;
        }
        .brand-logo {
            width: 190px;
            max-width: 34vw;
            filter: brightness(0) invert(1);
            opacity: 0.96;
        }
        .brand-copy {
            min-width: 0;
        }
        .brand-kicker {
            display: inline-flex;
            align-items: center;
            color: #f2f3f5 !important;
            font-size: 0.76rem;
            font-weight: 800;
            letter-spacing: 0.08em;
            text-transform: uppercase;
            margin-bottom: 0.32rem;
        }
        .brand-hero .brand-title,
        section.brand-hero div.brand-copy div.brand-title {
            color: #ffffff !important;
            font-size: 2rem;
            font-weight: 850;
            line-height: 1.05;
            margin: 0;
        }
        .brand-hero .brand-subtitle,
        section.brand-hero div.brand-copy p.brand-subtitle {
            color: #f4f5f6 !important;
            font-size: 1rem;
            line-height: 1.45;
            margin: 0.45rem 0 0;
            max-width: 46rem;
        }
        .section-title {
            color: #25272b;
            font-size: 1.12rem;
            font-weight: 800;
            margin: 0 0 0.2rem;
        }
        .section-caption {
            color: #5a5e63;
            font-size: 0.86rem;
            margin: 0 0 0.8rem;
        }
        .package-checklist {
            display: grid;
            gap: 0.5rem;
            margin: 0.85rem 0 0.6rem;
        }
        .package-check {
            display: grid;
            grid-template-columns: 5.1rem minmax(0, 1fr);
            gap: 0.65rem;
            align-items: start;
            border: 1px solid #d6d8dc;
            border-radius: 10px;
            padding: 0.65rem 0.75rem;
            background: #f8f9fa;
        }
        .package-badge {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            min-height: 1.7rem;
            border-radius: 999px;
            color: #ffffff !important;
            font-size: 0.72rem;
            font-weight: 850;
            text-transform: uppercase;
            letter-spacing: 0.03em;
        }
        .package-badge.ok { background: #176b43; }
        .package-badge.warn { background: #985d00; }
        .package-badge.info { background: #55585c; }
        .package-check-title {
            color: #25272b !important;
            font-weight: 800;
            line-height: 1.25;
        }
        .package-check-detail {
            color: #5a5e63 !important;
            font-size: 0.84rem;
            line-height: 1.35;
            margin-top: 0.1rem;
        }
        .json-import-card {
            display: grid;
            grid-template-columns: 3.2rem minmax(0, 1fr);
            align-items: center;
            gap: 0.75rem;
            border: 1px solid #b9d7c6;
            border-radius: 12px;
            padding: 0.9rem 1rem;
            margin: 0.45rem 0 0.7rem;
            background: linear-gradient(180deg, #ffffff 0%, #f4fbf7 100%);
            box-shadow: 0 8px 24px rgba(23, 107, 67, 0.08);
        }
        .json-import-card .json-file-icon {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            width: 2.75rem;
            height: 2.75rem;
            border-radius: 10px;
            font-size: 1.05rem;
            font-weight: 900;
            border: 1px solid #8bc5a5;
            background: #dff4e8;
            color: #176b43 !important;
        }
        .json-import-card .json-title {
            color: #25272b !important;
            font-weight: 850;
            line-height: 1.25;
        }
        .json-import-card .json-detail {
            color: #5a5e63 !important;
            font-size: 0.86rem;
            line-height: 1.35;
            margin-top: 0.1rem;
            word-break: break-word;
        }
        .json-import-card.ok {
            border-color: #8bc5a5;
            background: #f0faf4;
        }
        .json-import-card.ok .json-file-icon {
            border-color: #6bb58d;
            background: #dff4e8;
            color: #176b43 !important;
        }
        .json-import-card.error {
            border-color: #f0a7a7;
            background: #fff1f1;
        }
        .json-import-card.error .json-file-icon {
            border-color: #e87b7b;
            background: #ffe1e1;
            color: #b42318 !important;
        }
        .json-import-card.selected {
            border-color: #8bc5a5;
            background: linear-gradient(180deg, #ffffff 0%, #f0faf4 100%);
        }
        .st-key-btn_remover_pacote_arquivo button {
            background: #d92d20 !important;
            border: 1px solid #b42318 !important;
            color: #ffffff !important;
            box-shadow: 0 8px 18px rgba(217, 45, 32, 0.22) !important;
        }
        .st-key-btn_remover_pacote_arquivo button:hover {
            background: #b42318 !important;
            border-color: #912018 !important;
        }
        .st-key-btn_remover_pacote_arquivo button *,
        .st-key-btn_remover_pacote_arquivo button svg,
        .st-key-btn_remover_pacote_arquivo button path {
            color: #ffffff !important;
            stroke: #ffffff !important;
        }
        .generate-callout {
            border: 1px solid #b7c9bd;
            border-radius: 12px;
            padding: 0.9rem 1rem;
            margin: 0.5rem 0 0.85rem;
            background: linear-gradient(180deg, #f6fbf8 0%, #eef8f2 100%);
        }
        .generate-callout-title {
            color: #174d32 !important;
            font-weight: 850;
            line-height: 1.25;
        }
        .generate-callout-detail {
            color: #375846 !important;
            font-size: 0.9rem;
            line-height: 1.35;
            margin-top: 0.2rem;
        }
        .status-strip {
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 0.65rem;
            margin-bottom: 0.2rem;
        }
        .status-item {
            border: 1px solid #d6d8dc;
            border-radius: 10px;
            padding: 0.75rem 0.85rem;
            background: linear-gradient(180deg, #ffffff 0%, #f5f6f7 100%);
        }
        .status-label {
            display: block;
            color: #656a70;
            font-size: 0.72rem;
            text-transform: uppercase;
            letter-spacing: 0.04em;
            margin-bottom: 0.15rem;
        }
        .status-value {
            color: #25272b;
            font-weight: 800;
            font-size: 1.18rem;
            word-break: break-word;
        }
        .stTabs [data-baseweb="tab-list"] {
            gap: 0.35rem;
            border-bottom: 1px solid #d7d9dd;
        }
        .stTabs [data-baseweb="tab"] {
            color: #3e4248;
            font-weight: 650;
        }
        .stTabs [aria-selected="true"] {
            color: #3f4247 !important;
        }
        [data-testid="stAlert"] {
            border-radius: 10px;
        }
        [data-testid="stAlert"],
        [data-testid="stAlert"] *,
        [data-testid="stFileUploaderDropzone"],
        [data-testid="stFileUploaderDropzone"] * {
            color: #303236 !important;
        }
        [data-testid="stFileUploaderDropzone"] {
            background: linear-gradient(180deg, #ffffff 0%, #f6f7f8 100%) !important;
            border: 1px dashed #b8bdc4 !important;
            border-radius: 14px !important;
            min-height: 5.8rem;
            padding: 0.85rem !important;
            box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.82);
        }
        [data-testid="stFileUploaderFile"] {
            border: 1px solid #87bea1 !important;
            border-radius: 14px !important;
            background: #f4fbf7 !important;
            padding: 0.48rem 0.55rem !important;
            box-shadow: 0 10px 22px rgba(23, 107, 67, 0.08) !important;
        }
        [data-testid="stFileUploaderFile"] [data-testid="stFileUploaderFileName"],
        [data-testid="stFileUploaderFile"] [data-testid="stFileUploaderFileSize"],
        [data-testid="stFileUploaderFile"] span,
        [data-testid="stFileUploaderFile"] div {
            color: #263b2f !important;
            font-weight: 650 !important;
        }
        [data-testid="stFileUploaderFile"] svg,
        [data-testid="stFileUploaderFile"] svg *,
        [data-testid="stFileUploaderFile"] path {
            color: #176b43 !important;
            stroke: #176b43 !important;
        }
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] ~ button,
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] + button,
        [data-testid="stFileUploaderDropzone"] button[aria-label*="Add"],
        [data-testid="stFileUploaderDropzone"] button[title*="Add"],
        [data-testid="stFileUploaderDropzone"] button[aria-label*="Adicionar"],
        [data-testid="stFileUploaderDropzone"] button[title*="Adicionar"] {
            display: none !important;
            visibility: hidden !important;
            pointer-events: none !important;
        }
        [data-testid="stFileUploaderFile"] button,
        [data-testid="stFileUploaderDeleteBtn"],
        [data-testid="stFileUploaderFile"] button[aria-label*="Remove"],
        [data-testid="stFileUploaderFile"] button[title*="Remove"],
        [data-testid="stFileUploaderFile"] button[aria-label*="Excluir"],
        [data-testid="stFileUploaderFile"] button[title*="Excluir"] {
            background: #d92d20 !important;
            border-color: #d92d20 !important;
            color: #ffffff !important;
            min-width: 2.1rem !important;
            width: 2.1rem !important;
            height: 2.1rem !important;
            border-radius: 999px !important;
            box-shadow: 0 6px 14px rgba(217, 45, 32, 0.22) !important;
        }
        [data-testid="stFileUploaderFile"] button:hover,
        [data-testid="stFileUploaderDeleteBtn"]:hover {
            background: #b42318 !important;
            border-color: #b42318 !important;
        }
        [data-testid="stFileUploaderFile"] button *,
        [data-testid="stFileUploaderDeleteBtn"] *,
        [data-testid="stFileUploaderFile"] button svg,
        [data-testid="stFileUploaderDeleteBtn"] svg,
        [data-testid="stFileUploaderFile"] button path,
        [data-testid="stFileUploaderDeleteBtn"] path {
            color: #ffffff !important;
            stroke: #ffffff !important;
        }
        [data-testid="stFileUploaderDropzone"] button {
            background: #ffffff !important;
            border: 1px solid #b8bdc4 !important;
            color: #25272b !important;
            min-height: 2.6rem;
            border-radius: 10px !important;
            font-weight: 800 !important;
            box-shadow: 0 6px 16px rgba(37, 39, 43, 0.08) !important;
        }
        [data-testid="stFileUploaderDropzone"] button:hover {
            border-color: #176b43 !important;
            background: #f4fbf7 !important;
            color: #174d32 !important;
        }
        [data-testid="stFileUploaderDropzone"] button *,
        [data-testid="stFileUploaderDropzone"] button svg,
        [data-testid="stFileUploaderDropzone"] button svg *,
        [data-testid="stFileUploaderDropzone"] button path {
            color: #25272b !important;
            fill: none !important;
            stroke: #5a5e63 !important;
        }
        [data-testid="stFileUploaderDropzone"] button:hover *,
        [data-testid="stFileUploaderDropzone"] button:hover svg,
        [data-testid="stFileUploaderDropzone"] button:hover path {
            color: #174d32 !important;
            stroke: #176b43 !important;
        }
        [data-testid="stFileUploader"]:has([data-testid="stFileUploaderFile"])
        [data-testid="stFileUploaderDropzone"] {
            background: #f4fbf7 !important;
            border-style: solid !important;
            border-color: #87bea1 !important;
        }
        [data-testid="stFileUploader"]:has([data-testid="stFileUploaderFile"])
        [data-testid="stFileUploaderDropzone"] > button,
        [data-testid="stFileUploader"]:has([data-testid="stFileUploaderFile"])
        [data-testid="stFileUploaderDropzone"] > div > button:not([data-testid="stFileUploaderDeleteBtn"]) {
            display: none !important;
            visibility: hidden !important;
            pointer-events: none !important;
        }
        [data-testid="stFileUploaderFile"] > div:first-child,
        [data-testid="stFileUploaderFile"] > div:first-child svg,
        [data-testid="stFileUploaderFile"] > div:first-child path {
            background: #dff4e8 !important;
            color: #176b43 !important;
            border-color: #8bc5a5 !important;
            stroke: #176b43 !important;
        }
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] button,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"] {
            display: inline-flex !important;
            visibility: visible !important;
            pointer-events: auto !important;
            align-items: center !important;
            justify-content: center !important;
            background: #d92d20 !important;
            border: 1px solid #b42318 !important;
            color: #ffffff !important;
            min-width: 2.35rem !important;
            width: 2.35rem !important;
            height: 2.35rem !important;
            border-radius: 10px !important;
            box-shadow: 0 6px 14px rgba(217, 45, 32, 0.22) !important;
        }
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] button:hover,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"]:hover {
            background: #b42318 !important;
            border-color: #912018 !important;
        }
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] button *,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"] *,
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] button svg,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"] svg,
        [data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderFile"] button path,
        [data-testid="stFileUploader"] [data-testid="stFileUploaderDeleteBtn"] path {
            background: transparent !important;
            color: #ffffff !important;
            fill: none !important;
            stroke: #ffffff !important;
        }
        label, .stTextInput label, .stTextArea label, .stFileUploader label {
            color: #25272b !important;
            font-weight: 700 !important;
            line-height: 1.25 !important;
        }
        iframe[title*="streamlit_drawable_canvas"] {
            max-width: 100% !important;
        }
        @media (max-width: 640px) {
            .block-container {
                padding: 0.8rem 0.75rem 2.6rem !important;
                max-width: 100% !important;
            }
            [data-testid="stToolbar"],
            [data-testid="stDecoration"],
            [data-testid="stStatusWidget"],
            [data-testid="stAppDeployButton"],
            [data-testid="stHeaderActionElements"],
            [data-testid="stMainMenu"],
            div[class*="stActionButton"],
            div[class*="viewerBadge"],
            a[href*="streamlit.io"] {
                display: none !important;
                visibility: hidden !important;
                pointer-events: none !important;
            }
            .brand-hero {
                align-items: flex-start;
                flex-direction: column;
                padding: 1.05rem;
                gap: 0.8rem;
                border-radius: 14px;
            }
            .brand-logo { width: 150px; max-width: 70vw; }
            .brand-hero .brand-title,
            section.brand-hero div.brand-copy div.brand-title {
                font-size: 1.55rem !important;
                line-height: 1.08 !important;
            }
            .section-title {
                font-size: 1.08rem;
                line-height: 1.2;
            }
            .section-caption {
                font-size: 0.92rem;
                line-height: 1.35;
            }
            div.stButton > button,
            div.stDownloadButton > button {
                width: 100%;
                min-height: 3.1rem;
                padding: 0.55rem 0.75rem;
                font-size: 1rem;
            }
            .stTextInput input,
            .stTextArea textarea,
            [data-baseweb="input"] input,
            [data-baseweb="textarea"] textarea {
                min-height: 2.9rem;
                font-size: 1rem !important;
            }
            [data-testid="stFileUploaderDropzone"] {
                min-height: 5.2rem;
                padding: 0.75rem !important;
            }
            [data-testid="stFileUploaderDropzone"] button {
                width: 100%;
                min-height: 2.75rem;
            }
            [data-testid="stFileUploaderDropzone"] small,
            [data-testid="stFileUploaderDropzone"] span,
            [data-testid="stFileUploaderDropzone"] p {
                font-size: 0.88rem !important;
                line-height: 1.25 !important;
            }
            [data-testid="stHorizontalBlock"] {
                gap: 0.65rem !important;
            }
            [data-testid="column"] {
                min-width: 0 !important;
            }
            .status-strip { grid-template-columns: 1fr; }
            .package-check { grid-template-columns: 1fr; }
            .package-badge { width: fit-content; padding: 0 0.72rem; }
            .stTabs [data-baseweb="tab-list"] {
                overflow-x: auto;
                flex-wrap: nowrap;
                scrollbar-width: thin;
            }
            .stTabs [data-baseweb="tab"] {
                min-width: max-content;
                padding-left: 0.75rem;
                padding-right: 0.75rem;
            }
            iframe[title*="streamlit_drawable_canvas"] {
                width: 100% !important;
                max-width: 100% !important;
            }
        }
        @media (max-width: 380px) {
            .brand-hero .brand-title,
            section.brand-hero div.brand-copy div.brand-title {
                font-size: 1.35rem !important;
            }
            .package-check-detail {
                font-size: 0.8rem;
            }
        }
    </style>
    """,
    unsafe_allow_html=True,
)


def ativar_wake_lock_audio_mobile() -> None:
    from streamlit.components.v1 import html as render_component_html

    render_component_html(
        """
        <script>
        (() => {
            let parentDoc;
            let parentNav;
            try {
                parentDoc = window.parent.document || document;
                parentNav = window.parent.navigator || navigator;
            } catch (erro) {
                parentDoc = document;
                parentNav = navigator;
            }

            if (window.__agresWakeLockInstalled) {
                return;
            }
            window.__agresWakeLockInstalled = true;

            let wakeLock = null;
            let manterAtivoAte = 0;
            const QUATRO_HORAS = 4 * 60 * 60 * 1000;

            function registrarEstado(valor) {
                try {
                    parentDoc.documentElement.setAttribute("data-agres-wake-lock", valor);
                } catch (erro) {}
            }

            async function liberarWakeLock() {
                if (!wakeLock) {
                    return;
                }
                try {
                    await wakeLock.release();
                } catch (erro) {}
                wakeLock = null;
                registrarEstado("released");
            }

            async function solicitarWakeLock() {
                if (!parentNav || !("wakeLock" in parentNav) || wakeLock) {
                    if (!parentNav || !("wakeLock" in parentNav)) {
                        registrarEstado("unsupported");
                    }
                    return;
                }
                if (parentDoc.visibilityState && parentDoc.visibilityState !== "visible") {
                    return;
                }
                try {
                    wakeLock = await parentNav.wakeLock.request("screen");
                    registrarEstado("active");
                    wakeLock.addEventListener("release", () => {
                        wakeLock = null;
                        registrarEstado("released");
                    });
                } catch (erro) {
                    registrarEstado("blocked");
                }
            }

            function textoDoAlvo(alvo) {
                let texto = "";
                let atual = alvo && alvo.nodeType === 1 ? alvo : alvo?.parentElement;
                for (let i = 0; atual && i < 8; i += 1) {
                    texto += " " + (atual.innerText || "");
                    texto += " " + (atual.getAttribute?.("aria-label") || "");
                    texto += " " + (atual.getAttribute?.("title") || "");
                    texto += " " + (atual.getAttribute?.("data-testid") || "");
                    atual = atual.parentElement;
                }
                return texto.toLowerCase();
            }

            function pareceControleDeAudio(alvo) {
                const texto = textoDoAlvo(alvo);
                return /audio|microfone|microphone|gravar|gravando|record|recording|trecho|parar|stop/.test(texto);
            }

            function ocultarIconesStreamlitMobile() {
                const seletores = [
                    "[data-testid='stToolbar']",
                    "[data-testid='stDecoration']",
                    "[data-testid='stStatusWidget']",
                    "[data-testid='stAppDeployButton']",
                    "[data-testid='stHeaderActionElements']",
                    "[data-testid='stMainMenu']",
                    "div[class*='stActionButton']",
                    "div[class*='viewerBadge']",
                    "a[href*='streamlit.io']",
                    "a[href*='share.streamlit.io']",
                    "button[title*='Deploy']",
                    "button[title*='Share']",
                    "button[title*='Manage']",
                    "button[aria-label*='Deploy']",
                    "button[aria-label*='Share']",
                    "button[aria-label*='Manage']"
                ];

                const documentos = [parentDoc];
                try {
                    if (window.top?.document && window.top.document !== parentDoc) {
                        documentos.push(window.top.document);
                    }
                } catch (erro) {}

                for (const docAlvo of documentos) {
                    const janela = docAlvo.defaultView || window.parent || window;
                    const largura = janela.innerWidth || window.innerWidth || 0;
                    const altura = janela.innerHeight || window.innerHeight || 0;
                    if (largura > 760) {
                        continue;
                    }

                    for (const seletor of seletores) {
                        docAlvo.querySelectorAll(seletor).forEach((elemento) => {
                            elemento.style.setProperty("display", "none", "important");
                            elemento.style.setProperty("visibility", "hidden", "important");
                            elemento.style.setProperty("pointer-events", "none", "important");
                        });
                    }

                    docAlvo.querySelectorAll("a[href*='streamlit.io'], a[href*='share.streamlit.io']").forEach((elemento) => {
                        elemento.style.setProperty("display", "none", "important");
                        elemento.style.setProperty("visibility", "hidden", "important");
                        elemento.style.setProperty("pointer-events", "none", "important");
                    });

                    docAlvo.querySelectorAll("button, a, [role='button'], div").forEach((elemento) => {
                        const estilo = janela.getComputedStyle(elemento);
                        const caixa = elemento.getBoundingClientRect();
                        const texto = textoDoAlvo(elemento);
                        const flutuanteInferior =
                            (estilo.position === "fixed" || estilo.position === "sticky") &&
                            caixa.bottom > altura - 170 &&
                            caixa.right > largura - 220 &&
                            caixa.width <= 180 &&
                            caixa.height <= 140;
                        const pareceStreamlit = /hosted with streamlit|streamlit|deploy|toolbar|manage|rerun|running|settings|share|fork|github/.test(texto);
                        if (flutuanteInferior && pareceStreamlit) {
                            elemento.style.setProperty("display", "none", "important");
                            elemento.style.setProperty("visibility", "hidden", "important");
                            elemento.style.setProperty("pointer-events", "none", "important");
                        }
                    });
                }
            }

            function melhorarUploadersJson() {
                const documentos = [parentDoc];
                try {
                    if (window.top?.document && window.top.document !== parentDoc) {
                        documentos.push(window.top.document);
                    }
                } catch (erro) {}

                for (const docAlvo of documentos) {
                    docAlvo.querySelectorAll("[data-testid='stFileUploader']").forEach((uploader) => {
                        const arquivo = uploader.querySelector("[data-testid='stFileUploaderFile']");
                        if (!arquivo) {
                            return;
                        }

                        const dropzone = uploader.querySelector("[data-testid='stFileUploaderDropzone']");
                        if (dropzone) {
                            dropzone.style.setProperty("background", "#f4fbf7", "important");
                            dropzone.style.setProperty("border-color", "#87bea1", "important");
                            dropzone.style.setProperty("border-style", "solid", "important");
                        }

                        arquivo.style.setProperty("background", "#f4fbf7", "important");
                        arquivo.style.setProperty("border", "1px solid #87bea1", "important");
                        arquivo.style.setProperty("border-radius", "14px", "important");

                        const iconeArquivo = Array.from(arquivo.querySelectorAll("svg")).find(
                            (icone) => !icone.closest("button")
                        );
                        if (iconeArquivo) {
                            const caixaIcone = iconeArquivo.parentElement;
                            if (caixaIcone) {
                                caixaIcone.style.setProperty("background", "#dff4e8", "important");
                                caixaIcone.style.setProperty("border-radius", "10px", "important");
                                caixaIcone.style.setProperty("color", "#176b43", "important");
                            }
                            iconeArquivo.style.setProperty("color", "#176b43", "important");
                            iconeArquivo.style.setProperty("stroke", "#176b43", "important");
                            iconeArquivo.querySelectorAll("path").forEach((path) => {
                                path.style.setProperty("stroke", "#176b43", "important");
                            });
                        }

                        uploader.querySelectorAll("button").forEach((botao) => {
                            if (arquivo.contains(botao)) {
                                botao.style.setProperty("display", "inline-flex", "important");
                                botao.style.setProperty("visibility", "visible", "important");
                                botao.style.setProperty("pointer-events", "auto", "important");
                                botao.style.setProperty("background", "#d92d20", "important");
                                botao.style.setProperty("border", "1px solid #b42318", "important");
                                botao.style.setProperty("color", "#ffffff", "important");
                                botao.style.setProperty("border-radius", "10px", "important");
                                botao.style.setProperty("min-width", "2.35rem", "important");
                                botao.style.setProperty("width", "2.35rem", "important");
                                botao.style.setProperty("height", "2.35rem", "important");
                                botao.setAttribute("title", "Excluir pacote selecionado");
                                botao.setAttribute("aria-label", "Excluir pacote selecionado");
                                botao.querySelectorAll("*").forEach((elemento) => {
                                    elemento.style.setProperty("background", "transparent", "important");
                                    elemento.style.setProperty("color", "#ffffff", "important");
                                    elemento.style.setProperty("stroke", "#ffffff", "important");
                                });
                                return;
                            }
                            botao.style.setProperty("display", "none", "important");
                            botao.style.setProperty("visibility", "hidden", "important");
                            botao.style.setProperty("pointer-events", "none", "important");
                        });
                    });

                    docAlvo.querySelectorAll("button").forEach((botao) => {
                        const texto = (botao.textContent || "").trim().toLowerCase();
                        if (!texto.includes("remover pacote")) {
                            return;
                        }
                        botao.style.setProperty("background", "#d92d20", "important");
                        botao.style.setProperty("border", "1px solid #b42318", "important");
                        botao.style.setProperty("color", "#ffffff", "important");
                        botao.style.setProperty("box-shadow", "0 8px 18px rgba(217, 45, 32, 0.22)", "important");
                        botao.setAttribute("title", "Excluir o pacote selecionado");
                        botao.querySelectorAll("*").forEach((elemento) => {
                            elemento.style.setProperty("color", "#ffffff", "important");
                        });
                    });
                }
            }

            function marcarUsoDeAudio(evento) {
                if (!pareceControleDeAudio(evento.target)) {
                    return;
                }
                manterAtivoAte = Date.now() + QUATRO_HORAS;
                solicitarWakeLock();
            }

            parentDoc.addEventListener("click", marcarUsoDeAudio, true);
            parentDoc.addEventListener("touchstart", marcarUsoDeAudio, true);
            parentDoc.addEventListener("pointerdown", marcarUsoDeAudio, true);
            ocultarIconesStreamlitMobile();
            melhorarUploadersJson();
            window.setInterval(() => {
                ocultarIconesStreamlitMobile();
                melhorarUploadersJson();
            }, 1000);
            parentDoc.addEventListener("visibilitychange", () => {
                if (parentDoc.visibilityState === "visible" && Date.now() < manterAtivoAte) {
                    solicitarWakeLock();
                } else {
                    liberarWakeLock();
                }
            });

            window.setInterval(() => {
                if (Date.now() < manterAtivoAte && parentDoc.visibilityState === "visible") {
                    solicitarWakeLock();
                } else if (Date.now() >= manterAtivoAte) {
                    liberarWakeLock();
                }
            }, 30000);
        })();
        </script>
        """,
        height=0,
        width=0,
    )


try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    if not GOOGLE_API_KEY or GOOGLE_API_KEY.strip() == "cole_sua_chave_aqui":
        raise ValueError("GOOGLE_API_KEY ainda não foi preenchida no arquivo .streamlit/secrets.toml.")
    MODELO_GEMINI = str(st.secrets.get("GEMINI_MODEL", "gemini-2.5-flash")).removeprefix("models/")
    genai_client = genai.Client(api_key=GOOGLE_API_KEY)
except Exception:
    st.error("⚠️ Erro crítico: chave GOOGLE_API_KEY não configurada nos Secrets do Streamlit.")
    st.stop()


# ==========================================
# 2. Tratamento de texto e validação
# ==========================================
def data_atual_brasil() -> date:
    try:
        return datetime.now(ZoneInfo("America/Sao_Paulo")).date()
    except Exception:
        return date.today()


def imagem_data_uri(caminho: Path) -> str:
    if not caminho.exists():
        return ""
    conteudo = base64.b64encode(caminho.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{conteudo}"


def normalizar_busca(texto: str) -> str:
    texto = unicodedata.normalize("NFD", texto or "")
    texto = "".join(caractere for caractere in texto if unicodedata.category(caractere) != "Mn")
    return texto.lower()


def valor_para_texto(valor) -> str:
    if valor is None:
        return ""
    if isinstance(valor, list):
        return "\n".join(valor_para_texto(item) for item in valor if valor_para_texto(item))
    if isinstance(valor, dict):
        linhas = []
        for chave, item in valor.items():
            texto_item = valor_para_texto(item)
            if texto_item:
                linhas.append(f"{chave}: {texto_item}")
        return "\n".join(linhas)
    return str(valor)


def limpar_texto(valor) -> str:
    texto = valor_para_texto(valor).replace("\r\n", "\n").replace("\r", "\n")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r" *\n *", "\n", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    texto = texto.strip(" \n\t;")
    if normalizar_busca(texto) in {"null", "none", "n/a", "nao informado", "nao informada"}:
        return ""
    return texto


def dicionario_ou_vazio(valor) -> dict:
    return valor if isinstance(valor, dict) else {}


def lista_ou_vazia(valor) -> list:
    return valor if isinstance(valor, list) else []


def texto_ou_padrao(valor, padrao="Não informado") -> str:
    texto = limpar_texto(valor)
    return texto if texto else padrao


TITULOS_RELATO_REMOVER = {
    "contexto do atendimento",
    "descricao do problema",
    "diagnostico inicial",
    "acoes corretivas",
    "acoes executadas",
    "procedimentos executados",
    "configuracoes e calibracoes relevantes",
    "configuracoes e calibracoes",
    "testes complementares",
    "validacao do sistema",
    "resultado final",
    "conclusao tecnica",
    "intervencoes fisicas e diagnosticos registrados",
}


def finalizar_frase(texto: str) -> str:
    texto = limpar_texto(texto).rstrip(" .;:")
    if not texto:
        return ""
    return texto if texto.endswith((".", "!", "?")) else f"{texto}."


def limpar_relato_narrativo(valor) -> str:
    texto = limpar_texto(valor)
    if not texto:
        return ""

    texto = re.sub(r"\*\*(.*?)\*\*", r"\1", texto)
    texto = re.sub(r"__(.*?)__", r"\1", texto)
    paragrafos = []
    linhas_paragrafo = []

    def fechar_paragrafo() -> None:
        if linhas_paragrafo:
            paragrafo = " ".join(linhas_paragrafo)
            paragrafo = re.sub(r"\s+([,.;:])", r"\1", paragrafo)
            if paragrafo:
                paragrafos.append(finalizar_frase(paragrafo))
            linhas_paragrafo.clear()

    for linha in texto.split("\n"):
        linha = linha.strip()
        if not linha:
            fechar_paragrafo()
            continue

        linha = re.sub(r"^\s*[-*•]+\s*", "", linha).strip()
        linha = linha.strip("*_ ")
        titulo_normalizado = normalizar_busca(linha.rstrip(":"))
        titulo_curto = linha.endswith(":") and len(linha.split()) <= 8
        if titulo_normalizado in TITULOS_RELATO_REMOVER or titulo_curto:
            fechar_paragrafo()
            continue

        linhas_paragrafo.append(linha)

    fechar_paragrafo()
    return "\n\n".join(paragrafos)


def dividir_itens(texto: str) -> list[str]:
    itens = []
    for linha in limpar_texto(texto).split("\n"):
        for parte in re.split(r"\s*[;•]\s*", linha):
            item = parte.strip(" -–—\t")
            if item:
                itens.append(item)
    return itens


def dividir_frases_tecnicas(texto: str) -> list[str]:
    itens = []
    for item in dividir_itens(texto):
        partes = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9])", item)
        itens.extend(parte.strip(" -–—\t") for parte in partes if parte.strip(" -–—\t"))
    return itens


def formatar_topicos_tecnicos(texto: str) -> str:
    topicos = []
    vistos = set()
    for item in dividir_frases_tecnicas(texto):
        item = finalizar_frase(item)
        chave = normalizar_busca(item)
        if item and chave not in vistos:
            topicos.append(f"• {item}")
            vistos.add(chave)
    return "\n".join(topicos)


def formatar_equipamentos_agres(texto: str) -> str:
    texto = limpar_texto(texto)
    if not texto or normalizar_busca(texto) == "nao informado":
        return ""

    linhas = []
    vistos = set()
    rotulos = (
        ("tela instalada", "Tela instalada"),
        ("tela", "Tela instalada"),
        ("equipamento", "Equipamento"),
        ("ecu", "ECU"),
        ("compensador de terreno", "Compensador de terreno"),
        ("compensador", "Compensador de terreno"),
        ("piloto hidraulico", "Piloto Hidráulico"),
        ("piloto hidráulico", "Piloto Hidráulico"),
        ("modelo", "Modelo"),
        ("s/n", "S/N"),
        ("sn", "S/N"),
        ("numero de serie", "S/N"),
        ("número de série", "S/N"),
        ("serial", "S/N"),
        ("aplicacao", "Aplicação"),
        ("aplicação", "Aplicação"),
        ("versao de aplicacao", "Aplicação"),
        ("versão de aplicação", "Aplicação"),
        ("sistema", "Sistema"),
        ("versao de sistema", "Sistema"),
        ("versão de sistema", "Sistema"),
        ("versao de software da tela", "Software da tela"),
        ("versão de software da tela", "Software da tela"),
        ("versao de software da ecu", "SW"),
        ("versão de software da ecu", "SW"),
        ("sw", "SW"),
        ("hw", "HW"),
        ("versao do compensador", "Versão do compensador"),
        ("versão do compensador", "Versão do compensador"),
        ("comando de pulverizacao", "Comando de pulverização"),
        ("comando de pulverização", "Comando de pulverização"),
        ("comando", "Comando"),
        ("cabeamento", "Cabeamento"),
    )

    partes = []
    for linha in texto.replace("•", "\n").split("\n"):
        for parte in re.split(r"\s*;\s*", linha):
            parte = parte.strip(" -–—\t")
            if parte:
                partes.append(parte)

    for item in partes:
        item = re.sub(r"\s+", " ", item).strip(" ,.;")
        if not item:
            continue

        if ":" in item:
            rotulo, valor = item.split(":", 1)
            rotulo_limpo = normalizar_busca(rotulo).strip()
            valor = valor.strip(" ,.;")
            valor = normalizar_modelos_equipamentos_agres(valor)
            rotulo_final = next((final for chave, final in rotulos if rotulo_limpo == chave), rotulo.strip())
            item = f"{rotulo_final}: {valor}" if valor else f"{rotulo_final}:"
        else:
            item_busca = normalizar_busca(item)
            item = normalizar_modelos_equipamentos_agres(item)
            if item_busca.startswith(("sn ", "s n ")):
                item = re.sub(r"^(?:s\s*n|sn)\s+", "S/N: ", item, flags=re.I)

        chave = normalizar_busca(item)
        if chave and chave not in vistos:
            if linhas and re.match(r"^(ECU|Tela instalada|Equipamento|Compensador de terreno|Piloto Hidráulico)\b", item):
                linhas.append("")
            linhas.append(item)
            vistos.add(chave)

    return "\n".join(linhas).strip()


def selecionar_imagens_equipamentos_agres(caminhos_cabecalho: dict, evidencias: dict) -> list[dict]:
    candidatos = []
    if caminhos_cabecalho.get("info_equip"):
        candidatos.append(
            {
                "caminho": caminhos_cabecalho["info_equip"],
                "origem": "2. Cabeçalho do Relatório - Informações do Equipamento / Equipamento Agres",
            }
        )

    for indice, caminho in enumerate(evidencias.get("fotos_equipamento", []) or [], start=1):
        candidatos.append(
            {
                "caminho": caminho,
                "origem": f"3. Evidências Fotográficas - Equipamento Agres, foto {indice}",
            }
        )

    for indice, caminho in enumerate(evidencias.get("fotos_configuracao", []) or [], start=1):
        candidatos.append(
            {
                "caminho": caminho,
                "origem": f"3. Evidências Fotográficas - Configurações, foto {indice}",
            }
        )

    selecionadas = []
    vistos = set()
    for item in candidatos:
        caminho = item.get("caminho")
        if not caminho:
            continue
        caminho = Path(caminho)
        chave = str(caminho.resolve()) if caminho.exists() else str(caminho)
        if chave in vistos or not caminho.exists() or caminho.suffix.lower().lstrip(".") not in EXTENSOES_IMAGEM:
            continue
        selecionadas.append({"caminho": caminho, "origem": item["origem"]})
        vistos.add(chave)
        if len(selecionadas) >= MAX_IMAGENS_EQUIPAMENTOS_IA:
            break
    return selecionadas


def normalizar_imagens_equipamentos_ia(imagens: list | None) -> list[dict]:
    imagens_validas = []
    vistos = set()
    for indice, item in enumerate(imagens or [], start=1):
        if isinstance(item, dict):
            caminho = item.get("caminho") or item.get("path") or item.get("arquivo")
            origem = limpar_texto(item.get("origem", "")) or f"Imagem de equipamento Agres {indice}"
        else:
            caminho = item
            origem = f"Imagem de equipamento Agres {indice}"
        if not caminho:
            continue
        caminho = Path(caminho)
        chave = str(caminho.resolve()) if caminho.exists() else str(caminho)
        if chave in vistos or not caminho.exists() or caminho.suffix.lower().lstrip(".") not in EXTENSOES_IMAGEM:
            continue
        imagens_validas.append({"caminho": caminho, "origem": origem})
        vistos.add(chave)
        if len(imagens_validas) >= MAX_IMAGENS_EQUIPAMENTOS_IA:
            break
    return imagens_validas


def contem_termo(texto: str, termos: tuple[str, ...]) -> bool:
    texto_normalizado = normalizar_busca(texto)
    return any(termo in texto_normalizado for termo in termos)


def normalizar_marcador_servico(valor) -> tuple[str, list[str]]:
    texto = limpar_texto(valor)
    if not texto:
        return "", []

    texto_normalizado = normalizar_busca(texto)
    marcador = ""
    if texto.strip().upper() == "X" or any(
        termo in texto_normalizado
        for termo in (
            "sim",
            "realizado",
            "realizada",
            "executado",
            "executada",
            "suporte",
            "instalacao",
            "treinamento",
            "validacao",
            "homologacao",
        )
    ):
        marcador = "X"

    detalhes = [] if texto.strip().upper() == "X" else [texto]
    return marcador, detalhes


def normalizar_tipo_atendimento(valor, padrao: str = "") -> str:
    texto = normalizar_busca(valor)
    if not texto:
        return padrao
    if texto in TIPOS_ATENDIMENTO:
        return texto
    if "valid" in texto or "homolog" in texto:
        return "validacao_homologacao"
    if "instal" in texto:
        return "instalacao"
    if "trein" in texto:
        return "treinamento"
    if "suporte" in texto or "assist" in texto or "diagnost" in texto:
        return "suporte"
    return padrao


def normalizar_tipos_atendimento(valor, padrao=None, limite: int = 2) -> list[str]:
    tipos: list[str] = []

    def adicionar(tipo: str) -> None:
        tipo_normalizado = normalizar_tipo_atendimento(tipo)
        if tipo_normalizado and tipo_normalizado not in tipos:
            tipos.append(tipo_normalizado)

    if isinstance(valor, dict):
        for chave in CAMPOS_SERVICO:
            if valor.get(chave):
                adicionar(chave)
    elif isinstance(valor, (list, tuple, set)):
        for item in valor:
            adicionar(str(item))
    else:
        texto_original = limpar_texto(valor)
        texto = normalizar_busca(texto_original)
        if texto in TIPOS_ATENDIMENTO:
            adicionar(texto)
        else:
            if "valid" in texto or "homolog" in texto:
                adicionar("validacao_homologacao")
            if "suporte" in texto or "assist" in texto or "diagnost" in texto:
                adicionar("suporte")
            if "instal" in texto:
                adicionar("instalacao")
            if "trein" in texto:
                adicionar("treinamento")

    if not tipos and padrao is not None:
        tipos = normalizar_tipos_atendimento(padrao, None, limite)

    if "validacao_homologacao" in tipos:
        return ["validacao_homologacao"]

    tipos = sorted(tipos, key=lambda tipo: ORDEM_TIPOS_ATENDIMENTO.index(tipo) if tipo in ORDEM_TIPOS_ATENDIMENTO else 99)
    if "treinamento" in tipos:
        primario = next((tipo for tipo in tipos if tipo in {"suporte", "instalacao"}), "")
        return ([primario] if primario else []) + ["treinamento"]
    if "suporte" in tipos and "instalacao" in tipos:
        return [tipos[0]]
    return tipos[:limite]


def escolher_tipos_atendimento(dados: dict, marcadores: dict | None = None, limite: int = 2) -> list[str]:
    tipos = normalizar_tipos_atendimento(dados.get("tipos_atendimento") or dados.get("tipo_atendimento"), limite=limite)
    if tipos:
        return tipos

    marcadores = marcadores or {}
    marcados = [
        campo
        for campo in CAMPOS_SERVICO
        if marcadores.get(campo) == "X" or dados.get(campo) == "X"
    ]
    if marcados:
        return marcados[:limite]

    texto = normalizar_busca(
        "\n".join(
            limpar_texto(dados.get(campo, ""))
            for campo in ("objetivos", "relato", "equipamentos", "maquinas", "configuracoes")
        )
    )
    inferidos: list[str] = []
    if re.search(r"\b(validacao|validar|validado|homologacao|homologar|homologado)\b", texto):
        inferidos.append("validacao_homologacao")
    if re.search(r"\b(suporte|falha|diagnostico|correcao|ajuste|manutencao)\b", texto):
        inferidos.append("suporte")
    if re.search(r"\b(instalacao|instalar|instalado|montagem|chicote|fixacao)\b", texto):
        inferidos.append("instalacao")
    if re.search(r"\b(treinamento|treinar|orientacao operacional|capacitacao)\b", texto):
        inferidos.append("treinamento")

    return inferidos[:limite] or ["suporte"]


def escolher_tipo_atendimento_unico(dados: dict, marcadores: dict | None = None) -> str:
    return escolher_tipos_atendimento(dados, marcadores, limite=1)[0]


def aplicar_tipos_atendimento(dados: dict, tipos_preferidos=None) -> dict:
    tipos = normalizar_tipos_atendimento(tipos_preferidos, limite=2) or escolher_tipos_atendimento(dados)
    dados["tipos_atendimento"] = tipos
    dados["tipo_atendimento"] = tipos[0] if tipos else ""
    for campo in CAMPOS_SERVICO:
        dados[campo] = "X" if campo in tipos else ""
    return dados


def aplicar_tipo_atendimento_unico(dados: dict, tipo_preferido: str = "") -> dict:
    return aplicar_tipos_atendimento(dados, tipo_preferido)


def escrever_celula_servico(cell: _Cell, texto: str, tamanho: int = 14) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(texto)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(tamanho)


def aplicar_tipo_atendimento_word(caminho_docx: Path, dados: dict) -> None:
    tipos = normalizar_tipos_atendimento(dados.get("tipos_atendimento") or dados.get("tipo_atendimento"), limite=2)
    if "validacao_homologacao" not in tipos:
        return

    outros_tipos = [tipo for tipo in tipos if tipo != "validacao_homologacao"]
    documento = Document(caminho_docx)
    for tabela in documento.tables:
        if not tabela.rows:
            continue
        primeira_linha = tabela.rows[0]
        textos = [normalizar_busca(celula.text) for celula in primeira_linha.cells]
        if not any("suporte" in texto for texto in textos) or not any("treinamento" in texto for texto in textos):
            continue

        celulas = primeira_linha.cells
        try:
            if outros_tipos:
                celula_validacao = celulas[0].merge(celulas[2])
                escrever_celula_servico(celula_validacao, "VALIDAÇÃO/HOMOLOGAÇÃO", tamanho=11)
                escrever_celula_servico(celulas[3], "X", tamanho=14)
                escrever_celula_servico(celulas[4], TIPOS_ATENDIMENTO[outros_tipos[0]].upper(), tamanho=11)
                escrever_celula_servico(celulas[5], "X", tamanho=14)
            else:
                celula_rotulo = celulas[0].merge(celulas[4])
                celula_marcador = celulas[5]
                escrever_celula_servico(celula_rotulo, "VALIDAÇÃO/HOMOLOGAÇÃO", tamanho=13)
                escrever_celula_servico(celula_marcador, "X", tamanho=14)
        except Exception:
            escrever_celula_servico(celulas[0], "VALIDAÇÃO/HOMOLOGAÇÃO", tamanho=10)
            escrever_celula_servico(celulas[1], "X", tamanho=14)
            if outros_tipos and len(celulas) >= 4:
                escrever_celula_servico(celulas[2], TIPOS_ATENDIMENTO[outros_tipos[0]].upper(), tamanho=10)
                escrever_celula_servico(celulas[3], "X", tamanho=14)
            for celula in celulas[2:]:
                if not outros_tipos or celula not in celulas[2:4]:
                    escrever_celula_servico(celula, "", tamanho=13)
        documento.save(caminho_docx)
        return


def extrair_link_localizacao(texto: str) -> str:
    match = re.search(
        r"https?://[^\s<>()\"]*(?:maps\.app\.goo\.gl|goo\.gl/maps|google\.[^\s<>()\"]*/maps|maps\.google\.[^\s<>()\"]*)[^\s<>()\"]*",
        limpar_texto(texto),
        flags=re.I,
    )
    return match.group(0).rstrip(".,;") if match else ""


def filtrar_campo_curto(texto: str, termos_bloqueados: tuple[str, ...]) -> tuple[str, list[str]]:
    itens_validos = []
    itens_para_relato = []

    for item in dividir_itens(texto):
        if contem_termo(item, termos_bloqueados):
            itens_para_relato.append(item)
        else:
            itens_validos.append(item)

    return "\n".join(itens_validos).strip(), itens_para_relato


def adicionar_ao_relato(relato: str, itens: list[str]) -> str:
    if not itens:
        return limpar_relato_narrativo(relato)

    relato_base = limpar_relato_narrativo(relato)
    relato_normalizado = normalizar_busca(relato_base)
    itens_novos = [item for item in itens if normalizar_busca(item) not in relato_normalizado]
    if not itens_novos:
        return relato_base

    detalhes = "; ".join(limpar_texto(item).rstrip(" .;") for item in itens_novos if limpar_texto(item))
    complemento = finalizar_frase(f"Foram registrados ainda os seguintes pontos técnicos: {detalhes}")
    return limpar_relato_narrativo((relato_base + "\n\n" + complemento).strip() if relato_base else complemento)


def contar_palavras_relato(texto: str) -> int:
    return len(re.findall(r"\b[\wÀ-ÿ]+\b", limpar_texto(texto)))


def resumir_equipamentos_para_relato(texto: str) -> str:
    linhas_importantes = []
    vistos = set()
    rotulos_importantes = (
        "tela instalada",
        "equipamento",
        "ecu",
        "compensador de terreno",
        "modelo",
        "s/n",
        "sn",
        "numero de serie",
        "número de série",
        "aplicacao",
        "aplicação",
        "sistema",
        "sw",
        "hw",
        "versao do compensador",
        "versão do compensador",
    )
    for linha in limpar_texto(texto).replace("•", "\n").split("\n"):
        linha = linha.strip(" -–—\t")
        if not linha or ":" not in linha:
            continue
        rotulo, valor = linha.split(":", 1)
        rotulo_normalizado = normalizar_busca(rotulo)
        valor = limpar_texto(valor)
        if not valor:
            continue
        if rotulo_normalizado not in rotulos_importantes:
            continue
        item = f"{rotulo.strip()}: {valor}"
        chave = normalizar_busca(item)
        if chave not in vistos:
            linhas_importantes.append(item)
            vistos.add(chave)
        if len(linhas_importantes) >= 12:
            break
    return "; ".join(linhas_importantes)


def montar_complemento_rastreabilidade(dados: dict) -> str:
    campos = (
        ("cliente_local", "Cliente/local"),
        ("localizacao_maps", "Localização"),
        ("equipamentos", "Equipamento"),
        ("maquinas", "Máquina/implemento"),
        ("objetivos", "Objetivo"),
        ("acompanhantes", "Acompanhantes"),
    )
    partes = []
    for campo, rotulo in campos:
        valor = limpar_texto(dados.get(campo, ""))
        if not valor or normalizar_busca(valor) == "nao informado":
            continue
        if campo == "equipamentos":
            valor = resumir_equipamentos_para_relato(valor)
        else:
            valor = re.sub(r"\n+", "; ", valor.replace("•", "")).strip(" ;")
        if valor:
            partes.append(f"{rotulo}: {valor}")

    if not partes:
        return ""

    return finalizar_frase(
        "Para fins de rastreabilidade técnica do atendimento, também ficaram registrados os seguintes dados complementares: "
        + "; ".join(partes)
    )


def enriquecer_relato_quando_curto(dados: dict) -> str:
    relato = limpar_relato_narrativo(dados.get("relato", ""))
    if contar_palavras_relato(relato) >= 180:
        return relato

    complemento = montar_complemento_rastreabilidade(dados)
    if not complemento:
        return relato

    relato_normalizado = normalizar_busca(relato)
    if normalizar_busca(complemento[:120]) in relato_normalizado:
        return relato

    return limpar_relato_narrativo(f"{relato}\n\n{complemento}" if relato else complemento)


def normalizar_dados_relatorio(dados: dict) -> dict:
    dados_normalizados = {campo: limpar_texto(dados.get(campo, "")) for campo in CAMPOS_RELATORIO}
    detalhes_para_relato = []
    marcadores_servico = {}

    for campo in CAMPOS_SERVICO:
        marcador, detalhes = normalizar_marcador_servico(dados_normalizados[campo])
        marcadores_servico[campo] = marcador
        detalhes_para_relato.extend(f"{ROTULOS_CAMPOS[campo]}: {detalhe}" for detalhe in detalhes)

    aplicar_tipos_atendimento(
        dados_normalizados,
        dados_normalizados.get("tipos_atendimento")
        or dados_normalizados.get("tipo_atendimento")
        or escolher_tipos_atendimento(dados_normalizados, marcadores_servico),
    )

    configuracoes, itens_config_relato = filtrar_campo_curto(
        dados_normalizados["configuracoes"],
        TERMOS_INTERVENCAO_FISICA,
    )
    calibracoes, itens_calibracao_relato = filtrar_campo_curto(
        dados_normalizados["calibracoes"],
        TERMOS_NAO_CALIBRACAO,
    )

    dados_normalizados["configuracoes"] = formatar_topicos_tecnicos(configuracoes) or "Não informado"
    dados_normalizados["calibracoes"] = formatar_topicos_tecnicos(calibracoes) or "Não informado"
    dados_normalizados["equipamentos"] = formatar_equipamentos_agres(dados_normalizados["equipamentos"])
    dados_normalizados["relato"] = adicionar_ao_relato(
        dados_normalizados["relato"],
        itens_config_relato + itens_calibracao_relato + detalhes_para_relato,
    )
    dados_normalizados["relato"] = enriquecer_relato_quando_curto(dados_normalizados)

    if not dados_normalizados["relato"]:
        dados_normalizados["relato"] = "Não informado nos áudios ou observações encaminhadas."

    if not dados_normalizados["localizacao_maps"]:
        dados_normalizados["localizacao_maps"] = extrair_link_localizacao(
            "\n".join(
                dados_normalizados.get(campo, "")
                for campo in ("cliente_local", "objetivos", "relato")
            )
        )

    dados_normalizados["data_visita"] = dados_normalizados["data_visita"] or data_atual_brasil().strftime("%d/%m/%Y")
    dados_normalizados["tecnicos"] = dados_normalizados["tecnicos"] or ""

    for campo in (
        "cliente_local",
        "localizacao_maps",
        "equipamentos",
        "maquinas",
        "objetivos",
        "acompanhantes",
        "responsavel_revenda_fabrica",
        "documento_revenda_fabrica",
        "responsavel_fazenda",
        "documento_fazenda",
    ):
        dados_normalizados[campo] = texto_ou_padrao(dados_normalizados[campo])

    return dados_normalizados


# ==========================================
# 3. Inteligência artificial
# ==========================================
def montar_prompt(contexto_manual: str = "") -> str:
    hoje = data_atual_brasil().strftime("%d/%m/%Y")
    contexto_manual = limpar_texto(contexto_manual)
    bloco_contexto = (
        f"\nANOTAÇÕES COMPLEMENTARES INFORMADAS PELO TÉCNICO:\n{contexto_manual}\n"
        if contexto_manual
        else ""
    )

    return f"""
Você é redator técnico da Agres e deve transformar áudios, anotações e fotos de atendimento de campo em dados para um relatório formal.

Use português técnico, claro e objetivo, porém completo e minucioso. Reescreva falas informais em linguagem profissional, sem inventar dados, versões, medidas, peças ou conclusões que não estejam no material recebido.
Não resuma o atendimento. Preserve o máximo possível de informações técnicas citadas nos áudios/anotações, incluindo nomes, datas, local, cliente, máquina, implemento, equipamento, versões, números de série, sintomas, hipóteses, testes, tentativas, parâmetros, valores, componentes, decisões, dificuldades, pendências e conclusão.
{bloco_contexto}
REGRAS DE CLASSIFICAÇÃO DOS CAMPOS:
1. tipo_atendimento: retornar a opção principal entre "suporte", "instalacao", "treinamento" ou "validacao_homologacao".
2. suporte, instalacao, treinamento e validacao_homologacao: retornar "X" no tipo principal identificado e retornar "" nos demais campos. Quando o pacote offline trouxer dois tipos selecionados pelo técnico, o sistema aplicará essa seleção após o processamento da IA.
   Use "validacao_homologacao" quando o atendimento tiver objetivo de validar, homologar, aprovar funcionamento, acompanhar testes de aceitação ou confirmar desempenho de equipamento/sistema.
3. data_visita: preserve intervalo de datas quando o atendimento ocorrer em mais de um dia, por exemplo "19 a 21/01/2026".
4. cliente_local: informar cliente, cidade/UF, revenda, fábrica e propriedade quando existirem.
5. localizacao_maps: informar somente o link do Google Maps, Maps ou coordenadas da fazenda quando existirem; caso contrário, retornar "".
6. equipamentos: organizar somente equipamentos, sistemas e componentes da Agres. As imagens dos campos "2. Cabeçalho do Relatório - Equipamento Agres/Informações do Equipamento" e "3. Evidências Fotográficas - Equipamento Agres" são fonte primária para este campo. Use OCR visual nas fotos/telas recebidas para ler etiquetas, menus, telas de versão, plaquetas, número de série e QR Code quando legível. Não incluir máquina, trator ou implemento de outros fabricantes neste campo.
   Padronizar o campo "equipamentos" em blocos, uma informação por linha, sem tópicos e sem texto corrido. Usar os rótulos abaixo quando a informação existir:
   Tela instalada: <modelo da tela ou equipamento Agres>
   S/N: <número de série da tela/equipamento Agres>
   Aplicação: <versão de aplicação/software da tela>
   Sistema: <versão de sistema da tela>

   ECU: <modelo da ECU>
   S/N: <número de série da ECU>
   SW: <versão de software da ECU>
   HW: <versão de hardware da ECU>

   Compensador de terreno: <modelo ou descrição>
   Modelo: <ANP21 ou ANP40>
   S/N: <número de série do compensador>
   Versão do compensador: <versão>

   Modelos conhecidos de tela/equipamento: agroNave 7, agroNave 12, agroNave 12W, isoView, SprayRate e Agronave.
   Modelos do ISOVIEW: ISO 30 (Navegação), ISO31FP (Pulverização), ISO31OFP (Fruticultura), ISO32FP (Pulverização + Piloto), ISO32OFP (Fruticultura + Piloto), ISO33E (Piloto Elétrico), ISO33H (Piloto Hidráulico), ISO34 (Adubação), ISO35 (Adubação + Piloto), ISO36 (Monitor de Plantio) e ISO37 (Monitor de Plantio + Piloto).
   Ao registrar um modelo ISOVIEW no relatório, escrever somente o código do modelo, sem a descrição entre parênteses. Exemplo: "ISO33H", não "ISO33H (Piloto Hidráulico)".
   Modelos conhecidos de compensador: ANP21 e ANP40.
   Modelos conhecidos de ECU: isoBox Sprayer (Pulverização) e isoBox Spreader (Adubação).
   Se existir mais de uma tela, ECU, compensador ou módulo Agres, repetir o bloco para cada equipamento identificado, sem sobrescrever informações.
   Preservar números de série exatamente como aparecem na foto, incluindo zeros à esquerda, letras, hífens e maiúsculas/minúsculas relevantes.
   Se a foto mostrar "Tela instalada", "Versão de Aplicação", "Versão de Sistema", "Número de série", "ECU", "SW", "HW", "ANP21" ou "ANP40", esses dados devem ir obrigatoriamente para o campo "equipamentos" quando legíveis.
   Se a foto estiver ilegível ou a informação não tiver sido citada, não inventar. Apenas omitir a linha específica ou retornar "Não informado" quando nenhum equipamento Agres for identificado.
7. maquinas: organizar máquinas, tratores e implementos de outros fabricantes em linhas com fabricante, modelo, comando de válvulas e características relevantes. Exemplos como Baldan Avolla 2500 pertencem exclusivamente a este campo.
8. objetivos: escrever somente o objetivo principal do atendimento, em uma ou duas frases.
9. configuracoes: incluir somente parâmetros de sistema, software, tela, ECU, controlador, seções, geometria, versões, módulos habilitados, ganhos ou ajustes feitos em menus. Retornar em tópicos objetivos, uma linha por item, sem texto corrido. Quando houver valores, use o padrão "Parâmetro: valor".
10. calibracoes: incluir somente calibrações, aferições e validações com valores, medidas, sensores, vazão, largura, offset, angulação ou parâmetros numéricos. Retornar em tópicos objetivos, uma linha por item, sem texto corrido.
11. tecnicos: retornar "". O Técnico Responsável Agres é definido exclusivamente pela seleção realizada na coleta offline.
12. acompanhantes: informar técnicos, operadores, proprietários, consultores, representantes ou demais pessoas que acompanharam o atendimento. Pessoas citadas apenas como acompanhantes nunca devem ser incluídas no campo "tecnicos".
13. responsavel_revenda_fabrica: informar o nome do responsável da revenda, fábrica ou Agres que validou/acompanhou o atendimento, quando mencionado.
14. documento_revenda_fabrica: retornar sempre "".
15. responsavel_fazenda: informar o nome do responsável da fazenda, cliente, operador, encarregado ou proprietário que validou/acompanhou o atendimento, quando mencionado.
16. documento_fazenda: retornar sempre "".
17. relato: concentrar a narrativa técnica e cronológica do que foi realizado em campo. Cabos, chicotes, conectores, soldas, conversores PNP/NPN, pinagem, relés, terminadores CAN, suportes físicos, falhas, diagnósticos, testes, correções, pendências e recomendações pertencem ao relato quando fizerem parte da ação executada.
   Não copiar para o relato listas de parâmetros, telas de menu, configurações, calibrações, valores de seção, largura, bicos, offset, ganhos, tabelas de versão ou demais transcrições de OCR que já estejam nos campos "equipamentos", "configuracoes" ou "calibracoes".
   No relato, usar dados de identificação do equipamento somente de forma resumida quando forem relevantes para rastreabilidade, por exemplo modelo, número de série, versão da tela/ECU/compensador e tipo de ECU.
18. nome_arquivo_sugerido: retornar "". O nome final é montado automaticamente pelo sistema a partir do Técnico Responsável Agres selecionado na coleta.

PADRÃO DO RELATO:
- Escrever em terceira pessoa.
- Escrever em formato narrativo, como relato técnico do que foi realizado em campo.
- Não usar subtítulos, tópicos, listas, markdown, negrito com **texto**, enumeração ou blocos separados por títulos.
- Descrever em parágrafos corridos o contexto do atendimento, problema informado, diagnóstico inicial, procedimentos executados, problemas encontrados, correções aplicadas, testes de funcionamento, resultado final, pendências, recomendações e conclusão técnica.
- Não condensar várias ações em uma única frase genérica. Quando houver sequência de atividades, descrever a ordem de execução e o motivo técnico de cada etapa.
- Quando forem citados modelo, número de série ou versão de tela, ECU ou compensador, pode mencionar esses dados no relato de forma resumida para rastreabilidade. Parâmetros, configurações, calibrações, seções, bicos, offsets, larguras, vazões, pressões, ganhos, menus e listas extraídas das fotos devem ficar somente nos campos específicos.
- Não iniciar parágrafos do relato com "Configurações", "Calibrações", "Parâmetros", "Equipamentos" ou outro título de campo. O relato deve permanecer narrativo.
- Separar o relato em 6 a 10 parágrafos quando houver material suficiente, usando uma linha em branco entre parágrafos. Não retornar tudo em um único bloco.
- Quando houver áudio ou complemento técnico com conteúdo suficiente, escrever preferencialmente um relato com no mínimo 250 palavras.
- Em atendimentos de vários dias, separar a sequência por data ou por etapa.
- Se alguma informação técnica tiver sido mencionada de forma incerta, registrar como "foi informado" ou "foi relatado", sem transformar em certeza absoluta.
- Informar "Não informado" nos campos textuais quando o dado não for mencionado.

LEITURA DAS FOTOS:
- Quando houver fotos de identificação do equipamento ou telas de configuração, analisar visualmente os textos legíveis para preencher "equipamentos", "configuracoes" e "calibracoes".
- As fotos de "2. Cabeçalho - Equipamento Agres/Informações do Equipamento" e "3. Evidência Fotográfica - Equipamento Agres" devem ser analisadas antes do áudio para identificação técnica.
- Priorizar no campo "equipamentos" os dados de identificação Agres: modelo da tela/equipamento, número de série, versão de software/aplicação/sistema, ECU, SW/HW, compensador, ANP21/ANP40 e versões.
- Se uma foto mostrar a tela de sistema com campos como "Versão de Aplicação", "Versão de Sistema", "Número de série", "SW" ou "HW", transcrever exatamente no campo correto.
- Não classificar como equipamento Agres nomes de máquinas ou implementos de terceiros, como Baldan, Avolla, Jacto, Massey, John Deere, Case, New Holland, Stara ou similares; esses nomes pertencem ao campo "maquinas".

Retorne apenas um JSON válido, sem markdown e sem comentários, com exatamente esta estrutura:
{{
    "tipo_atendimento": "",
    "suporte": "",
    "instalacao": "",
    "treinamento": "",
    "validacao_homologacao": "",
    "data_visita": "{hoje}",
    "tecnicos": "",
    "cliente_local": "",
    "localizacao_maps": "",
    "equipamentos": "",
    "maquinas": "",
    "objetivos": "",
    "configuracoes": "",
    "calibracoes": "",
    "acompanhantes": "",
    "responsavel_revenda_fabrica": "",
    "documento_revenda_fabrica": "",
    "responsavel_fazenda": "",
    "documento_fazenda": "",
    "nome_arquivo_sugerido": "",
    "relato": ""
}}
"""


def extrair_json_resposta(texto: str) -> dict:
    texto_bruto = (texto or "").strip()
    inicio = texto_bruto.find("{")
    fim = texto_bruto.rfind("}")
    if inicio == -1 or fim == -1 or fim <= inicio:
        raise ValueError("A IA não retornou um JSON válido.")

    texto_json = texto_bruto[inicio : fim + 1]
    try:
        return json.loads(texto_json)
    except json.JSONDecodeError as erro:
        try:
            return json.loads(texto_json, strict=False)
        except json.JSONDecodeError:
            pass
        trecho = texto_json[:500]
        raise ValueError(f"JSON inválido retornado pela IA: {erro}. Trecho recebido: {trecho}") from erro


def processar_atendimento_completo(
    arquivos_audio_temp: list[Path],
    contexto_manual: str = "",
    imagens_equipamentos_agres: list[Path] | None = None,
) -> dict:
    materiais_para_ia = []
    arquivos_api = []

    try:
        def enviar_arquivo_para_ia(caminho: Path, rotulo: str):
            temp_file = genai_client.files.upload(file=str(caminho))
            arquivos_api.append(temp_file)
            for _ in range(60):
                estado = str(getattr(getattr(temp_file, "state", None), "name", getattr(temp_file, "state", ""))).upper()
                if "PROCESSING" not in estado:
                    break
                time.sleep(1)
                temp_file = genai_client.files.get(name=temp_file.name)
            estado_final = str(getattr(getattr(temp_file, "state", None), "name", getattr(temp_file, "state", ""))).upper()
            if "FAILED" in estado_final or "PROCESSING" in estado_final:
                raise ValueError(f"{rotulo} {caminho.name} não pôde ser preparado pela IA. Exporte e tente novamente.")
            materiais_para_ia.append(temp_file)

        for audio in arquivos_audio_temp:
            enviar_arquivo_para_ia(audio, "O áudio")

        imagens_validas = normalizar_imagens_equipamentos_ia(imagens_equipamentos_agres)
        if imagens_validas:
            materiais_para_ia.append(
                "\nIMAGENS PRIORITÁRIAS PARA IDENTIFICAÇÃO DOS EQUIPAMENTOS AGRES:\n"
                "As próximas imagens vieram dos campos '2. Cabeçalho do Relatório - Equipamento Agres/Informações do Equipamento' "
                "e '3. Evidências Fotográficas - Equipamento Agres/Configurações'. Use-as como fonte primária para OCR técnico.\n"
                "Extraia somente informações legíveis de equipamentos Agres: modelo da tela/equipamento, número de série, aplicação/software da tela, sistema, ECU, SW, HW, compensador, ANP21/ANP40, versões e demais identificadores técnicos."
            )
        for indice, imagem in enumerate(imagens_validas, start=1):
            materiais_para_ia.append(
                f"\nIMAGEM PRIORITÁRIA {indice}/{len(imagens_validas)} - {imagem['origem']}:\n"
                "Leia textos, menus, etiquetas, telas de versão, QR Code/identificação e plaquetas. "
                "Se houver dados de tela, ECU ou compensador, preencher o campo 'equipamentos' com os rótulos padronizados."
            )
            enviar_arquivo_para_ia(imagem["caminho"], "A imagem prioritária de equipamento Agres")

        resposta = genai_client.models.generate_content(
            model=MODELO_GEMINI,
            contents=[montar_prompt(contexto_manual)] + materiais_para_ia,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.2,
                max_output_tokens=16384,
            ),
        )
        dados = extrair_json_resposta(resposta.text)
        return normalizar_dados_relatorio(dados)
    except Exception as erro:
        raise Exception(f"Erro na interpretação técnica dos dados: {erro}") from erro
    finally:
        for arquivo in arquivos_api:
            try:
                genai_client.files.delete(name=arquivo.name)
            except Exception:
                pass


# ==========================================
# 4. Documento Word
# ==========================================
def imagem_docx(doc: DocxTemplate, caminho, largura_mm: int):
    if caminho:
        return InlineImage(doc, str(caminho), width=Mm(largura_mm))
    return ""


def limpar_nome_arquivo(texto: str) -> str:
    primeira_linha = next((linha for linha in limpar_texto(texto).split("\n") if linha.strip()), "Atendimento")
    nome = re.sub(r'[\\/*?:"<>|]', "", primeira_linha)
    nome = re.sub(r"\s+", "_", nome).strip("_")
    return (nome or "Atendimento")[:80]


def limpar_nome_relatorio(texto: str) -> str:
    nome = limpar_texto(texto)
    nome = nome.replace("\n", " ")
    nome = re.sub(r'[\\/*?:"<>|]', "", nome)
    nome = re.sub(r"\s*-\s*", " - ", nome)
    nome = re.sub(r"\s+", " ", nome).strip(" .-_")
    return (nome or "RELATÓRIO DE ATENDIMENTO")[:150]


def componente_nome_arquivo(texto: str, padrao: str) -> str:
    nome = limpar_texto(texto).replace("\n", " ")
    nome = re.sub(r'[\\/*?:"<>|]', "", nome)
    nome = re.sub(r"\s*&\s*", "_E_", nome)
    nome = re.sub(r"[^0-9A-Za-zÀ-ÿ]+", "_", nome)
    nome = re.sub(r"_+", "_", nome).strip("_")
    return (nome or padrao).upper()


def data_para_nome_arquivo(data_visita: str, usar_data_atual: bool = True) -> str:
    texto = limpar_texto(data_visita)
    candidatos: list[date] = []

    def adicionar_data(ano: str | int, mes: str | int, dia: str | int) -> None:
        try:
            candidatos.append(date(int(ano), int(mes), int(dia)))
        except ValueError:
            return

    for match in re.finditer(r"\b\d{1,2}\s*(?:a|até|-)\s*(\d{1,2})/(\d{1,2})/(\d{4})\b", texto, flags=re.I):
        dia_final, mes, ano = match.groups()
        adicionar_data(ano, mes, dia_final)

    for match in re.finditer(r"\b\d{1,2}/\d{1,2}\s*(?:a|até|-)\s*(\d{1,2})/(\d{1,2})/(\d{4})\b", texto, flags=re.I):
        dia_final, mes_final, ano = match.groups()
        adicionar_data(ano, mes_final, dia_final)

    for match in re.finditer(r"\b(\d{1,2})/(\d{1,2})/(\d{4})\b", texto):
        dia, mes, ano = match.groups()
        adicionar_data(ano, mes, dia)

    for match in re.finditer(r"\b(\d{4})-(\d{1,2})-(\d{1,2})\b", texto):
        ano, mes, dia = match.groups()
        adicionar_data(ano, mes, dia)

    meses = {
        "janeiro": 1,
        "fevereiro": 2,
        "marco": 3,
        "março": 3,
        "abril": 4,
        "maio": 5,
        "junho": 6,
        "julho": 7,
        "agosto": 8,
        "setembro": 9,
        "outubro": 10,
        "novembro": 11,
        "dezembro": 12,
    }
    padrao_intervalo_extenso = (
        r"\b\d{1,2}\s*(?:a|até|-)\s*(\d{1,2})\s*(?:de\s+)?"
        r"(janeiro|fevereiro|mar[cç]o|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)"
        r"\s*(?:de\s+)?(\d{4})\b"
    )
    for match in re.finditer(padrao_intervalo_extenso, texto, flags=re.I):
        dia, mes_nome, ano = match.groups()
        mes = meses[mes_nome.lower().replace("ç", "c")]
        adicionar_data(ano, mes, dia)

    if candidatos:
        return max(candidatos).strftime("%Y%m%d")

    return data_atual_brasil().strftime("%Y%m%d") if usar_data_atual else ""


def data_final_pacote(texto: str) -> str:
    bruto = limpar_texto(texto)
    for candidato in re.findall(r"(?<!\d)(\d{8})(?!\d)", bruto):
        try:
            return datetime.strptime(candidato, "%Y%m%d").strftime("%Y%m%d")
        except ValueError:
            continue
    if bruto:
        candidatos_iso = re.findall(r"(?<!\d)(\d{4})-(\d{2})-(\d{2})(?!\d)", bruto)
        for ano, mes, dia in candidatos_iso:
            try:
                return date(int(ano), int(mes), int(dia)).strftime("%Y%m%d")
            except ValueError:
                continue
    return ""


def data_pacote_ou_texto(*valores: str) -> str:
    for valor in valores:
        data_pacote = data_final_pacote(valor)
        if data_pacote:
            return data_pacote
        data_texto = data_para_nome_arquivo(valor, usar_data_atual=False)
        if data_texto:
            return data_texto
    return ""


def formatar_data_pacote(data_texto: str) -> str:
    data_normalizada = data_final_pacote(data_texto)
    if not data_normalizada:
        return ""
    return datetime.strptime(data_normalizada, "%Y%m%d").strftime("%d/%m/%Y")


def formatar_periodo_atendimento(data_inicio: str, data_final: str) -> str:
    inicio = data_final_pacote(data_inicio)
    fim = data_final_pacote(data_final)
    if not inicio and not fim:
        return ""
    if not inicio or inicio == fim:
        return formatar_data_pacote(fim or inicio)
    data_ini = datetime.strptime(inicio, "%Y%m%d").date()
    data_fim = datetime.strptime(fim, "%Y%m%d").date()
    if data_ini > data_fim:
        data_ini, data_fim = data_fim, data_ini
    if data_ini.year == data_fim.year and data_ini.month == data_fim.month:
        return f"{data_ini.day:02d} a {data_fim.strftime('%d/%m/%Y')}"
    if data_ini.year == data_fim.year:
        return f"{data_ini.strftime('%d/%m')} a {data_fim.strftime('%d/%m/%Y')}"
    return f"{data_ini.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}"


UF_NOMES = {
    "acre": "AC",
    "alagoas": "AL",
    "amapa": "AP",
    "amazonas": "AM",
    "bahia": "BA",
    "ceara": "CE",
    "distrito federal": "DF",
    "espirito santo": "ES",
    "goias": "GO",
    "maranhao": "MA",
    "mato grosso": "MT",
    "mato grosso do sul": "MS",
    "minas gerais": "MG",
    "para": "PA",
    "paraiba": "PB",
    "parana": "PR",
    "pernambuco": "PE",
    "piaui": "PI",
    "rio de janeiro": "RJ",
    "rio grande do norte": "RN",
    "rio grande do sul": "RS",
    "rondonia": "RO",
    "roraima": "RR",
    "santa catarina": "SC",
    "sao paulo": "SP",
    "sergipe": "SE",
    "tocantins": "TO",
    "argentina": "ARGENTINA",
    "paraguai": "PARAGUAI",
    "uruguai": "URUGUAI",
}


def normalizar_uf_nome(texto: str) -> str:
    uf = limpar_texto(texto).strip(" .,-/")
    if re.fullmatch(r"[A-Za-z]{2}", uf):
        return uf.upper()
    return UF_NOMES.get(normalizar_busca(uf), uf.upper())


def limpar_cidade_nome(texto: str) -> str:
    cidade = limpar_texto(texto)
    if "," in cidade:
        cidade = cidade.rsplit(",", 1)[-1].strip()
    cidade = re.sub(r"^(cidade|local|localização|localizacao|propriedade|fazenda|revenda)\s*:\s*", "", cidade, flags=re.I)
    cidade = re.sub(r"^(?:localizado|localizada)\s+em\s+", "", cidade, flags=re.I)
    cidade = re.sub(r"^\bem\s+", "", cidade, flags=re.I)
    return cidade.strip(" .,-/")


def extrair_cidade_uf(cliente_local: str) -> tuple[str, str]:
    texto = limpar_texto(cliente_local)
    linhas = [linha.strip() for linha in texto.split("\n") if linha.strip()]

    for linha in linhas:
        match = re.search(r"(?:cidade(?:/uf)?|local(?: cliente)?|propriedade)\s*:\s*([^,\n/]+?)\s*/\s*([A-Za-z]{2})\b", linha, re.I)
        if match:
            return match.group(1).strip(), match.group(2).strip()

    for linha in linhas:
        match = re.search(r"\b([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s.'-]{2,}?)\s*/\s*([A-Za-z]{2})\b", linha)
        if match and not linha.lower().startswith(("http", "www")):
            return limpar_cidade_nome(match.group(1)), normalizar_uf_nome(match.group(2))

    for linha in linhas:
        match = re.search(
            r"\b([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s.'-]{2,}?)\s*,\s*"
            r"(acre|alagoas|amapá|amapa|amazonas|bahia|ceará|ceara|distrito federal|espírito santo|espirito santo|"
            r"goiás|goias|maranhão|maranhao|mato grosso|mato grosso do sul|minas gerais|pará|para|paraíba|paraiba|"
            r"paraná|parana|pernambuco|piauí|piaui|rio de janeiro|rio grande do norte|rio grande do sul|rondônia|rondonia|"
            r"roraima|santa catarina|são paulo|sao paulo|sergipe|tocantins|argentina|paraguai|uruguai)\b",
            linha,
            re.I,
        )
        if match and not linha.lower().startswith(("http", "www")):
            return limpar_cidade_nome(match.group(1)), normalizar_uf_nome(match.group(2))

    for linha in linhas:
        match = re.search(r"(?:^|[,;/])\s*([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s.'-]{2,}?)([A-Z]{2})\b", linha)
        if match and not linha.lower().startswith(("http", "www")):
            cidade = limpar_cidade_nome(match.group(1))
            uf = normalizar_uf_nome(match.group(2))
            if cidade and uf:
                return cidade, uf

    cidade = ""
    uf = ""
    for linha in linhas:
        match_cidade = re.search(r"cidade(?: revenda)?\s*:\s*(.+)", linha, re.I)
        if match_cidade and not cidade:
            cidade = match_cidade.group(1).strip()
        match_uf = re.search(r"\b(?:uf|estado)\s*:\s*([A-Za-z]{2})\b", linha, re.I)
        if match_uf:
            uf = match_uf.group(1).strip()

    if cidade and "/" in cidade:
        partes = [parte.strip() for parte in cidade.rsplit("/", 1)]
        if len(partes) == 2 and re.fullmatch(r"[A-Za-z]{2}", partes[1]):
            return partes[0], partes[1]

    if cidade and not uf:
        match_colado = re.search(r"(.+?)([A-Za-z]{2})$", cidade)
        if match_colado and len(match_colado.group(1).strip()) > 3:
            cidade = match_colado.group(1).strip(" ,-/")
            uf = match_colado.group(2)

    return cidade or "LOCAL NÃO INFORMADO", uf or "UF"


def tipo_atendimento_para_nome(dados: dict) -> str:
    tipos = normalizar_tipos_atendimento(dados.get("tipos_atendimento") or dados.get("tipo_atendimento")) or escolher_tipos_atendimento(dados)
    return "_".join(TIPOS_ATENDIMENTO_ARQUIVO.get(tipo, "ATENDIMENTO") for tipo in tipos)


def tipos_atendimento_para_texto(valor) -> str:
    tipos = normalizar_tipos_atendimento(valor, [])
    return " + ".join(TIPOS_ATENDIMENTO.get(tipo, tipo) for tipo in tipos) if tipos else "Não informado"


def equipamento_para_nome(dados: dict) -> str:
    texto = normalizar_busca(
        "\n".join(
            [
                dados.get("objetivos", ""),
                dados.get("equipamentos", ""),
                dados.get("configuracoes", ""),
                dados.get("calibracoes", ""),
                dados.get("relato", ""),
                dados.get("contexto_coleta", ""),
            ]
        )
    )
    sistemas_principais: list[tuple[int, str]] = []
    produtos_secundarios: list[tuple[int, str]] = []
    telas: list[tuple[int, str]] = []
    acessorios: list[tuple[int, str]] = []

    def adicionar_encontrados(destino: list[tuple[int, str]], padrao: str, formatador) -> None:
        for match in re.finditer(padrao, texto, flags=re.I):
            nome = formatador(match)
            if nome and nome not in {item[1] for item in destino}:
                destino.append((match.start(), nome))

    adicionar_encontrados(
        sistemas_principais,
        r"\bgeo\s*nave\s*([0-9]{1,3})(?:\s*(OFP|FPS|FP|C|S))?\b",
        lambda match: f"GEONAVE{match.group(1)}{(match.group(2) or '').upper()}",
    )
    adicionar_encontrados(
        sistemas_principais,
        r"\bagro\s*nave\s*([0-9]{1,3})\b|\bagn\s*([0-9]{1,3})\b",
        lambda match: f"AGRONAVE{match.group(1) or match.group(2)}",
    )
    adicionar_encontrados(
        sistemas_principais,
        r"\biso\s*(30|31|32|33|34|35|36|37)\s*(OFP|FP|E|H)?\b",
        lambda match: codigo_modelo_isoview(match.group(1), match.group(2) or ""),
    )
    adicionar_encontrados(
        sistemas_principais,
        r"\biso\s*([0-9]{1,3})(?:\s*(OFP|FPS|FP|C|S))?\b",
        lambda match: f"ISO{match.group(1)}{(match.group(2) or '').upper()}",
    )

    produtos_secundarios_config = [
        ("ISOBOX SPRAYER", r"\bisobox\s+sprayer\b"),
        ("ISOPILOT", r"\bisopilot\b"),
    ]
    for nome, padrao in produtos_secundarios_config:
        adicionar_encontrados(produtos_secundarios, padrao, lambda _match, nome=nome: nome)

    adicionar_encontrados(telas, r"\bisoview\b", lambda _match: "ISOVIEW")

    acessorios_config = [
        ("ANP40", r"\banp\s*40\b"),
        ("ANP21", r"\banp\s*21\b"),
    ]
    for nome, padrao in acessorios_config:
        adicionar_encontrados(acessorios, padrao, lambda _match, nome=nome: nome)

    for grupo in (sistemas_principais, produtos_secundarios, telas, acessorios):
        if grupo:
            return " ".join(nome for _, nome in sorted(grupo))

    return "EQUIPAMENTO AGRES"


def tecnico_para_nome_arquivo(dados: dict) -> str:
    texto = limpar_texto(dados.get("tecnico_agres_responsavel", ""))
    texto = re.sub(r"^(técnicos?|tecnicos?|responsáveis?|responsaveis?)\s*:\s*", "", texto, flags=re.I)
    partes = [
        parte.strip()
        for parte in re.split(r"\s*(?:,|;|\n|/|&|\be\b)\s*", texto, flags=re.I)
        if parte.strip()
    ]
    for parte in partes:
        primeiro_nome = parte.split()[0] if parte.split() else ""
        if primeiro_nome and normalizar_busca(primeiro_nome) not in {"tecnico", "tecnica", "agres"}:
            return componente_nome_arquivo(primeiro_nome, "TECNICO")
    return "TECNICO"


def gerar_nome_arquivo_relatorio(dados: dict) -> str:
    data_nome = data_final_pacote(dados.get("data_atendimento_final", "")) or data_para_nome_arquivo(dados.get("data_visita", ""))
    cidade, uf = extrair_cidade_uf(dados.get("cliente_local", ""))
    partes = [
        data_nome,
        "RELATÓRIO",
        "ATIVIDADES",
        componente_nome_arquivo(equipamento_para_nome(dados), "EQUIPAMENTO"),
        tecnico_para_nome_arquivo(dados),
        componente_nome_arquivo(cidade, "LOCAL"),
        componente_nome_arquivo(uf, "UF"),
    ]
    return "_".join(partes)[:150]


def linhas_metadados(texto: str) -> list[str]:
    return [linha.strip() for linha in limpar_texto(texto).split("\n") if linha.strip()]


def separar_metadados_figura(linha: str) -> tuple[str, str, str]:
    partes = [parte.strip() for parte in linha.split("|", 2)]
    titulo = partes[0] if len(partes) >= 1 else ""
    legenda = partes[1] if len(partes) >= 2 else ""
    fonte = partes[2] if len(partes) >= 3 else ""
    return titulo, legenda, fonte


def montar_metadados_figura(categoria: str, indice_foto: int, numero_figura: int, legendas_evidencias: dict) -> dict:
    configuracao = CATEGORIAS_EVIDENCIAS[categoria]
    linhas_categoria = linhas_metadados((legendas_evidencias or {}).get(categoria, ""))
    titulo_manual, legenda_manual, fonte_manual = ("", "", "")

    if indice_foto < len(linhas_categoria):
        titulo_manual, legenda_manual, fonte_manual = separar_metadados_figura(linhas_categoria[indice_foto])

    titulo_base = finalizar_frase(titulo_manual or configuracao["titulo_padrao"])
    if re.match(r"^figura\s+\d+\b", normalizar_busca(titulo_base)):
        titulo = titulo_base
    else:
        titulo = f"Figura {numero_figura} – {titulo_base}"

    legenda_base = finalizar_frase(legenda_manual or configuracao["legenda_padrao"])
    legenda = legenda_base if normalizar_busca(legenda_base).startswith(("legenda:", "nota:")) else f"Legenda: {legenda_base}"

    fonte_base = finalizar_frase(fonte_manual or f"O autor ({data_atual_brasil().year})")
    fonte = fonte_base if normalizar_busca(fonte_base).startswith("fonte:") else f"Fonte: {fonte_base}"

    return {"titulo": titulo, "legenda": legenda, "fonte": fonte}


def nome_assinatura_para_exibicao(texto: str) -> str:
    texto = limpar_texto(texto)
    if not texto or normalizar_busca(texto) == "nao informado":
        return "Nome: ______________________________________"
    return f"Nome: {texto}"


def formatar_run_assinatura(run, tamanho: int = 10, bold: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(tamanho)
    run.font.bold = bold


def formatar_paragrafo_assinatura(paragraph, space_after: int = 0) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.keep_together = True


def remover_bordas_tabela(table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    for borda_existente in tbl_pr.xpath("./w:tblBorders"):
        tbl_pr.remove(borda_existente)

    bordas = OxmlElement("w:tblBorders")
    for nome_borda in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borda = OxmlElement(f"w:{nome_borda}")
        borda.set(qn("w:val"), "nil")
        bordas.append(borda)
    tbl_pr.append(bordas)


def preencher_celula_assinatura(cell, titulo: str, nome: str, documento: str, caminho_assinatura: Path | None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""

    paragrafo_imagem = cell.paragraphs[0]
    formatar_paragrafo_assinatura(paragrafo_imagem, space_after=1)
    assinatura_inserida = False
    if caminho_assinatura and Path(caminho_assinatura).exists():
        try:
            paragrafo_imagem.add_run().add_picture(str(caminho_assinatura), width=Mm(TAM_ASSINATURA))
            assinatura_inserida = True
        except Exception:
            assinatura_inserida = False

    if not assinatura_inserida:
        run_espaco = paragrafo_imagem.add_run("\n\n")
        formatar_run_assinatura(run_espaco, tamanho=10)

    paragrafo_linha = cell.add_paragraph()
    formatar_paragrafo_assinatura(paragrafo_linha, space_after=1)
    formatar_run_assinatura(paragrafo_linha.add_run("______________________________________"), tamanho=10)

    titulo_limpo = limpar_papel_assinatura(titulo)
    if titulo_limpo:
        paragrafo_titulo = cell.add_paragraph()
        formatar_paragrafo_assinatura(paragrafo_titulo, space_after=1)
        formatar_run_assinatura(paragrafo_titulo.add_run(titulo_limpo), tamanho=10, bold=True)

    paragrafo_nome = cell.add_paragraph()
    formatar_paragrafo_assinatura(paragrafo_nome, space_after=1)
    formatar_run_assinatura(paragrafo_nome.add_run(nome_assinatura_para_exibicao(nome)), tamanho=9)


def inserir_elemento_antes_paragrafo(elemento, paragraph) -> None:
    parent = paragraph._p.getparent()
    elemento_parent = elemento.getparent()
    if elemento_parent is not None:
        elemento_parent.remove(elemento)
    parent.insert(parent.index(paragraph._p), elemento)


def criar_elemento_quebra_pagina():
    quebra = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    quebra.append(run)
    return quebra


def elemento_paragrafo_tem_conteudo(elemento) -> bool:
    texto = "".join(elemento.xpath(".//*[local-name()='t']/text()")).strip()
    if texto:
        return True
    return bool(elemento.xpath(".//*[local-name()='drawing' or local-name()='pict' or local-name()='br']"))


def elemento_tem_quebra_pagina(elemento) -> bool:
    return bool(elemento.xpath(".//*[local-name()='br' and @*[local-name()='type']='page']"))


def remover_paragrafos_vazios_antes(paragraph) -> None:
    parent = paragraph._p.getparent()
    anterior = paragraph._p.getprevious()
    while anterior is not None and anterior.tag.endswith("}p") and not elemento_paragrafo_tem_conteudo(anterior):
        remover = anterior
        anterior = anterior.getprevious()
        parent.remove(remover)


def paragrafo_tem_quebra_pagina_imediata_antes(paragraph) -> bool:
    anterior = paragraph._p.getprevious()
    if anterior is None or not anterior.tag.endswith("}p"):
        return False
    return elemento_tem_quebra_pagina(anterior)


def inserir_quebra_pagina_antes_paragrafo(paragraph) -> None:
    remover_paragrafos_vazios_antes(paragraph)
    anterior = paragraph._p.getprevious()
    if anterior is not None and elemento_tem_quebra_pagina(anterior):
        paragraph.paragraph_format.page_break_before = False
        return
    if not paragrafo_tem_quebra_pagina_imediata_antes(paragraph):
        paragraph.paragraph_format.page_break_before = True


def impedir_quebra_linha_tabela(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.widow_control = True


def inserir_assinaturas_docx(caminho_docx: Path, dados: dict, caminhos_assinaturas: list[dict] | dict | None = None) -> None:
    documento = Document(str(caminho_docx))
    if not caminhos_assinaturas:
        return

    if isinstance(caminhos_assinaturas, dict):
        assinaturas_documento = []
        for chave, configuracao in ASSINATURAS_RESPONSAVEIS.items():
            caminho = caminhos_assinaturas.get(chave)
            if caminho or dados.get(configuracao["campo"]):
                assinaturas_documento.append(
                    {
                        "nome": dados.get(configuracao["campo"], ""),
                        "representa": "",
                        "caminho": caminho,
                    }
                )
    else:
        assinaturas_documento = [
            {
                **assinatura,
                "nome": texto_assinatura(
                    assinatura,
                    "nome",
                    "responsavel",
                    "responsável",
                    "nome_responsavel",
                    "nomeResponsavel",
                ),
                "representa": limpar_papel_assinatura(
                    texto_assinatura(
                        assinatura,
                        "representa",
                        "funcao",
                        "função",
                        "cargo",
                        "papel",
                        "role",
                    )
                ),
            }
            for indice, assinatura in enumerate(caminhos_assinaturas or [], start=1)
            if isinstance(assinatura, dict)
        ]

    assinaturas_documento = [
        assinatura
        for assinatura in assinaturas_documento
        if assinatura.get("nome") or assinatura.get("representa") or assinatura.get("caminho")
    ]
    if not assinaturas_documento:
        return

    titulo = documento.add_paragraph()
    titulo.paragraph_format.page_break_before = True
    titulo.paragraph_format.keep_with_next = True
    titulo.paragraph_format.keep_together = True
    titulo.paragraph_format.space_before = Pt(0)
    titulo.paragraph_format.space_after = Pt(6)
    run_titulo = titulo.add_run("Assinaturas:")
    run_titulo.font.name = "Arial"
    run_titulo.font.size = Pt(11)
    run_titulo.font.bold = True

    linhas = (len(assinaturas_documento) + 1) // 2
    tabela = documento.add_table(rows=linhas, cols=2)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    remover_bordas_tabela(tabela)

    for indice, assinatura in enumerate(assinaturas_documento):
        linha = indice // 2
        coluna = indice % 2
        preencher_celula_assinatura(
            tabela.cell(linha, coluna),
            assinatura.get("representa", ""),
            assinatura.get("nome", ""),
            assinatura.get("documento", ""),
            assinatura.get("caminho"),
        )

    for row in tabela.rows:
        impedir_quebra_linha_tabela(row)

    documento.save(str(caminho_docx))


def gerar_docx(
    dados_json: dict,
    dicionario_evidencias: dict,
    caminhos_cabecalho: dict,
    pasta_saida: Path,
    legendas_evidencias: dict | None = None,
    caminhos_assinaturas: dict | None = None,
) -> Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Modelo não encontrado: {TEMPLATE_PATH.name}")

    doc = DocxTemplate(str(TEMPLATE_PATH))
    dados_render = dict(dados_json)

    dados_render["img_info_equipamento"] = imagem_docx(doc, caminhos_cabecalho.get("info_equip"), TAM_PLAQUETA)
    dados_render["img_maquina"] = imagem_docx(doc, caminhos_cabecalho.get("maquina"), TAM_MAQUINA)
    dados_render["img_implemento"] = imagem_docx(doc, caminhos_cabecalho.get("implemento"), TAM_MAQUINA)

    contador_figura = 1
    for categoria in CATEGORIAS_EVIDENCIAS:
        lista_fotos = []
        for indice_foto, foto_path in enumerate(dicionario_evidencias.get(categoria, [])):
            metadados = montar_metadados_figura(categoria, indice_foto, contador_figura, legendas_evidencias or {})
            lista_fotos.append(
                {
                    "titulo": metadados["titulo"],
                    "imagem": InlineImage(doc, str(foto_path), width=Mm(TAM_EVIDENCIA)),
                    "fonte": metadados["fonte"],
                    "legenda": metadados["legenda"],
                }
            )
            contador_figura += 1
        dados_render[categoria] = lista_fotos

    nome_arquivo = f"{gerar_nome_arquivo_relatorio(dados_render)}.docx"
    caminho_saida = pasta_saida / nome_arquivo
    doc.render(dados_render)
    doc.save(str(caminho_saida))
    aplicar_tipo_atendimento_word(caminho_saida, dados_render)
    aplicar_paragrafos_relato_docx(caminho_saida, dados_render.get("relato", ""))
    inserir_assinaturas_docx(caminho_saida, dados_render, caminhos_assinaturas)
    aplicar_formatacao_texto_tecnico(caminho_saida)
    aplicar_paginacao_abnt_figuras(caminho_saida)
    return caminho_saida


def nome_arquivo_seguro(texto: str, padrao: str = "arquivo") -> str:
    nome = normalizar_busca(texto).replace(" ", "_")
    nome = re.sub(r"[^a-z0-9_.-]+", "_", nome).strip("._-")
    return nome[:90] or padrao


def arcname_unico(arcname: str, usados: set[str]) -> str:
    if arcname not in usados:
        usados.add(arcname)
        return arcname

    caminho = Path(arcname)
    base = str(caminho.with_suffix(""))
    sufixo = caminho.suffix
    contador = 2
    while True:
        candidato = f"{base}_{contador}{sufixo}"
        if candidato not in usados:
            usados.add(candidato)
            return candidato
        contador += 1


def adicionar_arquivo_zip(zip_out: zipfile.ZipFile, caminho: Path | None, arcname: str, usados: set[str]) -> None:
    if not caminho:
        return
    caminho = Path(caminho)
    if caminho.exists() and caminho.is_file():
        zip_out.write(caminho, arcname_unico(arcname.replace("\\", "/"), usados))


def adicionar_fotos_atendimento_zip(
    zip_out: zipfile.ZipFile,
    caminhos_cabecalho: dict,
    dicionario_evidencias: dict,
    fotos_atendimento: list[Path] | None,
    legendas_evidencias: dict | None,
    caminhos_assinaturas: dict | None,
    usados: set[str],
) -> None:
    indice_zip = 1

    def adicionar_foto_unica(caminho: Path | None, nome_base: str, extensao_padrao: str = ".jpg") -> None:
        nonlocal indice_zip
        if not caminho:
            return
        caminho = Path(caminho)
        if not caminho.exists():
            return
        extensao = caminho.suffix.lower() or extensao_padrao
        nome = nome_arquivo_seguro(nome_base, f"foto_{indice_zip:02d}")
        adicionar_arquivo_zip(
            zip_out,
            caminho,
            f"FOTOS DO ATENDIMENTO/{indice_zip:02d}_{nome}{extensao}",
            usados,
        )
        indice_zip += 1

    cabecalho_nomes = {
        "info_equip": "informacoes_equipamento",
        "maquina": "maquina",
        "implemento": "implemento",
    }
    for chave, nome in cabecalho_nomes.items():
        adicionar_foto_unica(caminhos_cabecalho.get(chave), nome)

    contador = 1
    for categoria, configuracao in CATEGORIAS_EVIDENCIAS.items():
        for indice, caminho in enumerate(dicionario_evidencias.get(categoria, [])):
            caminho = Path(caminho)
            if not caminho.exists():
                continue
            metadados = montar_metadados_figura(categoria, indice, contador, legendas_evidencias or {})
            adicionar_foto_unica(caminho, f"figura_{contador:02d}_{metadados['titulo']}")
            contador += 1

    for caminho in fotos_atendimento or []:
        caminho = Path(caminho)
        if not caminho.exists():
            continue
        adicionar_foto_unica(caminho, caminho.stem)

    if isinstance(caminhos_assinaturas, dict):
        for chave, caminho in (caminhos_assinaturas or {}).items():
            adicionar_foto_unica(caminho, f"assinatura_{chave}", ".png")
    else:
        for indice, assinatura in enumerate(caminhos_assinaturas or [], start=1):
            nome_base = assinatura.get("nome") or assinatura.get("representa") or f"assinatura_{indice}"
            adicionar_foto_unica(assinatura.get("caminho"), f"assinatura_{indice}_{nome_base}", ".png")


def gerar_pacote_relatorio(
    arquivo_docx: Path,
    dicionario_evidencias: dict,
    fotos_atendimento: list[Path] | None,
    caminhos_cabecalho: dict,
    pasta_saida: Path,
    legendas_evidencias: dict | None,
    caminhos_assinaturas: dict | None,
) -> Path:
    nome_base = arquivo_docx.stem
    caminho_zip = pasta_saida / f"{nome_base} - WORD E FOTOS.zip"
    usados = set()

    with zipfile.ZipFile(caminho_zip, "w", compression=zipfile.ZIP_DEFLATED) as zip_out:
        adicionar_arquivo_zip(zip_out, arquivo_docx, arquivo_docx.name, usados)
        adicionar_fotos_atendimento_zip(
            zip_out,
            caminhos_cabecalho,
            dicionario_evidencias,
            fotos_atendimento,
            legendas_evidencias,
            caminhos_assinaturas,
            usados,
        )

    return caminho_zip


def chave_paragrafo(paragraph) -> str:
    return paragraph._p.getroottree().getpath(paragraph._p)


def iterar_paragrafos_word(parent, vistos=None):
    if vistos is None:
        vistos = set()

    if isinstance(parent, DocumentClass):
        for paragraph in parent.paragraphs:
            chave = chave_paragrafo(paragraph)
            if chave not in vistos:
                vistos.add(chave)
                yield paragraph
        for table in parent.tables:
            yield from iterar_paragrafos_word(table, vistos)
    elif isinstance(parent, Table):
        for row in parent.rows:
            for cell in row.cells:
                yield from iterar_paragrafos_word(cell, vistos)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            chave = chave_paragrafo(paragraph)
            if chave not in vistos:
                vistos.add(chave)
                yield paragraph
        for table in parent.tables:
            yield from iterar_paragrafos_word(table, vistos)


def paragrafo_tem_imagem(paragraph) -> bool:
    return paragraph._p.xpath(".//*[local-name()='drawing' or local-name()='pict']")


def formatar_paragrafo_relato(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    formato = paragraph.paragraph_format
    formato.left_indent = Mm(0)
    formato.right_indent = Mm(0)
    formato.first_line_indent = Mm(12.5)
    formato.space_before = Pt(0)
    formato.space_after = Pt(0)
    formato.line_spacing = 1.5
    formato.widow_control = True
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def dividir_relato_em_paragrafos(texto: str) -> list[str]:
    texto = limpar_relato_narrativo(texto)
    paragrafos = [paragrafo.strip() for paragrafo in re.split(r"\n{2,}", texto) if paragrafo.strip()]
    if len(paragrafos) > 1 or not paragrafos or len(paragrafos[0]) < 900:
        return paragrafos

    frases = re.split(r"(?<=[.!?])\s+(?=[A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9])", paragrafos[0])
    frases = [frase.strip() for frase in frases if frase.strip()]
    if len(frases) < 6:
        return paragrafos

    alvo = max(2, min(4, round(len(frases) / 3)))
    blocos = []
    for indice in range(0, len(frases), alvo):
        blocos.append(" ".join(frases[indice : indice + alvo]).strip())
    return blocos


def aplicar_paragrafos_relato_docx(caminho_docx: Path, relato: str) -> None:
    paragrafos_relato = dividir_relato_em_paragrafos(relato)
    if len(paragrafos_relato) <= 1:
        return

    documento = Document(caminho_docx)
    relato_normalizado = limpar_texto(relato)
    alvo = None
    for paragraph in iterar_paragrafos_word(documento):
        if paragrafo_tem_imagem(paragraph):
            continue
        if limpar_texto(paragraph.text) == relato_normalizado:
            alvo = paragraph
            break

    if alvo is None:
        trecho = limpar_texto(paragrafos_relato[0])[:80]
        for paragraph in iterar_paragrafos_word(documento):
            if trecho and trecho in limpar_texto(paragraph.text):
                alvo = paragraph
                break

    if alvo is None:
        return

    novos_paragrafos = []
    for paragrafo_texto in paragrafos_relato:
        novo = alvo.insert_paragraph_before()
        novo.style = alvo.style
        novo.add_run(paragrafo_texto)
        formatar_paragrafo_relato(novo)
        novos_paragrafos.append(novo)

    remover_paragrafo_word(alvo)
    documento.save(caminho_docx)


def aplicar_formatacao_texto_tecnico(caminho_docx: Path) -> None:
    documento = Document(caminho_docx)
    preservar_centralizado = {
        "RELATÓRIO DE ATENDIMENTO/ATIVIDADES",
        "SUPORTE",
        "INSTALAÇÃO",
        "TREINAMENTO",
        "VALIDAÇÃO/HOMOLOGAÇÃO",
        "X",
    }

    for paragraph in iterar_paragrafos_word(documento):
        texto = limpar_texto(paragraph.text)
        if not texto or paragrafo_tem_imagem(paragraph):
            continue
        texto_superior = texto.upper()
        if (
            texto_superior in preservar_centralizado
            or texto.startswith("Figura ")
            or texto.startswith("Fonte:")
            or texto.startswith("Legenda:")
        ):
            continue
        if paragraph.paragraph_format.first_line_indent:
            formatar_paragrafo_relato(paragraph)
            continue

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.widow_control = True

    documento.save(caminho_docx)


def formatar_paragrafo_figura(paragraph, keep_with_next: bool, tamanho_fonte: int | None = 10) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formato = paragraph.paragraph_format
    formato.keep_together = True
    formato.keep_with_next = keep_with_next
    formato.widow_control = True
    formato.line_spacing = 1
    formato.space_before = Pt(2)
    formato.space_after = Pt(2)
    if tamanho_fonte:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(tamanho_fonte)


def limpar_paragrafo_word(paragraph) -> None:
    p = paragraph._p
    propriedades = p.pPr
    for elemento in list(p):
        if propriedades is not None and elemento is propriedades:
            continue
        p.remove(elemento)


def remover_paragrafo_word(paragraph) -> None:
    elemento = paragraph._element
    pai = elemento.getparent()
    if pai is not None:
        pai.remove(elemento)


def copiar_runs_word(origem, destino) -> None:
    for elemento in origem._p:
        if (
            elemento.tag.endswith("}r") or elemento.tag.endswith("}hyperlink")
        ) and elemento.xpath(".//*[local-name()='drawing' or local-name()='pict']"):
            destino._p.append(deepcopy(elemento))


def adicionar_run_formatado(paragraph, texto: str, tamanho_fonte: int = 10):
    run = paragraph.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(tamanho_fonte)
    return run


def consolidar_bloco_figura(
    titulo_paragraph,
    imagem_paragraph,
    fonte_paragraph,
    legenda_paragraph,
    inserir_quebra_pagina: bool,
) -> None:
    titulo = limpar_texto(titulo_paragraph.text)
    fonte = limpar_texto(fonte_paragraph.text)
    legenda = limpar_texto(legenda_paragraph.text)
    imagem_runs = [
        deepcopy(elemento)
        for elemento in imagem_paragraph._p
        if (
            elemento.tag.endswith("}r") or elemento.tag.endswith("}hyperlink")
        ) and elemento.xpath(".//*[local-name()='drawing' or local-name()='pict']")
    ]

    limpar_paragrafo_word(titulo_paragraph)
    adicionar_run_formatado(titulo_paragraph, titulo, tamanho_fonte=10)
    titulo_paragraph.add_run().add_break()
    for imagem_run in imagem_runs:
        titulo_paragraph._p.append(imagem_run)
    titulo_paragraph.add_run().add_break()
    adicionar_run_formatado(titulo_paragraph, fonte, tamanho_fonte=8)
    titulo_paragraph.add_run().add_break()
    adicionar_run_formatado(titulo_paragraph, legenda, tamanho_fonte=8)
    titulo_paragraph.add_run().add_break()
    titulo_paragraph.add_run().add_break()
    if inserir_quebra_pagina:
        titulo_paragraph.add_run().add_break(WD_BREAK.PAGE)

    formatar_paragrafo_figura(titulo_paragraph, keep_with_next=False, tamanho_fonte=None)
    titulo_paragraph.paragraph_format.space_before = Pt(6)
    titulo_paragraph.paragraph_format.space_after = Pt(0)


def consolidar_figuras_em_blocos_unicos(paragrafos: list) -> None:
    total_figuras = sum(1 for paragraph in paragrafos if limpar_texto(paragraph.text).startswith("Figura "))
    removidos = set()
    numero_figura = 0

    for indice, paragraph in enumerate(paragrafos):
        if id(paragraph) in removidos:
            continue

        texto = limpar_texto(paragraph.text)
        if not texto.startswith("Figura "):
            continue

        numero_figura += 1
        imagem_paragraph = None
        fonte_paragraph = None
        legenda_paragraph = None
        paragrafos_para_remover = []

        for proximo in paragrafos[indice + 1 :]:
            proximo_texto = limpar_texto(proximo.text)
            proximo_tem_imagem = bool(paragrafo_tem_imagem(proximo))

            if proximo_texto.startswith("Figura "):
                break

            paragrafos_para_remover.append(proximo)
            if proximo_tem_imagem and imagem_paragraph is None:
                imagem_paragraph = proximo
            elif proximo_texto.startswith("Fonte:"):
                fonte_paragraph = proximo
            elif proximo_texto.startswith(("Legenda:", "Nota:")):
                legenda_paragraph = proximo
                break

        if not (imagem_paragraph and fonte_paragraph and legenda_paragraph):
            formatar_paragrafo_figura(paragraph, keep_with_next=True)
            continue

        for proximo in paragrafos[indice + 1 + len(paragrafos_para_remover) :]:
            proximo_texto = limpar_texto(proximo.text)
            if proximo_texto or paragrafo_tem_imagem(proximo):
                break
            paragrafos_para_remover.append(proximo)

        consolidar_bloco_figura(
            paragraph,
            imagem_paragraph,
            fonte_paragraph,
            legenda_paragraph,
            numero_figura % FIGURAS_POR_PAGINA == 0 and numero_figura < total_figuras,
        )

        for paragrafo_remover in paragrafos_para_remover:
            removidos.add(id(paragrafo_remover))
            remover_paragrafo_word(paragrafo_remover)


def aplicar_paginacao_abnt_figuras(caminho_docx: Path) -> None:
    documento = Document(str(caminho_docx))
    paragrafos = list(iterar_paragrafos_word(documento))
    consolidar_figuras_em_blocos_unicos(paragrafos)

    paragrafos = list(iterar_paragrafos_word(documento))
    quebra_fotos_inserida = False
    for indice, paragraph in enumerate(paragrafos):
        texto = limpar_texto(paragraph.text).rstrip(":")
        if texto == "Fotos":
            if not quebra_fotos_inserida:
                inserir_quebra_pagina_antes_paragrafo(paragraph)
                quebra_fotos_inserida = True
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.widow_control = True
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)

            for proximo in paragrafos[indice + 1 :]:
                proximo_texto = limpar_texto(proximo.text)
                if proximo_texto or paragrafo_tem_imagem(proximo):
                    break
                remover_paragrafo_word(proximo)

    paragrafos = list(iterar_paragrafos_word(documento))

    titulos_fotos = {
        "Identificação do Equipamento",
        "Instalação e Chicotes",
        "Configurações",
        "Configurações do Sistema",
        "Outros Registros",
        "Atividades Adicionais",
    }
    titulos_fotos_nova_pagina = {
        "Configurações",
        "Configurações do Sistema",
        "Outros Registros",
        "Atividades Adicionais",
    }
    for indice, paragraph in enumerate(paragrafos):
        texto = limpar_texto(paragraph.text)
        if texto not in titulos_fotos:
            continue
        remover_paragrafos_vazios_antes(paragraph)
        if texto in titulos_fotos_nova_pagina:
            inserir_quebra_pagina_antes_paragrafo(paragraph)
        for proximo in paragrafos[indice + 1 :]:
            proximo_texto = limpar_texto(proximo.text)
            if proximo_texto or paragrafo_tem_imagem(proximo):
                break
            remover_paragrafo_word(proximo)

    paragrafos = list(iterar_paragrafos_word(documento))

    dentro_bloco_figura = False
    numero_figura = 0
    for paragraph in paragrafos:
        texto = limpar_texto(paragraph.text)
        tem_imagem = bool(paragrafo_tem_imagem(paragraph))

        if texto.startswith("Figura ") and tem_imagem:
            numero_figura += 1
            formatar_paragrafo_figura(paragraph, keep_with_next=False)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            dentro_bloco_figura = False
            continue

        if texto.startswith("Figura "):
            dentro_bloco_figura = True
            numero_figura += 1
            formatar_paragrafo_figura(paragraph, keep_with_next=True)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            continue

        if dentro_bloco_figura and tem_imagem and "Fonte:" in texto and (
            "Legenda:" in texto or "Nota:" in texto
        ):
            formatar_paragrafo_figura(paragraph, keep_with_next=False, tamanho_fonte=None)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            dentro_bloco_figura = False
            continue

        if dentro_bloco_figura and tem_imagem:
            formatar_paragrafo_figura(paragraph, keep_with_next=True, tamanho_fonte=None)
            continue

        if dentro_bloco_figura and texto.startswith("Fonte:"):
            formatar_paragrafo_figura(paragraph, keep_with_next=True)
            continue

        if dentro_bloco_figura and (texto.startswith("Legenda:") or texto.startswith("Nota:")):
            formatar_paragrafo_figura(paragraph, keep_with_next=False)
            paragraph.paragraph_format.space_after = Pt(12)
            dentro_bloco_figura = False
            continue

        if dentro_bloco_figura and not texto:
            formatar_paragrafo_figura(paragraph, keep_with_next=True)
            continue

        if texto in titulos_fotos:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.widow_control = True
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)

    documento.save(str(caminho_docx))


def normalizar_imagem_para_docx(conteudo: bytes, caminho_saida: Path, padronizar_figura: bool = False) -> Path:
    if len(conteudo) > MAX_IMAGE_BYTES:
        raise ValueError(f"A imagem excede o limite de {MAX_IMAGE_BYTES // (1024 * 1024)} MB.")
    try:
        with Image.open(BytesIO(conteudo)) as imagem_original:
            if imagem_original.width * imagem_original.height > MAX_IMAGE_PIXELS:
                raise ValueError("A imagem possui resolução excessiva e não pode ser processada com segurança.")
            imagem = ImageOps.exif_transpose(imagem_original)

            if imagem.mode in ("RGBA", "LA") or (imagem.mode == "P" and "transparency" in imagem.info):
                imagem_rgba = imagem.convert("RGBA")
                fundo = Image.new("RGB", imagem_rgba.size, "white")
                fundo.paste(imagem_rgba, mask=imagem_rgba.getchannel("A"))
                imagem = fundo
            else:
                imagem = imagem.convert("RGB")

            if padronizar_figura:
                quadro = Image.new("RGB", FIGURA_CANVAS_PX, "white")
                imagem.thumbnail(FIGURA_CANVAS_PX, Image.Resampling.LANCZOS)
                x = (FIGURA_CANVAS_PX[0] - imagem.width) // 2
                y = FIGURA_CANVAS_PX[1] - imagem.height
                quadro.paste(imagem, (x, y))
                imagem = quadro

            caminho_limpo = caminho_saida.with_suffix(".jpg")
            imagem.save(caminho_limpo, format="JPEG", quality=90, optimize=True, progressive=False)
            return caminho_limpo
    except UnidentifiedImageError as erro:
        raise ValueError("Uma das imagens enviadas não pôde ser lida. Tente reenviar em JPG ou PNG.") from erro
    except OSError as erro:
        raise ValueError("Uma das imagens está incompleta ou com metadados inválidos. Tente reenviar a foto ou tirar uma nova captura.") from erro


def imagem_precisa_normalizacao(caminho: Path, padronizar_figura: bool = False) -> bool:
    if not caminho.exists():
        return True

    try:
        with Image.open(caminho) as imagem:
            if padronizar_figura and imagem.size != FIGURA_CANVAS_PX:
                return True
            if imagem.mode not in ("RGB", "L"):
                return True
            return False
    except (UnidentifiedImageError, OSError):
        return True


def normalizar_imagem_salva(caminho: Path, padronizar_figura: bool = False) -> Path:
    caminho_final = caminho.with_suffix(".jpg")
    if imagem_precisa_normalizacao(caminho_final, padronizar_figura):
        return normalizar_imagem_para_docx(caminho.read_bytes(), caminho, padronizar_figura)
    return caminho_final


def salvar_upload(
    uploaded_file,
    pasta_temp: Path,
    prefixo: str,
    extensoes_permitidas: set[str],
    extensao_padrao: str,
    padronizar_figura: bool = False,
) -> Path | None:
    if uploaded_file is None:
        return None

    nome_original = getattr(uploaded_file, "name", "")
    extensao = Path(nome_original).suffix.lower().lstrip(".")
    if extensao not in extensoes_permitidas:
        extensao = extensao_padrao

    conteudo = uploaded_file.getvalue()
    caminho = pasta_temp / f"{prefixo}_{uuid.uuid4().hex[:8]}.{extensao}"
    if extensoes_permitidas == EXTENSOES_IMAGEM:
        return normalizar_imagem_para_docx(conteudo, caminho, padronizar_figura)

    caminho.write_bytes(conteudo)
    return caminho


def novo_manifesto_rascunho() -> dict:
    return {
        "tipo_atendimento": "",
        "tipos_atendimento": [],
        "tecnico_agres_responsavel": "",
        "data_atendimento_inicio": "",
        "data_atendimento_final": "",
        "localizacao_maps": "",
        "audios": [],
        "cabecalho": {"info_equip": None, "maquina": None, "implemento": None},
        "evidencias": {categoria: [] for categoria in CATEGORIAS_EVIDENCIAS},
        "fotos_atendimento": [],
        "observacoes": "",
        "legendas_evidencias": {categoria: "" for categoria in CATEGORIAS_EVIDENCIAS},
        "responsaveis": {chave: "" for chave in ASSINATURAS_RESPONSAVEIS},
        "documentos": {chave: "" for chave in ASSINATURAS_RESPONSAVEIS},
        "assinaturas": {chave: None for chave in ASSINATURAS_RESPONSAVEIS},
        "assinaturas_habilitadas": True,
        "quantidade_assinaturas": MIN_ASSINATURAS,
        "assinaturas_lista": assinaturas_padrao_lista(MIN_ASSINATURAS),
    }


def obter_id_rascunho() -> str:
    draft_url = st.query_params.get("draft", "")
    if isinstance(draft_url, list):
        draft_url = draft_url[0] if draft_url else ""

    if "draft_id" in st.session_state:
        draft_id = st.session_state.draft_id
    elif re.fullmatch(r"[0-9a-f]{12}", str(draft_url)):
        draft_id = str(draft_url)
    else:
        draft_id = uuid.uuid4().hex[:12]

    st.session_state.draft_id = draft_id
    st.query_params["draft"] = draft_id
    return draft_id


def pasta_rascunho_atual() -> Path:
    draft_dir = DRAFTS_DIR / obter_id_rascunho()
    draft_dir.mkdir(parents=True, exist_ok=True)
    if not st.session_state.get("rascunhos_antigos_verificados"):
        limpar_rascunhos_antigos(draft_dir)
        st.session_state.rascunhos_antigos_verificados = True
    return draft_dir


def caminho_manifesto(draft_dir: Path) -> Path:
    return draft_dir / "manifest.json"


def caminho_backup_manifesto(draft_dir: Path) -> Path:
    return draft_dir / "manifest.backup.json"


def limpar_rascunhos_antigos(draft_atual: Path | None = None) -> None:
    limite = time.time() - (DRAFT_RETENTION_DAYS * 24 * 60 * 60)
    for pasta in DRAFTS_DIR.iterdir():
        if not pasta.is_dir() or (draft_atual and pasta.resolve() == draft_atual.resolve()):
            continue
        try:
            if pasta.stat().st_mtime < limite:
                shutil.rmtree(pasta)
        except OSError:
            continue


def normalizar_manifesto_carregado(dados: dict) -> dict:
    if not isinstance(dados, dict):
        return novo_manifesto_rascunho()
    manifesto = novo_manifesto_rascunho()
    manifesto.update(dados)
    padrao = novo_manifesto_rascunho()
    manifesto["tecnico_agres_responsavel"] = normalizar_tecnico_agres_responsavel(
        manifesto.get("tecnico_agres_responsavel", "")
    )
    manifesto["audios"] = lista_ou_vazia(manifesto.get("audios"))
    manifesto["cabecalho"] = {**padrao["cabecalho"], **dicionario_ou_vazio(manifesto.get("cabecalho"))}
    manifesto["evidencias"] = {**padrao["evidencias"], **dicionario_ou_vazio(manifesto.get("evidencias"))}
    manifesto["evidencias"] = {
        categoria: lista_ou_vazia(manifesto["evidencias"].get(categoria))
        for categoria in CATEGORIAS_EVIDENCIAS
    }
    manifesto["fotos_atendimento"] = lista_ou_vazia(manifesto.get("fotos_atendimento"))
    manifesto["legendas_evidencias"] = {
        **padrao["legendas_evidencias"],
        **dicionario_ou_vazio(manifesto.get("legendas_evidencias")),
    }
    manifesto["responsaveis"] = {**padrao["responsaveis"], **dicionario_ou_vazio(manifesto.get("responsaveis"))}
    manifesto["documentos"] = {**padrao["documentos"], **dicionario_ou_vazio(manifesto.get("documentos"))}
    manifesto["assinaturas"] = {**padrao["assinaturas"], **dicionario_ou_vazio(manifesto.get("assinaturas"))}
    manifesto["assinaturas_habilitadas"] = bool(manifesto.get("assinaturas_habilitadas", True))
    manifesto["assinaturas_lista"] = normalizar_assinaturas_lista(manifesto)
    manifesto["quantidade_assinaturas"] = len(manifesto["assinaturas_lista"])
    return manifesto


def carregar_manifesto(draft_dir: Path) -> dict:
    for caminho in (caminho_manifesto(draft_dir), caminho_backup_manifesto(draft_dir)):
        if not caminho.exists():
            continue
        try:
            dados = json.loads(caminho.read_text(encoding="utf-8"))
            if isinstance(dados, dict):
                return normalizar_manifesto_carregado(dados)
        except (OSError, json.JSONDecodeError):
            continue
    return novo_manifesto_rascunho()


def salvar_manifesto(draft_dir: Path, manifesto: dict) -> None:
    draft_dir.mkdir(parents=True, exist_ok=True)
    caminho = caminho_manifesto(draft_dir)
    temporario = draft_dir / "manifest.tmp.json"
    backup = caminho_backup_manifesto(draft_dir)
    conteudo = json.dumps(manifesto, ensure_ascii=False, indent=2)
    temporario.write_text(conteudo, encoding="utf-8")
    json.loads(temporario.read_text(encoding="utf-8"))
    if caminho.exists():
        try:
            json.loads(caminho.read_text(encoding="utf-8"))
            shutil.copy2(caminho, backup)
        except (OSError, json.JSONDecodeError):
            pass
    temporario.replace(caminho)


def caminhos_referenciados_manifesto(manifesto: dict) -> set[str]:
    caminhos: set[str] = set()

    def visitar(valor) -> None:
        if isinstance(valor, dict):
            caminho = valor.get("path")
            if isinstance(caminho, str):
                caminhos.add(caminho.replace("\\", "/"))
            for item in valor.values():
                visitar(item)
        elif isinstance(valor, list):
            for item in valor:
                visitar(item)

    visitar(manifesto)
    return caminhos


def remover_arquivos_orfaos_rascunho(draft_dir: Path, manifesto: dict) -> None:
    referenciados = caminhos_referenciados_manifesto(manifesto)
    preservados = {"manifest.json", "manifest.backup.json", "manifest.tmp.json"}
    for caminho in draft_dir.rglob("*"):
        if not caminho.is_file():
            continue
        relativo = caminho.relative_to(draft_dir).as_posix()
        if relativo not in preservados and relativo not in referenciados:
            try:
                caminho.unlink()
            except OSError:
                continue
    for pasta in sorted((item for item in draft_dir.rglob("*") if item.is_dir()), reverse=True):
        try:
            pasta.rmdir()
        except OSError:
            continue


def resolver_arquivo_rascunho(draft_dir: Path, item) -> Path | None:
    if not item:
        return None
    rel_path = item.get("path") if isinstance(item, dict) else str(item)
    caminho = (draft_dir / rel_path).resolve()
    try:
        caminho.relative_to(draft_dir.resolve())
    except ValueError:
        return None
    return caminho if caminho.exists() else None


def salvar_arquivo_rascunho(
    uploaded_file,
    draft_dir: Path,
    subpasta: str,
    prefixo: str,
    extensoes_permitidas: set[str],
    extensao_padrao: str,
    padronizar_figura: bool = False,
) -> dict | None:
    if uploaded_file is None:
        return None

    conteudo = uploaded_file.getvalue()
    if not conteudo:
        return None
    limite = MAX_IMAGE_BYTES if extensoes_permitidas == EXTENSOES_IMAGEM else MAX_AUDIO_BYTES
    if len(conteudo) > limite:
        raise ValueError(f"O arquivo excede o limite de {limite // (1024 * 1024)} MB.")

    nome_original = getattr(uploaded_file, "name", f"{prefixo}.{extensao_padrao}")
    extensao = Path(nome_original).suffix.lower().lstrip(".")
    if extensao not in extensoes_permitidas:
        extensao = extensao_padrao

    pasta = draft_dir / subpasta
    pasta.mkdir(parents=True, exist_ok=True)
    digest = sha1(conteudo).hexdigest()[:12]
    caminho = pasta / f"{prefixo}_{digest}.{extensao}"

    if extensoes_permitidas == EXTENSOES_IMAGEM:
        caminho_final = caminho.with_suffix(".jpg")
        if imagem_precisa_normalizacao(caminho_final, padronizar_figura):
            caminho = normalizar_imagem_para_docx(conteudo, caminho, padronizar_figura)
        else:
            caminho = caminho_final
    elif not caminho.exists():
        caminho.write_bytes(conteudo)

    return {
        "name": nome_original,
        "path": str(caminho.relative_to(draft_dir)),
        "size": len(conteudo),
    }


def decodificar_data_url(
    data_url: str,
    mime_type_padrao: str = "application/octet-stream",
    limite_bytes: int | None = None,
) -> tuple[bytes, str]:
    bruto = str(data_url or "").strip()
    if not bruto:
        raise ValueError("Arquivo offline inválido: conteúdo em base64 vazio.")

    mime_type = mime_type_padrao.lower()
    match = re.match(r"^data:([^;,]+)?(?:;[^,]*)?;base64,(.*)$", bruto, flags=re.DOTALL)
    if match:
        mime_type = (match.group(1) or mime_type).lower()
        conteudo_base64 = match.group(2)
    else:
        conteudo_base64 = bruto

    conteudo_base64 = re.sub(r"\s+", "", conteudo_base64)
    if not conteudo_base64:
        raise ValueError("Arquivo offline inválido: conteúdo em base64 não encontrado.")
    tamanho_estimado = (len(conteudo_base64) * 3) // 4
    if limite_bytes and tamanho_estimado > limite_bytes:
        raise ValueError(f"O arquivo excede o limite de {limite_bytes // (1024 * 1024)} MB.")
    conteudo_base64 += "=" * ((4 - len(conteudo_base64) % 4) % 4)

    try:
        conteudo = base64.b64decode(conteudo_base64, validate=True)
    except Exception as erro:
        raise ValueError("Arquivo offline inválido: base64 corrompido ou incompleto.") from erro
    if limite_bytes and len(conteudo) > limite_bytes:
        raise ValueError(f"O arquivo excede o limite de {limite_bytes // (1024 * 1024)} MB.")
    return conteudo, mime_type


def extensao_item_offline(item: dict, mime_type: str, extensoes_permitidas: set[str], extensao_padrao: str) -> str:
    nome_original = str(item.get("name") or "")
    extensao = Path(nome_original).suffix.lower().lstrip(".")
    if not extensao:
        extensao = EXTENSAO_POR_MIME.get(mime_type, extensao_padrao)
    if extensao == "jpeg":
        extensao = "jpg"
    return extensao if extensao in extensoes_permitidas else extensao_padrao


def salvar_item_offline_rascunho(
    item: dict | None,
    draft_dir: Path,
    subpasta: str,
    prefixo: str,
    extensoes_permitidas: set[str],
    extensao_padrao: str,
    padronizar_figura: bool = False,
) -> dict | None:
    if not isinstance(item, dict):
        return None

    conteudo_codificado = item.get("dataUrl") or item.get("data_url") or item.get("base64") or item.get("content")
    if not conteudo_codificado:
        return None

    mime_type_informado = str(item.get("type") or item.get("mime_type") or item.get("mimeType") or "application/octet-stream")
    limite = MAX_IMAGE_BYTES if extensoes_permitidas == EXTENSOES_IMAGEM else MAX_AUDIO_BYTES
    conteudo, mime_type = decodificar_data_url(conteudo_codificado, mime_type_informado, limite)
    if not conteudo:
        return None

    extensao = extensao_item_offline(item, mime_type, extensoes_permitidas, extensao_padrao)
    nome_original = item.get("name") or f"{prefixo}.{extensao}"
    pasta = draft_dir / subpasta
    pasta.mkdir(parents=True, exist_ok=True)
    digest = sha1(conteudo).hexdigest()[:12]
    caminho = pasta / f"{prefixo}_{digest}.{extensao}"

    if extensoes_permitidas == EXTENSOES_IMAGEM:
        caminho = normalizar_imagem_para_docx(conteudo, caminho, padronizar_figura)
    elif not caminho.exists():
        caminho.write_bytes(conteudo)

    return {
        "name": nome_original,
        "path": str(caminho.relative_to(draft_dir)),
        "size": len(conteudo),
    }


def salvar_item_offline_seguro(
    item: dict | None,
    draft_dir: Path,
    subpasta: str,
    prefixo: str,
    extensoes_permitidas: set[str],
    extensao_padrao: str,
    erros: list[str],
    rotulo: str,
    padronizar_figura: bool = False,
) -> dict | None:
    try:
        return salvar_item_offline_rascunho(
            item,
            draft_dir,
            subpasta,
            prefixo,
            extensoes_permitidas,
            extensao_padrao,
            padronizar_figura,
        )
    except Exception as erro:
        erros.append(f"{rotulo}: {erro}")
        return None


def importar_pacote_offline_json(texto_pacote: str, draft_dir: Path, manifesto: dict) -> dict:
    if len(texto_pacote) > MAX_PACKAGE_BYTES:
        raise ValueError(f"O pacote excede o limite de {MAX_PACKAGE_BYTES // (1024 * 1024)} MB.")
    try:
        pacote = json.loads(texto_pacote)
    except json.JSONDecodeError as erro:
        raise ValueError("Não foi possível ler o pacote offline. Exporte novamente pelo modo offline.") from erro

    if not isinstance(pacote, dict) or pacote.get("version") != 1:
        raise ValueError("Pacote offline incompatível com esta versão do aplicativo.")

    campos_objeto = ("cabecalho", "evidencias", "legendas_evidencias", "responsaveis", "documentos", "assinaturas")
    campos_lista = ("audios", "fotos_atendimento", "assinaturas_lista")
    for campo in campos_objeto:
        if campo in pacote and not isinstance(pacote[campo], dict):
            raise ValueError(f"Pacote offline corrompido: o campo '{campo}' possui formato inválido.")
    for campo in campos_lista:
        if campo in pacote and not isinstance(pacote[campo], list):
            raise ValueError(f"Pacote offline corrompido: o campo '{campo}' possui formato inválido.")
    for categoria, itens in dicionario_ou_vazio(pacote.get("evidencias")).items():
        if categoria in CATEGORIAS_EVIDENCIAS and not isinstance(itens, list):
            raise ValueError(f"Pacote offline corrompido: as fotos de '{categoria}' possuem formato inválido.")

    manifesto = novo_manifesto_rascunho()
    tipos_pacote = normalizar_tipos_atendimento(
        pacote.get("tipos_atendimento") or pacote.get("tipo_atendimento"),
        manifesto.get("tipos_atendimento", []),
    )
    manifesto["tipos_atendimento"] = tipos_pacote
    manifesto["tipo_atendimento"] = tipos_pacote[0] if tipos_pacote else ""
    manifesto["tecnico_agres_responsavel"] = normalizar_tecnico_agres_responsavel(
        pacote.get("tecnico_agres_responsavel", manifesto.get("tecnico_agres_responsavel", ""))
    )
    manifesto["data_atendimento_inicio"] = data_final_pacote(pacote.get("data_atendimento_inicio", ""))
    manifesto["data_atendimento_final"] = data_pacote_ou_texto(
        pacote.get("data_atendimento_final", ""),
        pacote.get("periodo_atendimento", ""),
        pacote.get("data_visita", ""),
        pacote.get("observacoes", ""),
    )
    manifesto["localizacao_maps"] = limpar_texto(pacote.get("localizacao_maps", ""))
    manifesto["observacoes"] = limpar_texto(pacote.get("observacoes", ""))

    for categoria, texto in dicionario_ou_vazio(pacote.get("legendas_evidencias")).items():
        if categoria in CATEGORIAS_EVIDENCIAS:
            manifesto["legendas_evidencias"][categoria] = texto or ""

    for chave, texto in dicionario_ou_vazio(pacote.get("responsaveis")).items():
        if chave in ASSINATURAS_RESPONSAVEIS:
            manifesto["responsaveis"][chave] = texto or ""

    for chave, texto in dicionario_ou_vazio(pacote.get("documentos")).items():
        if chave in ASSINATURAS_RESPONSAVEIS:
            manifesto["documentos"][chave] = texto or ""

    audios = []
    audios_recebidos = lista_ou_vazia(pacote.get("audios"))
    erros_audio = []
    erros_imagem = []
    imagens_recebidas = 0
    imagens_salvas = 0
    for indice, item in enumerate(audios_recebidos):
        try:
            audio = salvar_item_offline_rascunho(
                item,
                draft_dir,
                "audios",
                f"offline_audio_{indice}",
                EXTENSOES_AUDIO,
                "m4a",
            )
        except Exception as erro:
            erros_audio.append(f"Áudio {indice + 1}: {erro}")
            continue
        if audio:
            audios.append(audio)
    if audios or "audios" in pacote:
        manifesto["audios"] = audios

    manifesto["importacao_offline"] = {
        "audios_recebidos": len(audios_recebidos),
        "audios_salvos": len(audios),
        "erros_audio": erros_audio[:3],
        "exported_at": limpar_texto(pacote.get("exported_at", "")),
        "app_version": limpar_texto(pacote.get("app_version", "")),
    }

    for chave, item in dicionario_ou_vazio(pacote.get("cabecalho")).items():
        if chave in manifesto["cabecalho"]:
            manifesto["cabecalho"][chave] = None
            if item:
                imagens_recebidas += 1
            arquivo = salvar_item_offline_seguro(
                item,
                draft_dir,
                "cabecalho",
                chave,
                EXTENSOES_IMAGEM,
                "jpg",
                erros_imagem,
                f"Cabeçalho {chave}",
            )
            if arquivo:
                imagens_salvas += 1
                manifesto["cabecalho"][chave] = arquivo

    for categoria, itens in dicionario_ou_vazio(pacote.get("evidencias")).items():
        if categoria not in CATEGORIAS_EVIDENCIAS:
            continue
        evidencias = []
        for indice, item in enumerate(itens or []):
            if item:
                imagens_recebidas += 1
            arquivo = salvar_item_offline_seguro(
                item,
                draft_dir,
                categoria,
                f"offline_{categoria}_{indice}",
                EXTENSOES_IMAGEM,
                "jpg",
                erros_imagem,
                f"{CATEGORIAS_EVIDENCIAS[categoria]['nome']} {indice + 1}",
                padronizar_figura=True,
            )
            if arquivo:
                imagens_salvas += 1
                evidencias.append(arquivo)
        manifesto["evidencias"][categoria] = evidencias

    fotos_atendimento = []
    for indice, item in enumerate(lista_ou_vazia(pacote.get("fotos_atendimento"))):
        if item:
            imagens_recebidas += 1
        foto = salvar_item_offline_seguro(
            item,
            draft_dir,
            "fotos_atendimento",
            f"offline_foto_atendimento_{indice}",
            EXTENSOES_IMAGEM,
            "jpg",
            erros_imagem,
            f"Foto do atendimento {indice + 1}",
        )
        if foto:
            imagens_salvas += 1
            fotos_atendimento.append(foto)
    if fotos_atendimento or "fotos_atendimento" in pacote:
        manifesto["fotos_atendimento"] = fotos_atendimento

    if "assinaturas_habilitadas" in pacote:
        manifesto["assinaturas_habilitadas"] = bool(pacote.get("assinaturas_habilitadas"))

    assinaturas_lista = pacote.get("assinaturas_lista")
    if isinstance(assinaturas_lista, list):
        quantidade_bruta = pacote.get("quantidade_assinaturas")
        quantidade = quantidade_assinaturas_normalizada(len(assinaturas_lista) if quantidade_bruta is None else quantidade_bruta)
        registros = []
        for indice in range(1, quantidade + 1):
            padrao = assinatura_padrao(indice)
            item = assinaturas_lista[indice - 1] if indice - 1 < len(assinaturas_lista) and isinstance(assinaturas_lista[indice - 1], dict) else {}
            imagem = None
            imagem_item = item.get("imagem") or item.get("assinatura")
            if imagem_item:
                imagens_recebidas += 1
                imagem = salvar_item_offline_seguro(
                    imagem_item,
                    draft_dir,
                    "assinaturas",
                    padrao["id"],
                    EXTENSOES_IMAGEM,
                    "png",
                    erros_imagem,
                    f"Assinatura {indice}",
                )
                if imagem:
                    imagens_salvas += 1
            registros.append(
                {
                    "id": padrao["id"],
                    "nome": texto_assinatura(
                        item,
                        "nome",
                        "responsavel",
                        "responsável",
                        "nome_responsavel",
                        "nomeResponsavel",
                    ),
                    "representa": limpar_papel_assinatura(
                        texto_assinatura(
                            item,
                            "representa",
                            "funcao",
                            "função",
                            "cargo",
                            "papel",
                            "role",
                            padrao=padrao["representa"],
                        )
                    ),
                    "documento": texto_assinatura(item, "documento", "cpf_rg", "cpfRg", "cpf", "rg", "doc"),
                    "imagem": imagem,
                }
            )
        manifesto["quantidade_assinaturas"] = quantidade
        manifesto["assinaturas_lista"] = registros
        if registros and any(assinatura_tem_conteudo(assinatura) for assinatura in registros):
            manifesto["assinaturas_habilitadas"] = True
        sincronizar_assinaturas_legadas(manifesto)
    else:
        for chave, item in dicionario_ou_vazio(pacote.get("assinaturas")).items():
            if chave in ASSINATURAS_RESPONSAVEIS:
                manifesto["assinaturas"][chave] = None
                if item:
                    imagens_recebidas += 1
                assinatura = salvar_item_offline_seguro(
                    item,
                    draft_dir,
                    "assinaturas",
                    chave,
                    EXTENSOES_IMAGEM,
                    "png",
                    erros_imagem,
                    f"Assinatura {chave}",
                )
                if assinatura:
                    imagens_salvas += 1
                    manifesto["assinaturas"][chave] = assinatura
        manifesto["assinaturas_lista"] = normalizar_assinaturas_lista(manifesto)
        manifesto["quantidade_assinaturas"] = len(manifesto["assinaturas_lista"])
        sincronizar_assinaturas_legadas(manifesto)

    manifesto["importacao_offline"].update(
        {
            "imagens_recebidas": imagens_recebidas,
            "imagens_salvas": imagens_salvas,
            "erros_imagem": erros_imagem[:8],
        }
    )

    salvar_manifesto(draft_dir, manifesto)
    remover_arquivos_orfaos_rascunho(draft_dir, manifesto)
    shutil.copy2(caminho_manifesto(draft_dir), caminho_backup_manifesto(draft_dir))
    return manifesto


def extrair_texto_pacote_offline(uploaded_file) -> tuple[str, str]:
    conteudo = uploaded_file.getvalue()
    nome_upload = getattr(uploaded_file, "name", "pacote_relatorio.json") or "pacote_relatorio.json"
    if len(conteudo) > MAX_PACKAGE_BYTES:
        raise ValueError(f"O pacote excede o limite de {MAX_PACKAGE_BYTES // (1024 * 1024)} MB.")

    extensao = Path(nome_upload).suffix.lower()
    tipo_upload = (getattr(uploaded_file, "type", "") or "").lower()
    if extensao == ".zip" or "zip" in tipo_upload:
        try:
            with zipfile.ZipFile(BytesIO(conteudo)) as pacote_zip:
                candidatos = [
                    info
                    for info in pacote_zip.infolist()
                    if not info.is_dir() and Path(info.filename).suffix.lower() == ".json"
                ]
                if not candidatos:
                    raise ValueError("O ZIP não contém um arquivo JSON da coleta offline.")
                candidatos.sort(
                    key=lambda info: (
                        0 if re.match(r"^\d{8}_RELATORIO_ATIVIDADES_", Path(info.filename).name, re.I) else 1,
                        Path(info.filename).name.lower(),
                    )
                )
                entrada_json = candidatos[0]
                if entrada_json.file_size > MAX_PACKAGE_BYTES:
                    raise ValueError(f"O JSON dentro do ZIP excede o limite de {MAX_PACKAGE_BYTES // (1024 * 1024)} MB.")
                texto_pacote = pacote_zip.read(entrada_json).decode("utf-8-sig")
                return texto_pacote, Path(entrada_json.filename).name or nome_upload
        except zipfile.BadZipFile as erro:
            raise ValueError("Não foi possível abrir o ZIP. Exporte novamente pelo modo offline.") from erro
        except UnicodeDecodeError as erro:
            raise ValueError("Não foi possível ler o JSON dentro do ZIP. Exporte novamente pelo modo offline.") from erro

    try:
        return conteudo.decode("utf-8-sig"), nome_upload
    except UnicodeDecodeError as erro:
        raise ValueError("Não foi possível ler o pacote offline. Exporte novamente pelo modo offline.") from erro


def importar_pacote_offline(uploaded_file, draft_dir: Path, manifesto: dict) -> dict:
    texto_pacote, nome_referencia = extrair_texto_pacote_offline(uploaded_file)
    manifesto_importado = importar_pacote_offline_json(texto_pacote, draft_dir, manifesto)
    data_nome = data_final_pacote(nome_referencia) or data_final_pacote(getattr(uploaded_file, "name", ""))
    if data_nome and not manifesto_importado.get("data_atendimento_final"):
        manifesto_importado["data_atendimento_final"] = data_nome
        salvar_manifesto(draft_dir, manifesto_importado)
        shutil.copy2(caminho_manifesto(draft_dir), caminho_backup_manifesto(draft_dir))
    return manifesto_importado


def aplicar_manifesto_na_sessao(manifesto: dict) -> None:
    tipos = normalizar_tipos_atendimento(
        manifesto.get("tipos_atendimento") or manifesto.get("tipo_atendimento"),
        [],
    )
    st.session_state.tipos_atendimento = tipos
    st.session_state.tipo_atendimento = tipos[0] if tipos else ""
    st.session_state.localizacao_maps = manifesto.get("localizacao_maps", "")
    st.session_state.observacoes_texto = manifesto.get("observacoes", "")
    for categoria in CATEGORIAS_EVIDENCIAS:
        st.session_state[f"legenda_{categoria}"] = manifesto.get("legendas_evidencias", {}).get(categoria, "")
    for chave, configuracao in ASSINATURAS_RESPONSAVEIS.items():
        st.session_state[configuracao["campo"]] = manifesto.get("responsaveis", {}).get(chave, "")
        st.session_state[configuracao["campo_documento"]] = manifesto.get("documentos", {}).get(chave, "")
    st.session_state.assinaturas_habilitadas = manifesto.get("assinaturas_habilitadas", True)
    st.session_state.quantidade_assinaturas = quantidade_assinaturas_normalizada(manifesto.get("quantidade_assinaturas", MIN_ASSINATURAS))
    for assinatura in normalizar_assinaturas_lista(manifesto):
        assinatura_id = assinatura["id"]
        st.session_state[f"assinatura_nome_{assinatura_id}"] = assinatura.get("nome", "")
        st.session_state[f"assinatura_representa_{assinatura_id}"] = assinatura.get("representa", "")
        st.session_state[f"assinatura_documento_{assinatura_id}"] = assinatura.get("documento", "")


def uploaded_file_para_item_pacote(uploaded_file) -> dict | None:
    if uploaded_file is None:
        return None
    conteudo = uploaded_file.getvalue()
    if not conteudo:
        return None
    mime_type = getattr(uploaded_file, "type", None) or "application/octet-stream"
    return {
        "name": getattr(uploaded_file, "name", "arquivo"),
        "type": mime_type,
        "size": len(conteudo),
        "dataUrl": f"data:{mime_type};base64,{base64.b64encode(conteudo).decode('ascii')}",
    }


def canvas_para_item_pacote(canvas_result, nome: str) -> dict | None:
    dados = getattr(canvas_result, "image_data", None)
    if dados is None:
        return None

    try:
        imagem_rgba = Image.fromarray(dados.astype("uint8"), "RGBA")
    except Exception:
        return None

    pixels_assinatura = 0
    for vermelho, verde, azul, alfa in imagem_rgba.getdata():
        if alfa > 0 and min(vermelho, verde, azul) < 245:
            pixels_assinatura += 1
            if pixels_assinatura >= 50:
                break

    if pixels_assinatura < 50:
        return None

    fundo = Image.new("RGB", imagem_rgba.size, "white")
    fundo.paste(imagem_rgba, mask=imagem_rgba.getchannel("A"))
    buffer = BytesIO()
    fundo.save(buffer, format="PNG")
    conteudo = buffer.getvalue()
    return {
        "name": nome,
        "type": "image/png",
        "size": len(conteudo),
        "dataUrl": f"data:image/png;base64,{base64.b64encode(conteudo).decode('ascii')}",
    }


def renderizar_coleta_streamlit() -> None:
    st.markdown(
        "<p class='section-title'>Coleta de campo</p>",
        unsafe_allow_html=True,
    )

    with st.container(border=True):
        st.markdown("### 1. Relato técnico")
        coleta_tipos_atendimento = st.multiselect(
            "Tipo de Atendimento",
            options=list(TIPOS_ATENDIMENTO),
            default=normalizar_tipos_atendimento(st.session_state.get("tipos_atendimento"), []),
            format_func=lambda chave: TIPOS_ATENDIMENTO[chave],
            key="coleta_tipos_atendimento",
            max_selections=2,
        )
        coleta_tipos_atendimento = normalizar_tipos_atendimento(coleta_tipos_atendimento, [])
        coleta_tipo_atendimento = coleta_tipos_atendimento[0] if coleta_tipos_atendimento else ""
        coleta_localizacao_maps = st.text_input(
            "Link de Localização do Atendimento",
            key="coleta_localizacao_maps",
        )
        coleta_observacoes = st.text_area(
            "Complemento Técnico",
            height=150,
            key="coleta_observacoes",
        )
        coleta_audio = st.audio_input("Gravar áudio", key="coleta_audio")
        coleta_audios_upload = st.file_uploader(
            "Ou enviar áudios",
            type=list(EXTENSOES_AUDIO),
            accept_multiple_files=True,
            key="coleta_audios_upload",
        )

    with st.container(border=True):
        st.markdown("### 2. Cabeçalho do Relatório")
        col1, col2, col3 = st.columns(3)
        coleta_info = col1.file_uploader("Informações do Equipamento", type=list(EXTENSOES_IMAGEM), key="coleta_info")
        coleta_maquina = col2.file_uploader("Máquina", type=list(EXTENSOES_IMAGEM), key="coleta_maquina")
        coleta_implemento = col3.file_uploader("Implemento", type=list(EXTENSOES_IMAGEM), key="coleta_implemento")

    with st.container(border=True):
        st.markdown("### 3. Evidências Fotográficas")
        col_e1, col_e2 = st.columns(2)
        coleta_equipamento = col_e1.file_uploader("Equipamento Agres", type=list(EXTENSOES_IMAGEM), accept_multiple_files=True, key="coleta_equipamento")
        coleta_instalacao = col_e1.file_uploader("Instalação", type=list(EXTENSOES_IMAGEM), accept_multiple_files=True, key="coleta_instalacao")
        coleta_configuracao = col_e2.file_uploader("Configurações", type=list(EXTENSOES_IMAGEM), accept_multiple_files=True, key="coleta_configuracao")
        coleta_outros = col_e2.file_uploader("Outros registros", type=list(EXTENSOES_IMAGEM), accept_multiple_files=True, key="coleta_outros")
        coleta_fotos_atendimento = st.file_uploader("Fotos do Atendimento (Somente ZIP)", type=list(EXTENSOES_IMAGEM), accept_multiple_files=True, key="coleta_fotos_atendimento")

        coleta_legendas = {
            "fotos_equipamento": st.text_area("Legendas - Equipamento", height=80, key="coleta_leg_equipamento"),
            "fotos_instalacao": st.text_area("Legendas - Instalação", height=80, key="coleta_leg_instalacao"),
            "fotos_configuracao": st.text_area("Legendas - Configurações", height=80, key="coleta_leg_configuracao"),
            "fotos_outros": st.text_area("Legendas - Outros", height=80, key="coleta_leg_outros"),
        }

    with st.container(border=True):
        st.markdown("### 4. Assinaturas")
        col_nome1, col_nome2 = st.columns(2)
        with col_nome1:
            responsavel_revenda = st.text_input("Responsável da revenda/fábrica", key="coleta_resp_revenda")
        with col_nome2:
            responsavel_fazenda = st.text_input("Responsável da fazenda", key="coleta_resp_fazenda")
        responsaveis = {"revenda_fabrica": responsavel_revenda, "fazenda": responsavel_fazenda}
        documentos = {"revenda_fabrica": "", "fazenda": ""}
        assinaturas = {"revenda_fabrica": None, "fazenda": None}
        if st_canvas is not None:
            usar_coleta_ampliada = st.toggle(
                "Usar área de assinatura ampliada",
                value=False,
                key="coleta_usar_assinatura_ampliada",
            )
            largura_coleta = 760 if usar_coleta_ampliada else 330
            altura_coleta = 360 if usar_coleta_ampliada else 180
            sufixo_coleta = "ampliada" if usar_coleta_ampliada else "normal"
            colunas_coleta = (
                [st.container(border=True), st.container(border=True)]
                if usar_coleta_ampliada
                else list(st.columns(2))
            )
            for chave, titulo, coluna in (
                ("revenda_fabrica", "Assinatura revenda/fábrica", colunas_coleta[0]),
                ("fazenda", "Assinatura fazenda", colunas_coleta[1]),
            ):
                with coluna:
                    st.markdown(f"**{titulo}**")
                    assinaturas[chave] = st_canvas(
                        fill_color="rgba(255, 255, 255, 0)",
                        stroke_width=3,
                        stroke_color="#25272b",
                        background_color="#ffffff",
                        height=altura_coleta,
                        width=largura_coleta,
                        drawing_mode="freedraw",
                        display_toolbar=True,
                        key=f"coleta_ass_{chave}_{sufixo_coleta}",
                    )
        else:
            st.info("Assinatura na tela indisponível neste ambiente. Envie imagens das assinaturas.")
            assinaturas_upload = {
                "revenda_fabrica": st.file_uploader("Imagem assinatura revenda/fábrica", type=list(EXTENSOES_IMAGEM), key="coleta_ass_upload_revenda"),
                "fazenda": st.file_uploader("Imagem assinatura fazenda", type=list(EXTENSOES_IMAGEM), key="coleta_ass_upload_fazenda"),
            }

    audios = []
    if coleta_audio:
        item = uploaded_file_para_item_pacote(coleta_audio)
        if item:
            audios.append(item)
    audios.extend(
        item
        for arquivo in (coleta_audios_upload or [])
        if (item := uploaded_file_para_item_pacote(arquivo))
    )

    pacote = {
        "version": 1,
        "tipo_atendimento": coleta_tipo_atendimento,
        "tipos_atendimento": coleta_tipos_atendimento,
        "localizacao_maps": coleta_localizacao_maps or "",
        "observacoes": coleta_observacoes or "",
        "audios": audios,
        "cabecalho": {
            "info_equip": uploaded_file_para_item_pacote(coleta_info),
            "maquina": uploaded_file_para_item_pacote(coleta_maquina),
            "implemento": uploaded_file_para_item_pacote(coleta_implemento),
        },
        "evidencias": {
            "fotos_equipamento": [item for arquivo in (coleta_equipamento or []) if (item := uploaded_file_para_item_pacote(arquivo))],
            "fotos_instalacao": [item for arquivo in (coleta_instalacao or []) if (item := uploaded_file_para_item_pacote(arquivo))],
            "fotos_configuracao": [item for arquivo in (coleta_configuracao or []) if (item := uploaded_file_para_item_pacote(arquivo))],
            "fotos_outros": [item for arquivo in (coleta_outros or []) if (item := uploaded_file_para_item_pacote(arquivo))],
        },
        "fotos_atendimento": [item for arquivo in (coleta_fotos_atendimento or []) if (item := uploaded_file_para_item_pacote(arquivo))],
        "legendas_evidencias": coleta_legendas,
        "responsaveis": responsaveis,
        "documentos": documentos,
        "assinaturas": {},
    }

    if st_canvas is not None:
        pacote["assinaturas"] = {
            "revenda_fabrica": canvas_para_item_pacote(assinaturas["revenda_fabrica"], "assinatura_revenda_fabrica.png"),
            "fazenda": canvas_para_item_pacote(assinaturas["fazenda"], "assinatura_fazenda.png"),
        }
    else:
        pacote["assinaturas"] = {
            "revenda_fabrica": uploaded_file_para_item_pacote(assinaturas_upload.get("revenda_fabrica")),
            "fazenda": uploaded_file_para_item_pacote(assinaturas_upload.get("fazenda")),
        }

    pacote_texto = json.dumps(pacote, ensure_ascii=False)
    st.download_button(
        "Baixar pacote da coleta",
        data=pacote_texto.encode("utf-8"),
        file_name=f"agres_coleta_{data_atual_brasil().strftime('%Y%m%d')}.json",
        mime="application/json",
        type="primary",
        use_container_width=True,
    )
    with st.expander("Copiar pacote em texto"):
        st.text_area("Pacote para colar no app principal", value=pacote_texto, height=180)

    if st.button("Voltar ao app principal", use_container_width=True):
        st.query_params.clear()
        st.rerun()


def atualizar_lista_rascunho(manifesto: dict, chave: str, arquivos, draft_dir: Path, subpasta: str, extensoes: set[str], extensao_padrao: str) -> None:
    if not arquivos:
        return

    itens = []
    for indice, arquivo in enumerate(arquivos):
        item = salvar_arquivo_rascunho(arquivo, draft_dir, subpasta, f"{chave}_{indice}", extensoes, extensao_padrao)
        if item:
            itens.append(item)

    if itens:
        manifesto[chave] = itens


def salvar_assinatura_canvas(canvas_result, draft_dir: Path, chave: str) -> dict | None:
    dados = getattr(canvas_result, "image_data", None)
    if dados is None:
        return None

    try:
        imagem_rgba = Image.fromarray(dados.astype("uint8"), "RGBA")
    except Exception:
        return None

    pixels_assinatura = 0
    for vermelho, verde, azul, alfa in imagem_rgba.getdata():
        if alfa > 0 and min(vermelho, verde, azul) < 245:
            pixels_assinatura += 1
            if pixels_assinatura >= 50:
                break

    if pixels_assinatura < 50:
        return None

    pasta = draft_dir / "assinaturas"
    pasta.mkdir(parents=True, exist_ok=True)
    caminho = pasta / f"{chave}.png"
    fundo = Image.new("RGB", imagem_rgba.size, "white")
    fundo.paste(imagem_rgba, mask=imagem_rgba.getchannel("A"))
    fundo.save(caminho, format="PNG")

    return {
        "name": caminho.name,
        "path": str(caminho.relative_to(draft_dir)),
        "size": caminho.stat().st_size,
    }


def atualizar_rascunho_atual(
    draft_dir: Path,
    manifesto: dict,
    tipo_atendimento,
    localizacao_maps: str,
    audios,
    cabecalho: dict,
    evidencias_upload: dict,
    fotos_atendimento_upload,
    observacoes: str,
    legendas: dict,
    responsaveis: dict | None = None,
    documentos: dict | None = None,
    assinaturas_habilitadas: bool = True,
    quantidade_assinaturas: int = MIN_ASSINATURAS,
    assinaturas_form: dict | None = None,
    assinaturas_canvas: dict | None = None,
    assinaturas_upload: dict | None = None,
) -> dict:
    tipos = normalizar_tipos_atendimento(tipo_atendimento, [])
    manifesto["tipos_atendimento"] = tipos
    manifesto["tipo_atendimento"] = tipos[0] if tipos else ""
    manifesto["localizacao_maps"] = limpar_texto(localizacao_maps)

    if audios:
        manifesto["audios"] = [
            item
            for indice, audio in enumerate(audios)
            if (item := salvar_arquivo_rascunho(audio, draft_dir, "audios", f"audio_{indice}", EXTENSOES_AUDIO, "wav"))
        ]

    for chave, arquivo in cabecalho.items():
        item = salvar_arquivo_rascunho(arquivo, draft_dir, "cabecalho", chave, EXTENSOES_IMAGEM, "jpg")
        if item:
            manifesto["cabecalho"][chave] = item

    for categoria, arquivos in evidencias_upload.items():
        if arquivos:
            itens = []
            for indice, arquivo in enumerate(arquivos):
                item = salvar_arquivo_rascunho(
                    arquivo,
                    draft_dir,
                    categoria,
                    f"{categoria}_{indice}",
                    EXTENSOES_IMAGEM,
                    "jpg",
                    padronizar_figura=True,
                )
                if item:
                    itens.append(item)
            if itens:
                manifesto["evidencias"][categoria] = itens

    if fotos_atendimento_upload:
        itens = []
        for indice, arquivo in enumerate(fotos_atendimento_upload):
            item = salvar_arquivo_rascunho(
                arquivo,
                draft_dir,
                "fotos_atendimento",
                f"foto_atendimento_{indice}",
                EXTENSOES_IMAGEM,
                "jpg",
            )
            if item:
                itens.append(item)
        if itens:
            manifesto["fotos_atendimento"] = itens

    manifesto["observacoes"] = observacoes or ""

    for categoria, texto in (legendas or {}).items():
        if categoria in CATEGORIAS_EVIDENCIAS:
            manifesto["legendas_evidencias"][categoria] = texto or ""

    manifesto["assinaturas_habilitadas"] = bool(assinaturas_habilitadas)
    manifesto["quantidade_assinaturas"] = quantidade_assinaturas_normalizada(quantidade_assinaturas)
    assinaturas_atuais = {item["id"]: item for item in normalizar_assinaturas_lista(manifesto)}
    assinaturas_form = assinaturas_form or {}

    for indice in range(1, manifesto["quantidade_assinaturas"] + 1):
        assinatura_id = f"assinatura_{indice}"
        atual = assinaturas_atuais.get(assinatura_id, assinatura_padrao(indice))
        recebido = assinaturas_form.get(assinatura_id, {})
        assinaturas_atuais[assinatura_id] = {
            "id": assinatura_id,
            "nome": limpar_texto(recebido.get("nome", atual.get("nome", ""))),
            "representa": limpar_papel_assinatura(recebido.get("representa", atual.get("representa", ""))),
            "documento": limpar_texto(recebido.get("documento", atual.get("documento", ""))),
            "imagem": atual.get("imagem"),
        }

    for chave, resultado in (assinaturas_canvas or {}).items():
        if chave in assinaturas_atuais:
            item = salvar_assinatura_canvas(resultado, draft_dir, chave)
            if item:
                assinaturas_atuais[chave]["imagem"] = item
                st.session_state[f"editar_assinatura_{chave}"] = False

    for chave, arquivo in (assinaturas_upload or {}).items():
        if chave in assinaturas_atuais:
            item = salvar_arquivo_rascunho(arquivo, draft_dir, "assinaturas", chave, EXTENSOES_IMAGEM, "png")
            if item:
                assinaturas_atuais[chave]["imagem"] = item
                st.session_state[f"editar_assinatura_{chave}"] = False

    manifesto["assinaturas_lista"] = [
        assinaturas_atuais[f"assinatura_{indice}"]
        for indice in range(1, manifesto["quantidade_assinaturas"] + 1)
    ]

    sincronizar_assinaturas_legadas(manifesto)

    salvar_manifesto(draft_dir, manifesto)
    return manifesto


def caminhos_salvos_rascunho(draft_dir: Path, manifesto: dict) -> tuple[list[Path], dict, dict, list[Path], list[dict]]:
    audios = [caminho for item in manifesto.get("audios", []) if (caminho := resolver_arquivo_rascunho(draft_dir, item))]
    cabecalho = {
        chave: resolver_arquivo_rascunho(draft_dir, item)
        for chave, item in manifesto.get("cabecalho", {}).items()
    }
    evidencias = {
        categoria: [
            normalizar_imagem_salva(caminho, padronizar_figura=True)
            for item in manifesto.get("evidencias", {}).get(categoria, [])
            if (caminho := resolver_arquivo_rascunho(draft_dir, item))
        ]
        for categoria in CATEGORIAS_EVIDENCIAS
    }
    fotos_atendimento = [
        normalizar_imagem_salva(caminho)
        for item in manifesto.get("fotos_atendimento", [])
        if (caminho := resolver_arquivo_rascunho(draft_dir, item))
    ]
    assinaturas = []
    assinaturas_normalizadas = normalizar_assinaturas_lista(manifesto)
    assinaturas_com_conteudo = any(assinatura_tem_conteudo(item) for item in assinaturas_normalizadas)
    if manifesto.get("assinaturas_habilitadas", True) or assinaturas_com_conteudo:
        for assinatura in assinaturas_normalizadas:
            caminho = resolver_arquivo_rascunho(draft_dir, assinatura.get("imagem"))
            if assinatura.get("nome") or assinatura.get("representa") or assinatura.get("documento") or caminho:
                assinaturas.append({**assinatura, "caminho": caminho})
    return audios, cabecalho, evidencias, fotos_atendimento, assinaturas


def contar_evidencias(manifesto: dict) -> int:
    return sum(len(itens or []) for itens in manifesto.get("evidencias", {}).values())


def contar_linhas_legenda_manual(texto: str) -> int:
    total = 0
    for linha in linhas_metadados_preservando_vazios(texto):
        titulo, legenda, fonte = separar_metadados_figura(linha)
        if linha and linha != "||" and (titulo or legenda or fonte):
            total += 1
    return total


def itens_checklist_pacote(manifesto: dict) -> list[dict]:
    total_audios = len(manifesto.get("audios", []) or [])
    total_cabecalho = sum(1 for item in manifesto.get("cabecalho", {}).values() if item)
    total_evidencias = contar_evidencias(manifesto)
    total_fotos_zip = len(manifesto.get("fotos_atendimento", []) or [])
    assinaturas = normalizar_assinaturas_lista(manifesto)
    assinaturas_com_imagem = sum(1 for item in assinaturas if item.get("imagem"))
    assinaturas_com_nome = sum(1 for item in assinaturas if limpar_texto(item.get("nome", "")))
    legendas_manuais = sum(
        contar_linhas_legenda_manual(texto)
        for texto in (manifesto.get("legendas_evidencias", {}) or {}).values()
    )
    texto_tecnico = limpar_texto(manifesto.get("observacoes", ""))
    localizacao = limpar_texto(manifesto.get("localizacao_maps", ""))
    tecnico_agres = normalizar_tecnico_agres_responsavel(manifesto.get("tecnico_agres_responsavel", ""))
    tipos_atendimento = normalizar_tipos_atendimento(
        manifesto.get("tipos_atendimento") or manifesto.get("tipo_atendimento"),
        [],
    )

    itens = [
        {
            "ok": bool(tipos_atendimento),
            "titulo": "Tipo de Atendimento",
            "detalhe": tipos_atendimento_para_texto(tipos_atendimento),
        },
        {
            "ok": bool(tecnico_agres),
            "titulo": "Técnico Responsável Agres",
            "detalhe": tecnico_agres if tecnico_agres else "Não informado. Exporte novamente após selecionar o técnico na coleta offline.",
        },
        {
            "ok": bool(total_audios or texto_tecnico),
            "titulo": "Relato técnico",
            "detalhe": (
                f"{total_audios} áudio(s) importado(s) e complemento escrito preenchido."
                if total_audios and texto_tecnico
                else f"{total_audios} áudio(s) importado(s)."
                if total_audios
                else "Complemento escrito preenchido, sem áudio importado."
                if texto_tecnico
                else "Sem áudio e sem complemento técnico. A IA terá pouco contexto para redigir o relato."
            ),
        },
        {
            "ok": bool(localizacao),
            "titulo": "Localização",
            "detalhe": "Link/coordenada importado do pacote offline." if localizacao else "Sem link de localização da fazenda.",
        },
        {
            "ok": total_cabecalho >= 1,
            "titulo": "Cabeçalho",
            "detalhe": f"{total_cabecalho}/3 foto(s) de cabeçalho importada(s).",
        },
        {
            "ok": total_evidencias >= 1,
            "titulo": "Evidências para o Word",
            "detalhe": f"{total_evidencias} foto(s) serão inserida(s) no relatório.",
        },
        {
            "ok": bool(total_fotos_zip or total_evidencias or total_cabecalho),
            "titulo": "Fotos no ZIP",
            "detalhe": f"{total_fotos_zip} foto(s) extras, além de cabeçalho/evidências/assinaturas.",
        },
        {
            "ok": True,
            "titulo": "Legendas",
            "detalhe": (
                f"{legendas_manuais} legenda(s) manual(is) importada(s); as demais usam o padrão técnico."
                if legendas_manuais
                else "Nenhuma legenda manual importada; o relatório usará legendas técnicas padrão."
            ),
            "tipo": "info",
        },
    ]

    if manifesto.get("assinaturas_habilitadas", True):
        itens.append(
            {
                "ok": assinaturas_com_imagem >= 1,
                "titulo": "Assinaturas",
                "detalhe": f"{assinaturas_com_imagem} assinatura(s) com imagem e {assinaturas_com_nome} nome(s) preenchido(s).",
            }
        )
    else:
        itens.append(
            {
                "ok": True,
                "titulo": "Assinaturas",
                "detalhe": "Assinaturas desabilitadas no pacote offline.",
                "tipo": "info",
            }
        )

    return itens


def renderizar_checklist_pacote(manifesto: dict) -> None:
    linhas = []
    for item in itens_checklist_pacote(manifesto):
        status = "ok" if item.get("ok") else "warn"
        if item.get("tipo") == "info":
            status = "info"
        rotulo = "Pronto" if status == "ok" else "Info" if status == "info" else "Ajustar"
        linhas.append(
            "<div class='package-check'>"
            f"<span class='package-badge {status}'>{escape(rotulo)}</span>"
            "<div>"
            f"<div class='package-check-title'>{escape(item['titulo'])}</div>"
            f"<div class='package-check-detail'>{escape(item['detalhe'])}</div>"
            "</div>"
            "</div>"
        )
    st.markdown("<div class='package-checklist'>" + "".join(linhas) + "</div>", unsafe_allow_html=True)


def manifesto_tem_conteudo_para_gerar(manifesto: dict) -> bool:
    if limpar_texto(manifesto.get("observacoes", "")) or manifesto.get("audios"):
        return True
    if limpar_texto(manifesto.get("localizacao_maps", "")):
        return True
    if any(manifesto.get("cabecalho", {}).values()):
        return True
    if any(itens for itens in manifesto.get("evidencias", {}).values()):
        return True
    if manifesto.get("fotos_atendimento"):
        return True
    if any(limpar_texto(texto) for texto in manifesto.get("legendas_evidencias", {}).values()):
        return True
    if any(limpar_texto(item.get("nome", "")) for item in normalizar_assinaturas_lista(manifesto)):
        return True
    if any(item.get("imagem") for item in normalizar_assinaturas_lista(manifesto)):
        return True
    return False


def contexto_manifesto_para_geracao(manifesto: dict) -> str:
    linhas = []
    observacoes = limpar_texto(manifesto.get("observacoes", ""))
    if observacoes:
        linhas.append(observacoes)

    tipos = normalizar_tipos_atendimento(
        manifesto.get("tipos_atendimento") or manifesto.get("tipo_atendimento"),
        [],
    )
    if tipos:
        linhas.append(f"Tipo de Atendimento selecionado: {tipos_atendimento_para_texto(tipos)}.")
    tecnico_agres = normalizar_tecnico_agres_responsavel(manifesto.get("tecnico_agres_responsavel", ""))
    if tecnico_agres:
        linhas.append(f"Técnico Responsável Agres: {tecnico_agres}.")
    periodo_atendimento = formatar_periodo_atendimento(
        manifesto.get("data_atendimento_inicio", ""),
        manifesto.get("data_atendimento_final", ""),
    )
    if periodo_atendimento:
        linhas.append(f"Período do atendimento: {periodo_atendimento}.")
    data_atendimento = data_final_pacote(manifesto.get("data_atendimento_final", ""))
    if data_atendimento:
        linhas.append(f"Data final do atendimento: {datetime.strptime(data_atendimento, '%Y%m%d').strftime('%d/%m/%Y')}.")
    localizacao = limpar_texto(manifesto.get("localizacao_maps", ""))
    if localizacao:
        linhas.append(f"Localização informada no Maps: {localizacao}.")

    cabecalho_fotos = sum(1 for item in manifesto.get("cabecalho", {}).values() if item)
    if cabecalho_fotos:
        linhas.append(f"Fotos de cabeçalho do relatório: {cabecalho_fotos} imagem(ns) anexada(s).")

    for categoria, configuracao in CATEGORIAS_EVIDENCIAS.items():
        quantidade = len(manifesto.get("evidencias", {}).get(categoria, []) or [])
        legenda = limpar_texto(manifesto.get("legendas_evidencias", {}).get(categoria, ""))
        if quantidade:
            linhas.append(f"{configuracao['nome']}: {quantidade} foto(s) anexada(s).")
        if legenda:
            linhas.append(f"Descrição das fotos de {configuracao['nome']}: {legenda}.")

    fotos_extras = len(manifesto.get("fotos_atendimento", []) or [])
    if fotos_extras:
        linhas.append(f"Fotos adicionais do atendimento: {fotos_extras} foto(s) anexada(s) somente para o ZIP.")

    for assinatura in normalizar_assinaturas_lista(manifesto):
        responsavel = limpar_texto(assinatura.get("nome", ""))
        representa = limpar_texto(assinatura.get("representa", ""))
        if responsavel:
            linhas.append(f"{representa or 'Responsável'}: {responsavel}.")

    if linhas and not observacoes:
        linhas.insert(0, "Coleta offline importada sem complemento técnico descritivo. Utilizar os metadados abaixo para compor o relatório sem inventar procedimentos não informados.")

    return limpar_texto("\n".join(linhas))


def limpar_rascunho_atual() -> None:
    draft_dir = DRAFTS_DIR / st.session_state.draft_id
    if draft_dir.exists():
        shutil.rmtree(draft_dir)
    st.session_state.draft_id = uuid.uuid4().hex[:12]
    st.query_params["draft"] = st.session_state.draft_id
    for chave in (
        "tipo_atendimento",
        "tipos_atendimento",
        "localizacao_maps",
        "observacoes_texto",
        "fotos_atendimento_zip",
        "assinaturas_habilitadas",
        "quantidade_assinaturas",
        *[f"legenda_{categoria}" for categoria in CATEGORIAS_EVIDENCIAS],
        *[configuracao["campo"] for configuracao in ASSINATURAS_RESPONSAVEIS.values()],
        *[configuracao["campo_documento"] for configuracao in ASSINATURAS_RESPONSAVEIS.values()],
    ):
        st.session_state.pop(chave, None)
    for chave in list(st.session_state):
        if chave.startswith(("assinatura_nome_", "assinatura_representa_", "assinatura_documento_", "editar_assinatura_", "canvas_assinatura_")):
            st.session_state.pop(chave, None)
    st.session_state.relatorio_pronto = None
    st.session_state.nome_arquivo_pronto = None
    st.session_state.pacote_zip_pronto = None
    st.session_state.nome_pacote_zip_pronto = None


def finalizar_importacao_offline(manifesto: dict) -> None:
    aplicar_manifesto_na_sessao(manifesto)
    st.session_state.importacao_offline_ok = True
    st.session_state.importacao_offline_stats = manifesto.get("importacao_offline", {})
    st.session_state.json_import_status = "ok"
    st.session_state.json_import_message = "Pacote importado e validado."
    st.rerun()


def atualizar_status_upload_json(chave_uploader: str) -> None:
    arquivo = st.session_state.get(chave_uploader)
    if arquivo:
        conteudo = arquivo.getvalue()
        if len(conteudo) > MAX_PACKAGE_BYTES:
            st.session_state.pop("json_upload_bytes", None)
            st.session_state.json_import_status = "error"
            st.session_state.json_import_message = (
                f"O pacote excede o limite de {MAX_PACKAGE_BYTES // (1024 * 1024)} MB. "
                "Remova arquivos muito pesados e exporte novamente."
            )
            return
        st.session_state.json_upload_bytes = conteudo
        st.session_state.json_upload_name = getattr(arquivo, "name", "pacote_relatorio.json")
        st.session_state.json_upload_type = getattr(arquivo, "type", "application/json")
        st.session_state.json_import_status = "selected"
        st.session_state.json_import_message = f"{getattr(arquivo, 'name', 'arquivo.json')} selecionado. Clique em Importar pacote."
    else:
        st.session_state.pop("json_upload_bytes", None)
        st.session_state.pop("json_upload_name", None)
        st.session_state.pop("json_upload_type", None)
        st.session_state.pop("json_import_status", None)
        st.session_state.pop("json_import_message", None)


def arquivo_json_em_memoria():
    conteudo = st.session_state.get("json_upload_bytes")
    if not conteudo:
        return None
    arquivo = BytesIO(conteudo)
    arquivo.name = st.session_state.get("json_upload_name", "pacote_relatorio.json")
    arquivo.type = st.session_state.get("json_upload_type", "application/json")
    return arquivo


def remover_upload_json() -> None:
    st.session_state.pop("json_upload_bytes", None)
    st.session_state.pop("json_upload_name", None)
    st.session_state.pop("json_upload_type", None)
    st.session_state.pop("json_import_status", None)
    st.session_state.pop("json_import_message", None)
    st.session_state.json_uploader_version = int(st.session_state.get("json_uploader_version", 0)) + 1


def formatar_tamanho_arquivo(tamanho: int) -> str:
    valor = float(max(0, tamanho or 0))
    for unidade in ("B", "KB", "MB", "GB"):
        if valor < 1024 or unidade == "GB":
            casas = 0 if unidade == "B" else 1
            return f"{valor:.{casas}f} {unidade}"
        valor /= 1024
    return f"{valor:.1f} GB"


def renderizar_status_upload_json(arquivo) -> None:
    estado = st.session_state.get("json_import_status", "empty")
    if estado == "error":
        st.markdown(
            """
            <style>
                [data-testid="stFileUploaderFile"] svg,
                [data-testid="stFileUploaderFile"] svg *,
                [data-testid="stFileUploaderFile"] path {
                    color: #b42318 !important;
                    stroke: #b42318 !important;
                }
                [data-testid="stFileUploaderFile"] {
                    border-color: #f0a7a7 !important;
                    background: #fff8f8 !important;
                }
            </style>
            """,
            unsafe_allow_html=True,
        )
    elif estado in {"ok", "selected"}:
        st.markdown(
            """
            <style>
                [data-testid="stFileUploaderFile"] svg,
                [data-testid="stFileUploaderFile"] svg *,
                [data-testid="stFileUploaderFile"] path {
                    color: #176b43 !important;
                    stroke: #176b43 !important;
                }
                [data-testid="stFileUploaderFile"] {
                    border-color: #8bc5a5 !important;
                    background: #fbfffc !important;
                }
            </style>
            """,
            unsafe_allow_html=True,
        )
    nome_arquivo = escape(getattr(arquivo, "name", "") or "")
    mensagem = escape(st.session_state.get("json_import_message", "Nenhum pacote importado."))
    classes = {
        "ok": "json-import-card ok",
        "error": "json-import-card error",
        "selected": "json-import-card selected",
    }
    titulos = {
        "ok": "Pacote importado corretamente",
        "error": "Erro ao importar pacote",
        "selected": "Pacote selecionado",
    }
    icones = {"ok": "✓", "error": "!", "selected": "ZIP" if nome_arquivo.lower().endswith(".zip") else "JSON"}
    classe = classes.get(estado, "json-import-card")
    titulo = escape(titulos.get(estado, "Aguardando pacote"))
    icone = escape(icones.get(estado, "JSON"))
    detalhe = mensagem if estado in {"ok", "error", "selected"} else "Selecione o JSON ou ZIP exportado pelo modo offline."
    if estado == "selected" and nome_arquivo:
        tamanho = len(st.session_state.get("json_upload_bytes") or b"")
        detalhe = f"{nome_arquivo} · {formatar_tamanho_arquivo(tamanho)} · pronto para importar."
    st.markdown(
        f"""
        <div class="{classe}">
            <div class="json-file-icon">{icone}</div>
            <div>
                <div class="json-title">{titulo}</div>
                <div class="json-detail">{detalhe}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def linhas_metadados_preservando_vazios(texto: str) -> list[str]:
    if texto is None:
        return []
    return [linha.strip() for linha in str(texto).replace("\r\n", "\n").replace("\r", "\n").split("\n")]


def renderizar_editor_legendas_fotos(categoria: str, arquivos_upload, manifesto: dict) -> str:
    chave_estado = f"legenda_{categoria}"
    texto_existente = st.session_state.get(
        chave_estado,
        manifesto.get("legendas_evidencias", {}).get(categoria, ""),
    )
    linhas_existentes = linhas_metadados_preservando_vazios(texto_existente)
    arquivos_upload = arquivos_upload or []
    itens_salvos = manifesto.get("evidencias", {}).get(categoria, []) or []
    total = len(arquivos_upload) if arquivos_upload else len(itens_salvos)

    if total == 0:
        st.caption("Sem fotos nesta categoria.")
        return texto_existente or ""

    linhas_finais = []
    for indice in range(total):
        linha_existente = linhas_existentes[indice] if indice < len(linhas_existentes) else ""
        titulo_existente, legenda_existente, fonte_existente = separar_metadados_figura(linha_existente)
        manual_existente = bool(linha_existente and linha_existente != "||" and (titulo_existente or legenda_existente or fonte_existente))
        nome_foto = (
            getattr(arquivos_upload[indice], "name", "")
            if arquivos_upload
            else (itens_salvos[indice].get("name", "") if isinstance(itens_salvos[indice], dict) else "")
        )
        rotulo_foto = nome_foto or f"Foto {indice + 1}"

        with st.container(border=True):
            st.markdown(f"**{indice + 1}. {rotulo_foto}**")
            usar_manual = st.checkbox(
                "Usar legenda personalizada nesta foto",
                value=manual_existente,
                key=f"usar_legenda_{categoria}_{indice}",
            )
            if usar_manual:
                titulo = st.text_input(
                    "Título da figura",
                    value=titulo_existente,
                    placeholder=CATEGORIAS_EVIDENCIAS[categoria]["titulo_padrao"],
                    key=f"titulo_legenda_{categoria}_{indice}",
                )
                legenda = st.text_area(
                    "Legenda",
                    value=legenda_existente,
                    placeholder=CATEGORIAS_EVIDENCIAS[categoria]["legenda_padrao"],
                    height=70,
                    key=f"texto_legenda_{categoria}_{indice}",
                )
                linhas_finais.append(f"{limpar_texto(titulo)} | {limpar_texto(legenda)} | {limpar_texto(fonte_existente)}")
            else:
                linhas_finais.append("||")

    return "\n".join(linhas_finais)


# ==========================================
# 5. Interface visual
# ==========================================
draft_dir = pasta_rascunho_atual()
manifesto_rascunho = carregar_manifesto(draft_dir)

if "observacoes_texto" not in st.session_state:
    st.session_state.observacoes_texto = manifesto_rascunho.get("observacoes", "")
if "tipo_atendimento" not in st.session_state:
    tipos_manifesto = normalizar_tipos_atendimento(
        manifesto_rascunho.get("tipos_atendimento") or manifesto_rascunho.get("tipo_atendimento"),
        [],
    )
    st.session_state.tipos_atendimento = tipos_manifesto
    st.session_state.tipo_atendimento = tipos_manifesto[0] if tipos_manifesto else ""
if "localizacao_maps" not in st.session_state:
    st.session_state.localizacao_maps = manifesto_rascunho.get("localizacao_maps", "")

for categoria in CATEGORIAS_EVIDENCIAS:
    chave_legenda = f"legenda_{categoria}"
    if chave_legenda not in st.session_state:
        st.session_state[chave_legenda] = manifesto_rascunho.get("legendas_evidencias", {}).get(categoria, "")

for chave, configuracao in ASSINATURAS_RESPONSAVEIS.items():
    campo = configuracao["campo"]
    if campo not in st.session_state:
        st.session_state[campo] = manifesto_rascunho.get("responsaveis", {}).get(chave, "")

if "assinaturas_habilitadas" not in st.session_state:
    st.session_state.assinaturas_habilitadas = manifesto_rascunho.get("assinaturas_habilitadas", True)
if "quantidade_assinaturas" not in st.session_state:
    st.session_state.quantidade_assinaturas = quantidade_assinaturas_normalizada(manifesto_rascunho.get("quantidade_assinaturas", MIN_ASSINATURAS))

for assinatura in normalizar_assinaturas_lista(manifesto_rascunho):
    assinatura_id = assinatura["id"]
    for campo, valor in (
        ("nome", assinatura.get("nome", "")),
        ("representa", assinatura.get("representa", "")),
    ):
        chave_estado = f"assinatura_{campo}_{assinatura_id}"
        if chave_estado not in st.session_state:
            st.session_state[chave_estado] = valor

logo_uri = imagem_data_uri(LOGO_PATH)
logo_html = f'<img class="brand-logo" src="{logo_uri}" alt="Agres">' if logo_uri else ""
st.markdown(
    f"""
    <section class="brand-hero">
        {logo_html}
        <div class="brand-copy">
            <div class="brand-title">Relatórios Técnicos Agres</div>
        </div>
    </section>
    """,
    unsafe_allow_html=True,
)

if st.session_state.pop("importacao_offline_ok", False):
    estatisticas_importacao = st.session_state.pop("importacao_offline_stats", {}) or {}
    audios_recebidos = int(estatisticas_importacao.get("audios_recebidos") or 0)
    audios_salvos = int(estatisticas_importacao.get("audios_salvos") or 0)
    imagens_recebidas = int(estatisticas_importacao.get("imagens_recebidas") or 0)
    imagens_salvas = int(estatisticas_importacao.get("imagens_salvas") or 0)
    if audios_recebidos:
        st.success(f"Pacote offline sincronizado. Áudios importados: {audios_salvos}/{audios_recebidos}.")
    else:
        st.success("Pacote offline sincronizado.")
    if imagens_recebidas:
        st.caption(f"Imagens importadas: {imagens_salvas}/{imagens_recebidas}.")
    if audios_recebidos and audios_salvos == 0:
        st.warning(
            "O pacote informa áudio, mas nenhum arquivo de áudio pôde ser salvo neste navegador. "
            "O relatório ainda pode ser gerado usando o texto e os dados coletados; para usar a fala, exporte o pacote novamente pelo iPad."
        )
    if imagens_recebidas and imagens_salvas < imagens_recebidas:
        st.warning(
            "Algumas imagens do pacote offline não puderam ser importadas. "
            "O relatório será gerado com as imagens válidas; reexporte o pacote se precisar recuperar alguma foto."
        )
    for mensagem in estatisticas_importacao.get("erros_audio", []) or []:
        st.caption(mensagem)
    for mensagem in estatisticas_importacao.get("erros_imagem", []) or []:
        st.caption(mensagem)

with st.container(border=True):
    st.markdown(
        "<p class='section-title'>1. Importar coleta offline</p>",
        unsafe_allow_html=True,
    )
    st.markdown(
        "<p class='section-caption'>Importe o pacote JSON ou ZIP para gerar o relatório técnico.</p>",
        unsafe_allow_html=True,
    )
    pacote_offline = arquivo_json_em_memoria()
    if pacote_offline is None:
        versao_uploader = int(st.session_state.get("json_uploader_version", 0))
        chave_uploader = f"pacote_offline_{versao_uploader}"
        st.file_uploader(
            "Selecione o pacote JSON ou ZIP",
            type=["json", "zip"],
            key=chave_uploader,
            on_change=atualizar_status_upload_json,
            args=(chave_uploader,),
        )
        if st.session_state.get("json_import_status") == "error":
            st.error(st.session_state.get("json_import_message", "Não foi possível carregar o pacote."))
    else:
        renderizar_status_upload_json(pacote_offline)
        coluna_importar, coluna_remover = st.columns([1.25, 0.75])
        importar_json = coluna_importar.button(
            "Importar pacote",
            type="primary",
            use_container_width=True,
            key="btn_importar_pacote_arquivo",
        )
        coluna_remover.button(
            "✕ Remover pacote",
            use_container_width=True,
            key="btn_remover_pacote_arquivo",
            on_click=remover_upload_json,
        )

        if importar_json:
            try:
                manifesto_rascunho = importar_pacote_offline(pacote_offline, draft_dir, manifesto_rascunho)
                finalizar_importacao_offline(manifesto_rascunho)
            except Exception as erro:
                st.session_state.json_import_status = "error"
                st.session_state.json_import_message = f"Não foi possível importar: {erro}"
                st.error(f"Erro ao sincronizar pacote offline: {erro}")

with st.container(border=True):
    st.markdown(
        "<p class='section-title'>2. Pacote importado</p>",
        unsafe_allow_html=True,
    )
    total_audios = len(manifesto_rascunho.get("audios", []) or [])
    total_cabecalho = sum(1 for item in manifesto_rascunho.get("cabecalho", {}).values() if item)
    total_evidencias = contar_evidencias(manifesto_rascunho)
    total_fotos_zip = len(manifesto_rascunho.get("fotos_atendimento", []) or [])
    total_assinaturas = sum(1 for item in normalizar_assinaturas_lista(manifesto_rascunho) if item.get("imagem"))

    col_resumo_1, col_resumo_2, col_resumo_3, col_resumo_4 = st.columns(4)
    col_resumo_1.metric("Áudios", total_audios)
    col_resumo_2.metric("Fotos do Word", total_evidencias)
    col_resumo_3.metric("Fotos do ZIP", total_cabecalho + total_evidencias + total_fotos_zip + total_assinaturas)
    col_resumo_4.metric("Assinaturas", total_assinaturas)

    renderizar_checklist_pacote(manifesto_rascunho)

    tipos_atendimento_salvos = normalizar_tipos_atendimento(
        manifesto_rascunho.get("tipos_atendimento") or manifesto_rascunho.get("tipo_atendimento"),
        [],
    )
    localizacao_maps_salva = manifesto_rascunho.get("localizacao_maps", "")
    tecnico_agres_salvo = normalizar_tecnico_agres_responsavel(
        manifesto_rascunho.get("tecnico_agres_responsavel", "")
    )
    st.caption(f"Tipo de Atendimento: {tipos_atendimento_para_texto(tipos_atendimento_salvos)}")
    if tecnico_agres_salvo:
        st.caption(f"Técnico Responsável Agres: {tecnico_agres_salvo}")
    periodo_salvo = formatar_periodo_atendimento(
        manifesto_rascunho.get("data_atendimento_inicio", ""),
        manifesto_rascunho.get("data_atendimento_final", ""),
    )
    if periodo_salvo:
        st.caption(f"Período do atendimento: {periodo_salvo}")
    if localizacao_maps_salva:
        st.caption(f"Localização: {localizacao_maps_salva}")
    importacao_offline = manifesto_rascunho.get("importacao_offline", {}) or {}
    if importacao_offline.get("exported_at") or importacao_offline.get("app_version"):
        partes_importacao = []
        if importacao_offline.get("exported_at"):
            partes_importacao.append(f"exportado em {importacao_offline['exported_at']}")
        if importacao_offline.get("app_version"):
            partes_importacao.append(f"versão offline {importacao_offline['app_version']}")
        st.caption("Pacote " + " | ".join(partes_importacao))

    observacoes_texto = st.text_area(
        "Ajuste opcional antes de gerar",
        height=110,
        key="observacoes_texto",
    )
    if limpar_texto(observacoes_texto) != limpar_texto(manifesto_rascunho.get("observacoes", "")):
        manifesto_rascunho["observacoes"] = observacoes_texto or ""
        salvar_manifesto(draft_dir, manifesto_rascunho)

    if st.button("Limpar pacote importado", use_container_width=True):
        st.session_state.pop("json_import_status", None)
        st.session_state.pop("json_import_message", None)
        limpar_rascunho_atual()
        st.rerun()


# ==========================================
# 6. Execução
# ==========================================
caminhos_audio_salvos, caminhos_cabecalho_salvos, evidencias_salvas, fotos_atendimento_salvas, assinaturas_salvas = caminhos_salvos_rascunho(
    draft_dir,
    manifesto_rascunho,
)
contexto_manifesto_salvo = contexto_manifesto_para_geracao(manifesto_rascunho) if manifesto_tem_conteudo_para_gerar(manifesto_rascunho) else ""
observacoes_salvas = (
    limpar_texto(observacoes_texto)
    or manifesto_rascunho.get("observacoes", "")
    or contexto_manifesto_salvo
)
legendas_salvas = {
    categoria: manifesto_rascunho.get("legendas_evidencias", {}).get(categoria, "")
    for categoria in CATEGORIAS_EVIDENCIAS
}
responsaveis_salvos = {
    chave: manifesto_rascunho.get("responsaveis", {}).get(chave, "")
    for chave in ASSINATURAS_RESPONSAVEIS
}
documentos_salvos = {
    chave: manifesto_rascunho.get("documentos", {}).get(chave, "")
    for chave in ASSINATURAS_RESPONSAVEIS
}
tipos_atendimento_salvos = normalizar_tipos_atendimento(
    manifesto_rascunho.get("tipos_atendimento") or manifesto_rascunho.get("tipo_atendimento"),
    [],
)
localizacao_maps_salva = manifesto_rascunho.get("localizacao_maps", "")
tecnico_agres_salvo = normalizar_tecnico_agres_responsavel(
    manifesto_rascunho.get("tecnico_agres_responsavel", "")
)
data_final_salva = data_final_pacote(manifesto_rascunho.get("data_atendimento_final", ""))

entrada_disponivel = (
    bool(caminhos_audio_salvos)
    or bool(limpar_texto(observacoes_salvas))
    or manifesto_tem_conteudo_para_gerar(manifesto_rascunho)
)
pronto_para_geracao = (
    entrada_disponivel
    and bool(tipos_atendimento_salvos)
    and bool(tecnico_agres_salvo)
    and bool(data_final_salva)
)

with st.container(border=True):
    st.markdown(
        "<p class='section-title'>3. Gerar relatório</p>",
        unsafe_allow_html=True,
    )

    if pronto_para_geracao:
        st.markdown(
            """
            <div class="generate-callout">
                <div class="generate-callout-title">Pacote pronto para geração</div>
                <div class="generate-callout-detail">Clique para gerar o Word e o ZIP com as fotos do atendimento.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    elif entrada_disponivel:
        if not tipos_atendimento_salvos:
            st.warning("Selecione o Tipo de Atendimento na coleta offline e exporte novamente o pacote.")
        elif not tecnico_agres_salvo:
            st.warning("Selecione o Técnico Responsável Agres na coleta offline e exporte novamente o pacote.")
        elif not data_final_salva:
            st.warning("Informe a Data Final do Atendimento na coleta offline e exporte novamente o pacote.")

    if st.button(
        "Gerar relatório técnico agora",
        type="primary",
        use_container_width=True,
        disabled=not pronto_para_geracao,
    ):
        st.session_state.relatorio_pronto = None
        st.session_state.pacote_zip_pronto = None
        st.session_state.nome_arquivo_pronto = None
        st.session_state.nome_pacote_zip_pronto = None

        try:
            with tempfile.TemporaryDirectory() as pasta_temp_raw:
                pasta_temp = Path(pasta_temp_raw)

                with st.status("Gerando relatório...", expanded=False) as status:
                    imagens_equipamentos_ia = selecionar_imagens_equipamentos_agres(
                        caminhos_cabecalho_salvos,
                        evidencias_salvas,
                    )
                    dados = processar_atendimento_completo(
                        caminhos_audio_salvos,
                        observacoes_salvas,
                        imagens_equipamentos_ia,
                    )
                    dados["localizacao_maps"] = localizacao_maps_salva or dados.get("localizacao_maps", "")
                    dados["contexto_coleta"] = observacoes_salvas
                    periodo_atendimento = formatar_periodo_atendimento(
                        manifesto_rascunho.get("data_atendimento_inicio", ""),
                        manifesto_rascunho.get("data_atendimento_final", ""),
                    )
                    if periodo_atendimento:
                        dados["data_visita"] = periodo_atendimento
                    dados["data_atendimento_final"] = manifesto_rascunho.get("data_atendimento_final", "")
                    if tecnico_agres_salvo:
                        dados["tecnico_agres_responsavel"] = tecnico_agres_salvo
                        dados["tecnicos"] = tecnico_agres_salvo
                    aplicar_tipos_atendimento(dados, tipos_atendimento_salvos)
                    for chave, configuracao in ASSINATURAS_RESPONSAVEIS.items():
                        responsavel = limpar_texto(responsaveis_salvos.get(chave, ""))
                        if responsavel:
                            dados[configuracao["campo"]] = responsavel

                    arquivo_final = gerar_docx(
                        dados,
                        evidencias_salvas,
                        caminhos_cabecalho_salvos,
                        pasta_temp,
                        legendas_salvas,
                        assinaturas_salvas,
                    )
                    pacote_final = gerar_pacote_relatorio(
                        arquivo_final,
                        evidencias_salvas,
                        fotos_atendimento_salvas,
                        caminhos_cabecalho_salvos,
                        pasta_temp,
                        legendas_salvas,
                        assinaturas_salvas,
                    )
                    st.session_state.relatorio_pronto = arquivo_final.read_bytes()
                    st.session_state.nome_arquivo_pronto = arquivo_final.name
                    st.session_state.pacote_zip_pronto = pacote_final.read_bytes()
                    st.session_state.nome_pacote_zip_pronto = pacote_final.name

                    status.update(label="Relatório finalizado!", state="complete", expanded=False)

        except Exception as erro:
            st.error(f"Erro no processamento: {erro}")
            st.exception(erro)
    elif not entrada_disponivel:
        st.caption("Importe o pacote para liberar a geração do relatório técnico.")

    if st.session_state.relatorio_pronto:
        st.success("✅ O laudo está pronto para download!")
        if st.session_state.pacote_zip_pronto:
            st.download_button(
                label="Baixar Word + fotos (.zip)",
                data=st.session_state.pacote_zip_pronto,
                file_name=st.session_state.nome_pacote_zip_pronto,
                mime="application/zip",
                type="primary",
                use_container_width=True,
            )
        st.download_button(
            label="Baixar somente Word",
            data=st.session_state.relatorio_pronto,
            file_name=st.session_state.nome_arquivo_pronto,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
